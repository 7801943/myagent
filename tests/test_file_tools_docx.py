import asyncio

from docx import Document

from myagent.tools.api import generate_schema
from myagent.tools.builtin.file_edit import file_edit, file_edit_table
from myagent.tools.builtin.file_read import file_read
from myagent.tools.builtin.file_write import file_write


def run_tool(coro):
    return asyncio.run(coro)


def test_file_tool_path_schema_allows_workspace_paths():
    for fn in (file_read, file_write, file_edit, file_edit_table):
        description = generate_schema(fn)["properties"]["path"]["description"]
        assert "绝对路径" in description
        assert "workspace 可见路径" in description
        assert "工具层解析" in description


def test_file_read_docx_outputs_paragraphs_and_tables_in_body_order(tmp_path):
    path = tmp_path / "ordered.docx"
    doc = Document()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"
    doc.add_paragraph("After table")
    doc.save(path)

    result = run_tool(file_read(str(path)))

    assert not result.is_error, result.content
    assert result.metadata["format"] == "docx"
    before = result.content.index("Before table")
    table_header = result.content.index("[表格 1]")
    row_1 = result.content.index("| Name | Value |")
    row_2 = result.content.index("| Alpha | 42 |")
    after = result.content.index("After table")
    assert before < table_header < row_1 < row_2 < after


def test_file_read_docx_table_only_document(tmp_path):
    path = tmp_path / "table_only.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Left"
    table.cell(0, 1).text = "Right"
    doc.save(path)

    result = run_tool(file_read(str(path)))

    assert not result.is_error, result.content
    assert "[表格 1]" in result.content
    assert "| Left | Right |" in result.content


def test_file_read_docx_line_range_can_slice_table_output(tmp_path):
    path = tmp_path / "slice.docx"
    doc = Document()
    doc.add_paragraph("Intro")
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "First"
    table.cell(1, 0).text = "Second"
    doc.add_paragraph("Outro")
    doc.save(path)

    result = run_tool(file_read(str(path), start_line_or_page=2, end_line_or_page=4))

    assert not result.is_error, result.content
    assert "Intro" not in result.content
    assert "2 | [表格 1]" in result.content
    assert "3 | | First |" in result.content
    assert "4 | | Second |" in result.content
    assert "Outro" not in result.content
