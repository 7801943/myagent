"""
FileRead / FileWrite 工具函数。
通过 JsonRpcProxy -> ExecutionEngine.execute_function() 执行。

支持格式：文本(txt/md/py/json/yaml/log 等)、CSV、PDF、XLSX/XLS、DOCX、图片
"""
import base64
import csv
import hashlib
import io
import json
import logging
import mimetypes
import os
import re
from copy import copy, deepcopy
from datetime import date, datetime
from pathlib import Path
from typing import Any

from myagent.tools.api import tool, ToolResult

logger = logging.getLogger(__name__)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 安全策略
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_DENIED_PATHS = {
    "/etc", "/root",
    "/sys", "/proc/sys", "/boot", "/dev",
}

_MAX_BASE64_BYTES = 20 * 1024 * 1024  # 20MB base64 输出上限
_PDF_DEFAULT_MAX_PAGES = 3            # 未指定页码范围时，默认最多渲染页数
_PDF_SCAN_FALLBACK_RATIO = 0.5        # 无文本页占比超过此阈值时，自动回退到 base64 渲染


def _check_path_safety(path: str) -> str | None:
    """检查路径是否被安全策略禁止。返回错误信息或 None。"""
    try:
        resolved = str(Path(path).resolve())
    except Exception:
        return f"路径解析失败: {path}"
    for denied_path in _DENIED_PATHS:
        if resolved.startswith(denied_path):
            return f"路径被安全策略禁止: {path} (匹配: {denied_path})"
    return None


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 文件类型检测
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_TEXT_EXTENSIONS = {
    ".txt", ".md", ".json", ".yaml", ".yml", ".toml",
    ".xml", ".html", ".css",
    ".ini", ".cfg", ".conf", ".env", ".gitignore",
    ".log", ".csv", ".tsv",
}

_SPREADSHEET_EXTENSIONS = {".xlsx", ".xls"}
_DOC_EXTENSIONS = {".docx", ".doc"}
_PDF_EXTENSIONS = {".pdf"}

_IMAGE_EXTENSIONS = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".svg": "image/svg+xml",
    ".ico": "image/x-icon",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
}


def _detect_file_type(path: Path) -> str:
    """
    根据扩展名检测文件类型。
    返回: "text" | "csv" | "xlsx" | "docx" | "pdf" | "image" | "binary"
    """
    ext = path.suffix.lower()

    if ext == ".csv":
        return "csv"
    if ext == ".tsv":
        return "csv"
    if ext in _SPREADSHEET_EXTENSIONS:
        return "xlsx"
    if ext in _DOC_EXTENSIONS:
        return "docx"
    if ext in _PDF_EXTENSIONS:
        return "pdf"
    if ext in _IMAGE_EXTENSIONS:
        return "image"
    if ext in _TEXT_EXTENSIONS:
        return "text"

    # 无扩展名或未知扩展名：尝试检测是否为文本
    if not ext:
        try:
            mime, _ = mimetypes.guess_type(str(path))
            if mime and mime.startswith("text/"):
                return "text"
        except Exception:
            pass

    return "binary"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 编码检测
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _detect_encoding(path: Path, hint: str | None = None) -> str:
    """检测文件编码。hint 优先，然后尝试 chardet，最后回退 utf-8。"""
    if hint:
        return hint

    # 尝试读取前 8KB 用于检测
    try:
        raw = path.open("rb").read(8192)
    except Exception:
        return "utf-8"

    # 检测 BOM
    if raw.startswith(b"\xef\xbb\xbf"):
        return "utf-8-sig"
    if raw.startswith(b"\xff\xfe"):
        return "utf-16-le"
    if raw.startswith(b"\xfe\xff"):
        return "utf-16-be"

    # 检测 null 字节（二进制文件）
    if b"\x00" in raw[:1024]:
        return "binary"

    # 尝试 utf-8
    try:
        raw.decode("utf-8")
        return "utf-8"
    except UnicodeDecodeError:
        pass

    # chardet 检测
    try:
        import chardet
        result = chardet.detect(raw)
        if result and result.get("encoding"):
            return result["encoding"]
    except ImportError:
        pass

    # 回退 latin-1（永远不会失败）
    return "latin-1"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 文本行号格式化
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _format_lines(
    lines: list[str],
    start_line: int,
    end_line: int | None,
    total_lines: int | None = None,
) -> tuple[str, dict[str, Any]]:
    """
    格式化行列表，始终添加行号。
    返回 (格式化后的文本, 元数据)。
    """
    display_total = total_lines or len(lines)
    width = len(str(display_total))

    output_parts: list[str] = []
    actual_count = 0

    for i, line in enumerate(lines):
        line_num = start_line + i
        if end_line is not None and line_num > end_line:
            break

        # 确保行以换行结尾
        if not line.endswith("\n"):
            line += "\n"

        output_parts.append(f"{line_num:>{width}} | {line}")
        actual_count += 1

    content = "".join(output_parts)

    meta = {
        "lines_output": actual_count,
        "start_line": start_line,
        "total_lines": total_lines,
    }

    return content, meta


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 各格式解析器
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

async def _read_text(
    path: Path,
    start_line: int | None,
    end_line: int | None,
    encoding: str | None,
) -> ToolResult:
    """读取纯文本文件。"""
    enc = _detect_encoding(path, encoding)
    if enc == "binary":
        return ToolResult(
            content=f"文件 {path.name} 似乎是二进制文件，请使用 output_format=\"base64\" 读取。",
            is_error=True,
        )

    try:
        with open(path, "r", encoding=enc, errors="replace") as f:
            all_lines = f.readlines()
    except PermissionError:
        return ToolResult(content=f"权限不足，无法读取: {path}", is_error=True)
    except Exception as e:
        return ToolResult(content=f"读取文件失败: {e}", is_error=True)

    total = len(all_lines)
    s = max(1, start_line or 1)
    e = min(end_line or total, total) if end_line else total

    if s > total:
        return ToolResult(
            content=f"文件共 {total} 行，start_line={s} 超出范围。",
            is_error=True,
            metadata={"total_lines": total},
        )

    sliced = all_lines[s - 1:e]
    content, meta = _format_lines(sliced, s, e, total)
    meta["path"] = str(path.resolve())
    meta["encoding"] = enc

    logger.info("file_read 纯文本完成: %s, 返回行数=%d", path, meta["lines_output"])
    return ToolResult(content=content, metadata=meta)


async def _read_csv(
    path: Path,
    start_line: int | None,
    end_line: int | None,
    encoding: str | None,
) -> ToolResult:
    """读取 CSV/TSV 文件。使用 csv 模块解析，按行输出。"""
    enc = _detect_encoding(path, encoding)
    if enc == "binary":
        enc = "utf-8"

    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","

    try:
        with open(path, "r", encoding=enc, errors="replace", newline="") as f:
            # 使用 csv.reader 处理各种引号和转义
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
    except PermissionError:
        return ToolResult(content=f"权限不足，无法读取: {path}", is_error=True)
    except Exception as e:
        return ToolResult(content=f"读取 CSV 失败: {e}", is_error=True)

    if not rows:
        return ToolResult(content="CSV 文件为空。", metadata={"path": str(path.resolve())})

    total = len(rows)
    s = max(1, start_line or 1)
    e = min(end_line or total, total) if end_line else total

    if s > total:
        return ToolResult(
            content=f"文件共 {total} 行，start_line={s} 超出范围。",
            is_error=True,
            metadata={"total_lines": total},
        )

    # 格式化每行为 | 分隔的列
    formatted_lines: list[str] = []
    for row in rows[s - 1:e]:
        formatted_lines.append(" | ".join(str(cell) for cell in row) + "\n")

    content, meta = _format_lines(formatted_lines, s, e, total)
    meta["path"] = str(path.resolve())
    meta["encoding"] = enc
    meta["format"] = "csv"

    logger.info("file_read CSV完成: %s, 返回行数=%d", path, meta["lines_output"])
    return ToolResult(content=content, metadata=meta)


async def _read_xlsx(
    path: Path,
    sheet_name: str | None,
    start_line: int | None,
    end_line: int | None,
    xlsx_range: str | None = None,
    render_mode: str = "values",
    row_mode: str = "text",
    include_tables: bool = True,
    include_merges: bool = True,
) -> ToolResult:
    """读取 XLSX/XLS 文件。多 Sheet 时先列出 Sheet 列表。"""
    try:
        import openpyxl
    except ImportError:
        return ToolResult(
            content="缺少 openpyxl 库，无法解析 XLSX。请安装: pip install openpyxl",
            is_error=True,
        )

    render_mode = (render_mode or "values").lower()
    row_mode = (row_mode or "text").lower()
    if render_mode not in ("values", "formulas", "both"):
        return ToolResult(
            content=f"不支持的 render_mode: {render_mode}。可选: values, formulas, both",
            is_error=True,
        )
    if row_mode not in ("text", "arrays", "objects"):
        return ToolResult(
            content=f"不支持的 row_mode: {row_mode}。可选: text, arrays, objects",
            is_error=True,
        )

    try:
        if render_mode == "both":
            wb = openpyxl.load_workbook(str(path), read_only=False, data_only=False)
            value_wb = openpyxl.load_workbook(str(path), read_only=False, data_only=True)
        else:
            wb = openpyxl.load_workbook(
                str(path), read_only=False, data_only=(render_mode == "values")
            )
            value_wb = None
    except Exception as e:
        return ToolResult(content=f"打开 XLSX 失败: {e}", is_error=True)

    try:
        sheet_names = wb.sheetnames

        # ── 未指定 sheet_name：列出所有工作表 ──
        if sheet_name is None:
            if len(sheet_names) == 1:
                ws = wb[sheet_names[0]]
                value_ws = value_wb[sheet_names[0]] if value_wb else None
                return await _format_xlsx_sheet(
                    ws, sheet_names[0], wb, path, start_line, end_line,
                    xlsx_range=xlsx_range, render_mode=render_mode,
                    row_mode=row_mode, value_ws=value_ws,
                    include_tables=include_tables, include_merges=include_merges,
                )

            # 多个 Sheet：列出信息
            lines = [f"文件: {path.name} (XLSX 工作簿)\n"]
            lines.append(f"包含 {len(sheet_names)} 个工作表：\n")
            for i, name in enumerate(sheet_names, 1):
                ws = wb[name]
                row_count = ws.max_row or 0
                col_count = ws.max_column or 0
                table_count = len(_worksheet_tables(ws))
                table_hint = f", {table_count} 个表格" if table_count else ""
                lines.append(f"  [{i}] \"{name}\" ({row_count} 行 × {col_count} 列{table_hint})\n")
            lines.append(
                "\n请指定 sheet_name 参数选择要读取的工作表。\n"
                f"示例: file_read(path=\"{path}\", sheet_name=\"{sheet_names[0]}\")\n"
            )

            profile = _profile_workbook(wb)
            wb.close()
            if value_wb:
                value_wb.close()
            return ToolResult(
                content="".join(lines),
                metadata={
                    "path": str(path.resolve()),
                    "format": "xlsx",
                    "sheet_count": len(sheet_names),
                    "sheet_names": sheet_names,
                    "workbook": profile,
                    "structure_token": _hash_json(profile),
                },
            )

        # ── 指定了 sheet_name ──
        if sheet_name not in sheet_names:
            wb.close()
            if value_wb:
                value_wb.close()
            return ToolResult(
                content=(
                    f"工作表 \"{sheet_name}\" 不存在。\n"
                    f"可用工作表: {', '.join(repr(n) for n in sheet_names)}"
                ),
                is_error=True,
                metadata={"sheet_names": sheet_names},
            )

        ws = wb[sheet_name]
        value_ws = value_wb[sheet_name] if value_wb else None
        return await _format_xlsx_sheet(
            ws, sheet_name, wb, path, start_line, end_line,
            xlsx_range=xlsx_range, render_mode=render_mode,
            row_mode=row_mode, value_ws=value_ws,
            include_tables=include_tables, include_merges=include_merges,
        )
    except Exception as e:
        wb.close()
        if 'value_wb' in locals() and value_wb:
            value_wb.close()
        return ToolResult(content=f"读取 XLSX 失败: {e}", is_error=True)


async def _format_xlsx_sheet(
    ws, sheet_name: str, wb, path: Path,
    start_line: int | None, end_line: int | None,
    xlsx_range: str | None = None,
    render_mode: str = "values",
    row_mode: str = "text",
    value_ws=None,
    include_tables: bool = True,
    include_merges: bool = True,
) -> ToolResult:
    """格式化一个 XLSX 工作表的内容。"""
    total = ws.max_row or 0
    col_count = ws.max_column or 0

    if total == 0:
        metadata = {"path": str(path.resolve()), "format": "xlsx", "sheet": sheet_name}
        metadata["workbook"] = _profile_workbook(wb)
        metadata["structure_token"] = _hash_json(metadata["workbook"])
        wb.close()
        if value_ws is not None:
            value_ws.parent.close()
        return ToolResult(content=f"工作表 \"{sheet_name}\" 为空。", metadata=metadata)

    if xlsx_range:
        min_col, min_row, max_col, max_row = _parse_a1_range(xlsx_range)
        s, e = min_row, max_row
    else:
        s = max(1, start_line or 1)
        e = min(end_line or total, total) if end_line else total
        min_col, min_row, max_col, max_row = 1, s, col_count, e

    if s > total:
        wb.close()
        if value_ws is not None:
            value_ws.parent.close()
        return ToolResult(
            content=f"工作表共 {total} 行，start_line={s} 超出范围。",
            is_error=True,
            metadata={"total_lines": total, "sheet": sheet_name},
        )

    display_lines: list[str] = []
    rows: list[list[Any]] = []
    for row_idx in range(min_row, max_row + 1):
        display_cells: list[str] = []
        row_values: list[Any] = []
        for col_idx in range(min_col, max_col + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            value_cell = value_ws.cell(row=row_idx, column=col_idx) if value_ws else None
            payload = _xlsx_cell_payload(cell, value_cell, render_mode)
            row_values.append(payload)
            display_cells.append(_xlsx_cell_display(payload))
        rows.append(row_values)
        display_lines.append(" | ".join(display_cells) + "\n")

    content, meta = _format_lines(display_lines, min_row, max_row, total)
    selected_range = _range_from_bounds(min_col, min_row, max_col, max_row)
    profile = _profile_workbook(wb)
    sheet_profile = next(
        (sheet for sheet in profile.get("sheets", []) if sheet.get("name") == sheet_name),
        {},
    )

    metadata_rows: Any = rows
    headers: list[str] | None = None
    if row_mode == "objects" and rows:
        headers = [str(_metadata_scalar(v) or "") for v in rows[0]]
        metadata_rows = []
        for row in rows[1:]:
            metadata_rows.append({headers[i]: _metadata_scalar(row[i]) for i in range(len(headers)) if headers[i]})
    elif row_mode == "arrays":
        metadata_rows = [[_metadata_scalar(v) for v in row] for row in rows]

    meta.update({
        "path": str(path.resolve()),
        "format": "xlsx",
        "sheet": sheet_name,
        "columns": col_count,
        "range": selected_range,
        "render_mode": render_mode,
        "row_mode": row_mode,
        "rows": metadata_rows,
        "workbook": profile,
        "structure_token": _hash_json(profile),
        "content_token": _content_token(ws, (min_col, min_row, max_col, max_row)),
    })
    if headers is not None:
        meta["headers"] = headers
    if include_tables:
        meta["tables"] = sheet_profile.get("tables", [])
    if include_merges:
        meta["merged_ranges"] = sheet_profile.get("merged_ranges", [])

    wb.close()
    if value_ws is not None:
        value_ws.parent.close()

    logger.info("file_read XLSX完成: %s, sheet=%s, 返回行数=%d", path, sheet_name, meta["lines_output"])
    return ToolResult(content=content, metadata=meta)


async def _read_pdf_text(
    path: Path,
    start_line: int | None,
    end_line: int | None,
) -> ToolResult:
    """使用 pdfplumber 提取 PDF 文本内容。"""
    # 抑制 pdfminer 的 DEBUG 日志洪水。
    # pdfminer 在解析 PDF 时会为每个 token/操作输出大量 DEBUG 日志
    # （nexttoken / add_results / nextobject / do_keyword / exec 等），
    # 一个普通 PDF 可能产生数千~数万条日志，导致：
    #   1. 子进程 stderr 被海量日志填满，产生严重 I/O 瓶颈
    #   2. 主进程 _drain_stderr 逐行读取并以 INFO 级别转发，进一步加剧延迟
    #   3. PDF 解析看似"卡住"，实际大部分时间花在日志 I/O 上
    # 将 pdfminer 相关 logger 提升到 WARNING 级别，彻底消除底层噪音。
    for _logger_name in ("pdfminer", "pdfminer.psparser", "pdfminer.pdfinterp",
                         "pdfminer.pdfpage", "pdfminer.converter",
                         "pdfminer.layout", "pdfminer.utils"):
        logging.getLogger(_logger_name).setLevel(logging.WARNING)

    try:
        import pdfplumber
    except ImportError:
        return ToolResult(
            content="缺少 pdfplumber 库，无法解析 PDF。请安装: pip install pdfplumber",
            is_error=True,
        )

    try:
        all_lines: list[str] = []
        page_count = 0

        with pdfplumber.open(str(path)) as pdf:
            page_count = len(pdf.pages)
            start_page = max(1, start_line or 1)
            end_page = min(end_line or page_count, page_count) if end_line else page_count

            if start_page > page_count:
                return ToolResult(
                    content=f"PDF 共 {page_count} 页，start_line={start_page} 超出范围。",
                    is_error=True,
                    metadata={"page_count": page_count},
                )

            for page_num in range(start_page, end_page + 1):
                page = pdf.pages[page_num - 1]
                text = page.extract_text() or ""
                # 添加页码标记
                all_lines.append(f"--- 第 {page_num} 页 ---\n")
                if text.strip():
                    for line in text.split("\n"):
                        all_lines.append(line + "\n")
                else:
                    all_lines.append("[此页无可提取文本，可能是扫描图片]\n")

        if not all_lines:
            return ToolResult(
                content="PDF 未能提取到任何文本。可能是扫描件/图片 PDF，"
                        "请使用 output_format=\"base64\" 以图片形式读取。",
                metadata={"page_count": page_count},
            )

        # ── 扫描页检测：如果大部分页无文本，自动回退到 base64 渲染 ──
        pages_read_count = end_page - start_page + 1
        empty_page_count = sum(
            1 for line in all_lines
            if line.strip() == "[此页无可提取文本，可能是扫描图片]"
        )
        # 每个 empty page 产生 2 行（页码标题 + 提示），实际空页数 = empty_page_count
        if (
            pages_read_count > 0
            and empty_page_count / pages_read_count > _PDF_SCAN_FALLBACK_RATIO
        ):
            logger.debug(
                "PDF 扫描页回退: 读取 %d 页中 %d 页无文本(%.0f%% > 阈值 %.0f%%)，回退到 base64 渲染",
                pages_read_count, empty_page_count,
                empty_page_count / pages_read_count * 100,
                _PDF_SCAN_FALLBACK_RATIO * 100,
            )
            return await _read_pdf_base64(path, start_line, end_line)

        # 把所有行视为一个整体，用页码范围代替行号范围
        total = len(all_lines)
        s = 1  # 已经只取了目标页的内容

        content, meta = _format_lines(all_lines, s, None, total)
        meta["path"] = str(path.resolve())
        meta["format"] = "pdf"
        meta["page_count"] = page_count
        meta["pages_read"] = f"{start_page}-{end_page}"

        logger.info("file_read PDF文本完成: %s, 读取页数=%d", path, end_page - start_page + 1)
        return ToolResult(content=content, metadata=meta)

    except Exception as e:
        return ToolResult(content=f"解析 PDF 失败: {e}", is_error=True)


async def _read_pdf_base64(
    path: Path,
    start_line: int | None,
    end_line: int | None,
) -> ToolResult:
    """使用 PyMuPDF 将 PDF 页渲染为 JPEG 图片并返回 base64。"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ToolResult(
            content="缺少 PyMuPDF 库，无法渲染 PDF。请安装: pip install PyMuPDF",
            is_error=True,
        )

    try:
        doc = fitz.open(str(path))
        page_count = len(doc)

        # 未指定范围时，默认只渲染前几页，避免生成过大响应
        if start_line is None and end_line is None:
            start_page = 1
            end_page = min(_PDF_DEFAULT_MAX_PAGES, page_count)
        else:
            start_page = max(1, start_line or 1)
            end_page = min(end_line or page_count, page_count) if end_line else page_count

        if start_page > page_count:
            doc.close()
            return ToolResult(
                content=f"PDF 共 {page_count} 页，start_line={start_page} 超出范围。",
                is_error=True,
                metadata={"page_count": page_count},
            )

        content_blocks: list[dict[str, Any]] = []
        total_bytes = 0

        for page_num in range(start_page, end_page + 1):
            page = doc[page_num - 1]
            # 渲染为 150 DPI 的 JPEG
            pix = page.get_pixmap(dpi=150)
            logger.debug(
                "PDF 第 %d/%d 页渲染: %dx%d (dpi=150, quality=75)",
                page_num, page_count, pix.width, pix.height,
            )
            img_bytes = pix.tobytes("jpeg", jpg_quality=75)
            b64_data = base64.b64encode(img_bytes).decode("ascii")
            total_bytes += len(img_bytes)

            content_blocks.append({
                "type": "image_base64",
                "data": b64_data,
                "media_type": "image/jpeg",
                "page": page_num,
            })

            if total_bytes > _MAX_BASE64_BYTES:
                # 超过大小限制，截断
                remaining = end_page - page_num
                desc = (
                    f"PDF 共 {page_count} 页，已渲染第 {start_page}-{page_num} 页为图片。"
                    f"（剩余 {remaining} 页因大小限制未输出，"
                    f"可调整 start_line/end_line 读取其他页）"
                )
                doc.close()
                logger.info("file_read PDF图片完成(截断): %s, 渲染页数=%d", path, page_num - start_page + 1)
                return ToolResult(
                    content=desc,
                    content_blocks=content_blocks,
                    metadata={
                        "path": str(path.resolve()),
                        "format": "pdf",
                        "page_count": page_count,
                        "pages_rendered": list(range(start_page, page_num + 1)),
                        "truncated": True,
                    },
                )

        doc.close()

        pages_hint = ""
        if page_count > end_page:
            pages_hint = (
                f"（仅渲染了前 {end_page - start_page + 1} 页，"
                f"共 {page_count} 页。可指定 start_line/end_line 读取更多页）"
            )
        desc = (
            f"PDF 共 {page_count} 页，已渲染第 {start_page}-{end_page} 页为图片"
            f"（共 {end_page - start_page + 1} 张）。{pages_hint}"
        )
        logger.info("file_read PDF图片完成: %s, 渲染页数=%d", path, end_page - start_page + 1)
        return ToolResult(
            content=desc,
            content_blocks=content_blocks,
            metadata={
                "path": str(path.resolve()),
                "format": "pdf",
                "page_count": page_count,
                "pages_rendered": list(range(start_page, end_page + 1)),
            },
        )

    except Exception as e:
        return ToolResult(content=f"渲染 PDF 失败: {e}", is_error=True)


async def _read_docx(
    path: Path,
    start_line: int | None,
    end_line: int | None,
) -> ToolResult:
    """读取 DOCX 文件。"""
    ext = path.suffix.lower()

    if ext == ".doc":
        # .doc 格式：python-docx 有限支持，提示用户
        return ToolResult(
            content=(
                "不支持旧版 .doc 格式。请将文件转换为 .docx 格式后重试。\n"
                "转换方法: 使用 LibreOffice 命令 `libreoffice --headless --convert-to docx <file.doc>`"
            ),
            is_error=True,
        )

    try:
        from docx import Document
    except ImportError:
        return ToolResult(
            content="缺少 python-docx 库，无法解析 DOCX。请安装: pip install python-docx",
            is_error=True,
        )

    try:
        doc = Document(str(path))
    except Exception as e:
        return ToolResult(content=f"打开 DOCX 失败: {e}", is_error=True)

    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            lines.append(text + "\n")
        else:
            lines.append("\n")  # 保留空行

    total = len(lines)
    if total == 0:
        return ToolResult(
            content="DOCX 文件内容为空。",
            metadata={"path": str(path.resolve()), "format": "docx"},
        )

    s = max(1, start_line or 1)
    e = min(end_line or total, total) if end_line else total

    if s > total:
        return ToolResult(
            content=f"文档共 {total} 行，start_line={s} 超出范围。",
            is_error=True,
            metadata={"total_lines": total},
        )

    sliced = lines[s - 1:e]
    content, meta = _format_lines(sliced, s, e, total)
    meta["path"] = str(path.resolve())
    meta["format"] = "docx"

    logger.info("file_read DOCX完成: %s, 返回行数=%d", path, meta["lines_output"])
    return ToolResult(content=content, metadata=meta)


# 大图片自动缩放阈值（超过此尺寸则按比例缩小）
_IMAGE_RESIZE_THRESHOLD = 2048  # 像素（长边）
_IMAGE_JPEG_QUALITY = 80


def _resize_image_if_needed(
    data: bytes, media_type: str
) -> tuple[bytes, str, int, int, bool]:
    """
    如果图片尺寸过大，按比例缩小以减少传输体积。

    Returns:
        (处理后的bytes, media_type, width, height, was_resized)
    """
    try:
        from PIL import Image
    except ImportError:
        return data, media_type, 0, 0, False

    try:
        img = Image.open(io.BytesIO(data))
        orig_width, orig_height = img.size
        was_resized = False

        max_dim = max(orig_width, orig_height)
        if max_dim > _IMAGE_RESIZE_THRESHOLD:
            scale = _IMAGE_RESIZE_THRESHOLD / max_dim
            new_width = int(orig_width * scale)
            new_height = int(orig_height * scale)
            logger.debug(
                "图像缩放: %dx%d -> %dx%d (阈值=%d, 缩放比=%.3f)",
                orig_width, orig_height, new_width, new_height,
                _IMAGE_RESIZE_THRESHOLD, scale,
            )
            img = img.resize((new_width, new_height), Image.LANCZOS)
            was_resized = True
        else:
            logger.debug("图像无需缩放: %dx%d (长边=%d, 阈值=%d)",
                         orig_width, orig_height, max_dim, _IMAGE_RESIZE_THRESHOLD)

        buf = io.BytesIO()
        if media_type == "image/jpeg" or was_resized:
            if img.mode in ("RGBA", "LA", "P"):
                img = img.convert("RGB")
            img.save(buf, format="JPEG", quality=_IMAGE_JPEG_QUALITY)
            output_media_type = "image/jpeg"
        else:
            img.save(buf, format="PNG")
            output_media_type = "image/png"

        final_width, final_height = img.size
        return buf.getvalue(), output_media_type, final_width, final_height, was_resized

    except Exception as e:
        logger.warning("图像处理异常，返回原始数据: %s", e)
        return data, media_type, 0, 0, False


async def _read_image_base64(path: Path) -> ToolResult:
    """读取图片文件并返回 base64 编码。大图自动缩放以减少传输体积。"""
    ext = path.suffix.lower()
    media_type = _IMAGE_EXTENSIONS.get(ext)

    if not media_type:
        mime, _ = mimetypes.guess_type(str(path))
        media_type = mime or "application/octet-stream"

    try:
        file_size = path.stat().st_size
        data = path.read_bytes()
    except PermissionError:
        return ToolResult(content=f"权限不足，无法读取: {path}", is_error=True)
    except Exception as e:
        return ToolResult(content=f"读取图片失败: {e}", is_error=True)

    # 自动缩放处理
    processed_data, final_media_type, width, height, was_resized = (
        _resize_image_if_needed(data, media_type)
    )

    # 最终大小检查
    if len(processed_data) > _MAX_BASE64_BYTES:
        return ToolResult(
            content=f"图片文件过大 ({file_size / 1024 / 1024:.1f}MB，"
                    f"缩放后仍超过 {_MAX_BASE64_BYTES / 1024 / 1024:.0f}MB 限制)。",
            is_error=True,
            metadata={
                "path": str(path.resolve()),
                "original_size": file_size,
                "dimensions": f"{width}x{height}" if width else "unknown",
            },
        )

    b64_data = base64.b64encode(processed_data).decode("ascii")

    size_info = f"{file_size} 字节"
    if was_resized:
        size_info += f"（已自动缩放，缩放后 {len(processed_data)} 字节）"
    dim_info = f"，尺寸 {width}x{height}" if width else ""

    logger.info("file_read 图片完成: %s, 尺寸=%s", path, f"{width}x{height}" if width else "unknown")
    return ToolResult(
        content=f"图片: {path.name} ({final_media_type}, {size_info}{dim_info})",
        content_blocks=[{
            "type": "image_base64",
            "data": b64_data,
            "media_type": final_media_type,
        }],
        metadata={
            "path": str(path.resolve()),
            "format": "image",
            "media_type": final_media_type,
            "original_size": file_size,
            "encoded_size": len(processed_data),
            "was_resized": was_resized,
            **({"width": width, "height": height} if width else {}),
        },
    )


async def _read_binary_base64(path: Path) -> ToolResult:
    """将任意二进制文件以 base64 编码返回。"""
    try:
        file_size = path.stat().st_size
        if file_size > _MAX_BASE64_BYTES:
            return ToolResult(
                content=f"文件过大 ({file_size / 1024 / 1024:.1f}MB > "
                        f"{_MAX_BASE64_BYTES / 1024 / 1024:.0f}MB 限制)。",
                is_error=True,
            )

        data = path.read_bytes()
    except PermissionError:
        return ToolResult(content=f"权限不足，无法读取: {path}", is_error=True)
    except Exception as e:
        return ToolResult(content=f"读取文件失败: {e}", is_error=True)

    b64_data = base64.b64encode(data).decode("ascii")

    mime, _ = mimetypes.guess_type(str(path))
    media_type = mime or "application/octet-stream"

    logger.info("file_read 二进制完成: %s, size=%d", path, file_size)
    return ToolResult(
        content=f"二进制文件: {path.name} ({media_type}, {file_size} 字节)",
        content_blocks=[{
            "type": "image_base64",
            "data": b64_data,
            "media_type": media_type,
        }],
        metadata={
            "path": str(path.resolve()),
            "format": "binary",
            "media_type": media_type,
            "file_size": file_size,
        },
    )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 主入口：file_read
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@tool(name="file_read",
      description=(
          "读取指定路径的文件内容。支持多种格式：\n"
          "- 文本文件(txt/md/py/json/yaml/log等): 按行输出，带行号\n"
          "- CSV/TSV: 解析后按行输出，列用 | 分隔\n"
          "- XLSX/XLS: 多 Sheet 时先列出 Sheet 信息，需指定 sheet_name 读取\n"
          "- PDF: 提取文本(默认) 或渲染为图片(base64模式)\n"
          "- DOCX: 提取段落文本，按行输出\n"
          "- 图片(png/jpg/gif等): 以 base64 编码返回\n"
          "参数 start_line_or_page / end_line_or_page 对文本/CSV/DOCX/XLSX 表示行号，对 PDF 表示页码"
      ))
async def file_read(
    path: str,
    sheet_name: str | None = None,
    start_line_or_page: int | None = None,
    end_line_or_page: int | None = None,
    output_format: str = "auto",
    encoding: str | None = None,
    xlsx_range: str | None = None,
    render_mode: str = "values",
    row_mode: str = "text",
    include_tables: bool = True,
    include_merges: bool = True,
) -> ToolResult:
    """
    读取文件内容，支持多种格式。始终显示行号。

    Args:
        path: 文件路径（支持绝对路径和相对路径）
        sheet_name: XLSX 工作表名称。仅当读取 XLSX 文件且有多个工作表时需要指定，未指定默认先返回所有表名
        start_line_or_page: 起始位置（从1开始，包含该行/页）。对文本/CSV/DOCX/XLSX 表示起始行号，对 PDF 表示起始页码。未指定则从第1行/页开始
        end_line_or_page: 结束位置（包含该行/页）。对文本/CSV/DOCX/XLSX 表示结束行号，对 PDF 表示结束页码。未指定则返回到文件末尾
        output_format: 输出格式，可选 "auto"（自动判断）、"text"（强制文本）、"base64"（强制二进制编码）。默认 "auto"
        encoding: 强制指定文件编码（如 "utf-8"、"gbk"）。默认自动检测
        xlsx_range: XLSX 专用 A1 区域（如 "A1:D20"）。指定后优先于 start_line_or_page/end_line_or_page
        render_mode: XLSX 专用读取模式，可选 "values"、"formulas"、"both"
        row_mode: XLSX 专用 metadata 行模式，可选 "text"、"arrays"、"objects"
        include_tables: XLSX metadata 是否包含 Excel tables 摘要
        include_merges: XLSX metadata 是否包含合并单元格摘要
    """
    # ── 参数校验 ──
    # 类型强制转换：LLM 可能传字符串形式的数字/布尔值
    if isinstance(start_line_or_page, str):
        try:
            start_line_or_page = int(start_line_or_page)
        except (ValueError, TypeError):
            return ToolResult(content=f"start_line_or_page 必须是整数，收到: {start_line_or_page!r}", is_error=True)
    if isinstance(end_line_or_page, str):
        try:
            end_line_or_page = int(end_line_or_page)
        except (ValueError, TypeError):
            return ToolResult(content=f"end_line_or_page 必须是整数，收到: {end_line_or_page!r}", is_error=True)
    if isinstance(output_format, str):
        output_format = output_format.lower()

    error = _check_path_safety(path)
    if error:
        return ToolResult(content=error, is_error=True)

    target = Path(path)
    if not target.exists():
        return ToolResult(content=f"文件不存在: {path}", is_error=True)
    if not target.is_file():
        return ToolResult(content=f"不是文件: {path}", is_error=True)

    # start_line_or_page / end_line_or_page 校验
    if start_line_or_page is not None and start_line_or_page < 1:
        start_line_or_page = 1
    if end_line_or_page is not None and start_line_or_page is not None and end_line_or_page < start_line_or_page:
        return ToolResult(
            content=f"end_line_or_page ({end_line_or_page}) 不能小于 start_line_or_page ({start_line_or_page})。",
            is_error=True,
        )

    # output_format 校验
    if output_format not in ("auto", "text", "base64"):
        return ToolResult(
            content=f"不支持的 output_format: {output_format}。可选: auto, text, base64",
            is_error=True,
        )

    # ── 文件类型检测 ──
    file_type = _detect_file_type(target)
    logger.info("file_read 开始: path=%s, file_type=%s, output_format=%s", path, file_type, output_format)

    # ── 路由到对应解析器 ──
    try:
        # 确定是否使用 base64 模式
        use_base64 = output_format == "base64"
        if output_format == "auto":
            use_base64 = file_type == "image" or file_type == "binary"

        if file_type == "text":
            if use_base64:
                return await _read_binary_base64(target)
            return await _read_text(target, start_line_or_page, end_line_or_page, encoding)

        elif file_type == "csv":
            if use_base64:
                return await _read_binary_base64(target)
            return await _read_csv(target, start_line_or_page, end_line_or_page, encoding)

        elif file_type == "xlsx":
            if use_base64:
                return await _read_binary_base64(target)
            return await _read_xlsx(
                target, sheet_name, start_line_or_page, end_line_or_page,
                xlsx_range=xlsx_range, render_mode=render_mode, row_mode=row_mode,
                include_tables=include_tables, include_merges=include_merges,
            )

        elif file_type == "docx":
            if use_base64:
                return await _read_binary_base64(target)
            return await _read_docx(target, start_line_or_page, end_line_or_page)

        elif file_type == "pdf":
            if use_base64:
                return await _read_pdf_base64(target, start_line_or_page, end_line_or_page)
            return await _read_pdf_text(target, start_line_or_page, end_line_or_page)

        elif file_type == "image":
            return await _read_image_base64(target)

        else:
            # binary 或未知类型
            return await _read_binary_base64(target)

    except Exception as e:
        return ToolResult(content=f"读取文件异常: {type(e).__name__}: {e}", is_error=True)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_write（保持不变）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@tool(name="file_write",
      description="将内容写入指定路径的文件。不存在则创建，已存在则覆盖。")
async def file_write(path: str, content: str,
                     append: bool = False) -> ToolResult:
    """
    将内容写入指定路径的文件。

    Args:
        path: 文件路径（支持绝对路径和相对路径）。不存在则创建，已存在则覆盖
        content: 要写入的文本内容
        append: 是否追加到文件末尾。默认为 False（覆盖写入）
    """
    logger.info("file_write 开始: path=%s, content_len=%d, append=%s", path, len(content), append)
    error = _check_path_safety(path)
    if error:
        return ToolResult(content=error, is_error=True)

    target = Path(path)
    try:
        target.parent.mkdir(parents=True, exist_ok=True)
        mode = "a" if append else "w"
        with open(target, mode, encoding="utf-8") as f:
            f.write(content)

        action = "追加" if append else "写入"
        logger.info("file_write 成功: %s %d 字符到 %s", action, len(content), target.resolve())
        return ToolResult(
            content=f"文件{action}成功: {target.resolve()} ({len(content)} 字符)",
            metadata={"path": str(target.resolve()),
                      "chars_written": len(content)},
        )
    except Exception as e:
        return ToolResult(content=f"写入文件失败: {e}", is_error=True)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit — 精确搜索替换编辑
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# ── 注释风格映射 ──
_COMMENT_STYLES: dict[str, tuple[str, str]] = {
    # (prefix, suffix) — 行注释用 (prefix, "")，块注释用 (prefix, suffix)
    ".css":  ("/* ", " */"),
    ".html": ("<!-- ", " -->"),
    ".xml":  ("<!-- ", " -->"),
    ".md":   ("<!-- ", " -->"),
    ".ini":  ("; ", ""),
    ".toml": ("# ", ""),
    ".yaml": ("# ", ""),
    ".yml":  ("# ", ""),
    ".conf": ("# ", ""),
}

# JSON 等不支持注释的格式
_NO_COMMENT_EXTENSIONS = {".json", ".csv", ".tsv"}


def _get_comment_style(path: Path) -> tuple[str, str] | None:
    """根据扩展名返回注释风格 (prefix, suffix)，不支持则返回 None。"""
    ext = path.suffix.lower()
    if ext in _NO_COMMENT_EXTENSIONS:
        return None
    return _COMMENT_STYLES.get(ext, ("# ", ""))  # 默认用 # 注释


def _build_diff_preview(
    original: str, replaced: str, match_pos: int
) -> str:
    """构建替换位置 ±3 行的 diff 预览。"""
    orig_lines = original.split("\n")
    # 找到匹配所在的行号
    before_text = original[:match_pos]
    match_line = before_text.count("\n")

    # 计算新文本中替换区域的行号（近似）
    repl_lines = replaced.split("\n")

    ctx_start = max(0, match_line - 3)
    ctx_end = min(len(repl_lines), match_line + 4)

    preview_lines = []
    for i in range(ctx_start, ctx_end):
        marker = ">" if match_line <= i < match_line + max(1, 1) else " "
        preview_lines.append(f"  {marker} {i + 1:>{len(str(ctx_end))}} | {repl_lines[i]}")

    return "\n".join(preview_lines)


def _atomic_write_text(path: Path, content: str) -> None:
    """原子写入纯文本文件。"""
    import tempfile
    tmp = tempfile.NamedTemporaryFile(
        mode="w", encoding="utf-8",
        dir=str(path.parent), suffix=".tmp", delete=False,
    )
    try:
        tmp.write(content)
        tmp.close()
        os.replace(tmp.name, str(path))
    except Exception:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass
        raise


def _atomic_write_binary(path: Path, data: bytes) -> None:
    """原子写入二进制文件。"""
    import tempfile
    tmp = tempfile.NamedTemporaryFile(
        dir=str(path.parent), suffix=".tmp", delete=False,
    )
    try:
        tmp.write(data)
        tmp.close()
        os.replace(tmp.name, str(path))
    except Exception:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass
        raise


# ── 纯文本编辑 ──

async def _edit_text(
    path: Path,
    target_content: str,
    replacement_content: str,
    start_line: int | None,
    end_line: int | None,
    allow_multiple: bool,
    highlight: str | None,
    comment: str | None,
) -> ToolResult:
    """纯文本文件的精确搜索替换。"""
    enc = _detect_encoding(path)
    if enc == "binary":
        return ToolResult(content="文件似乎是二进制格式，无法进行文本编辑。", is_error=True)

    try:
        with open(path, "r", encoding=enc, errors="replace") as f:
            original = f.read()
    except Exception as e:
        return ToolResult(content=f"读取文件失败: {e}", is_error=True)

    # ── 行号范围约束 ──
    search_text = original
    offset = 0  # 原始文本中的字符偏移
    if start_line is not None or end_line is not None:
        lines = original.split("\n")
        total_lines = len(lines)
        s = max(1, start_line or 1)
        e = min(end_line or total_lines, total_lines) if end_line else total_lines
        if s > total_lines:
            return ToolResult(
                content=f"文件共 {total_lines} 行，start_line={s} 超出范围。",
                is_error=True,
            )
        # 计算偏移量
        offset = sum(len(lines[i]) + 1 for i in range(s - 1))
        search_text = "\n".join(lines[s - 1:e])

    # ── 精确匹配 ──
    count = search_text.count(target_content)
    if count == 0:
        hint = ""
        if start_line is not None:
            hint = f"（在行 {start_line}-{end_line or '末尾'} 范围内）"
        return ToolResult(
            content=f"未找到目标内容{hint}。请使用 file_read 确认文件最新内容后重试。",
            is_error=True,
            metadata={"path": str(path.resolve()), "match_count": 0},
        )
    if count > 1 and not allow_multiple:
        return ToolResult(
            content=f"找到 {count} 个匹配，请提供更精确的 target_content，"
                    f"或使用 start_line/end_line 缩小范围，或设置 allow_multiple=True。",
            is_error=True,
            metadata={"path": str(path.resolve()), "match_count": count},
        )

    # ── 执行替换 ──
    if start_line is not None or end_line is not None:
        new_search_text = search_text.replace(target_content, replacement_content)
        lines = original.split("\n")
        total_lines = len(lines)
        s = max(1, start_line or 1)
        e = min(end_line or total_lines, total_lines) if end_line else total_lines
        new_lines = lines[:s - 1] + new_search_text.split("\n") + lines[e:]
        replaced = "\n".join(new_lines)
        actual_count = new_search_text.count(replacement_content) if replacement_content != target_content else count
    else:
        replaced = original.replace(target_content, replacement_content)
        actual_count = count

    # ── 批注（在替换位置上方添加注释）──
    comment_applied = False
    comment_hint = ""
    if comment:
        style = _get_comment_style(path)
        if style is None:
            comment_hint = "（该文件类型不支持注释，已忽略 comment 参数）"
        else:
            prefix, suffix = style
            comment_line = f"{prefix}[AGENT COMMENT] {comment}{suffix}\n"
            # 在第一个替换位置上方插入注释
            idx = replaced.find(replacement_content)
            if idx >= 0:
                # 找到该行的起始位置
                line_start = replaced.rfind("\n", 0, idx)
                line_start = line_start + 1 if line_start >= 0 else 0
                replaced = replaced[:line_start] + comment_line + replaced[line_start:]
                comment_applied = True

    # ── 原子写入 ──
    try:
        _atomic_write_text(path, replaced)
    except Exception as e:
        return ToolResult(content=f"写入文件失败: {e}", is_error=True)

    # ── 构建返回信息 ──
    match_pos = offset if offset else original.find(target_content)
    diff_preview = _build_diff_preview(original, replaced, max(0, match_pos))

    msg = f"替换成功: {path.name}，共 {actual_count} 处替换。"
    if highlight:
        msg += "（纯文本文件不支持标色，已忽略 highlight 参数）"
    if comment_hint:
        msg += comment_hint
    elif comment_applied:
        msg += " 已添加批注注释。"

    logger.info("file_edit 纯文本成功: path=%s, replacements=%d", path, actual_count)
    return ToolResult(
        content=msg + f"\n\n预览:\n{diff_preview}",
        metadata={
            "path": str(path.resolve()),
            "format": "text",
            "replacements": actual_count,
            "highlight": highlight,
            "comment_added": comment_applied,
        },
    )


# ── DOCX 编辑 ──

def _insert_docx_run_after(anchor_run: Any, text: str, template_run: Any | None = None) -> Any:
    """Insert a run after `anchor_run`, copying character formatting from `template_run`."""
    from docx.oxml import OxmlElement
    from docx.text.run import Run

    template_run = template_run or anchor_run
    new_r = OxmlElement("w:r")
    r_pr = template_run._r.rPr
    if r_pr is not None:
        new_r.append(deepcopy(r_pr))
    anchor_run._r.addnext(new_r)

    new_run = Run(new_r, anchor_run._parent)
    new_run.text = text
    return new_run


def _normalize_docx_edit_text(
    target_content: str,
    replacement_content: str,
) -> tuple[str, str]:
    """Tolerate a single trailing newline copied from file_read's paragraph output."""
    if (
        target_content.endswith("\n")
        and replacement_content.endswith("\n")
        and "\n" not in target_content[:-1]
        and "\n" not in replacement_content[:-1]
    ):
        return target_content[:-1], replacement_content[:-1]
    return target_content, replacement_content


async def _edit_docx(
    path: Path,
    target_content: str,
    replacement_content: str,
    allow_multiple: bool,
    highlight: str | None,
    comment: str | None,
) -> ToolResult:
    """DOCX 文件的精确搜索替换，保留格式。"""
    try:
        from docx import Document
    except ImportError:
        return ToolResult(
            content="缺少 python-docx 库。请安装: pip install python-docx",
            is_error=True,
        )

    try:
        doc = Document(str(path))
    except Exception as e:
        return ToolResult(content=f"打开 DOCX 失败: {e}", is_error=True)

    target_content, replacement_content = _normalize_docx_edit_text(
        target_content,
        replacement_content,
    )

    # ── 构建虚拟文本流及 Run 映射 ──
    virtual_parts: list[str] = []
    # 每个 entry: (para_idx, run_idx, char_start, char_end)
    run_map: list[tuple[int, int, int, int]] = []
    char_pos = 0

    for pi, para in enumerate(doc.paragraphs):
        if pi > 0:
            virtual_parts.append("\n")
            char_pos += 1
        for ri, run in enumerate(para.runs):
            text = run.text
            start = char_pos
            end = char_pos + len(text)
            if text:
                run_map.append((pi, ri, start, end))
                virtual_parts.append(text)
                char_pos = end

    virtual_text = "".join(virtual_parts)

    # ── 精确匹配 ──
    count = virtual_text.count(target_content)
    if count == 0:
        return ToolResult(
            content="未找到目标内容。请使用 file_read 确认文档最新内容后重试。",
            is_error=True,
            metadata={"path": str(path.resolve()), "match_count": 0},
        )
    if count > 1 and not allow_multiple:
        return ToolResult(
            content=f"找到 {count} 个匹配，请提供更精确的 target_content 或设置 allow_multiple=True。",
            is_error=True,
            metadata={"path": str(path.resolve()), "match_count": count},
        )

    matches: list[tuple[int, int, list[tuple[int, int, int, int]]]] = []
    search_start = 0
    while search_start < len(virtual_text):
        match_idx = virtual_text.find(target_content, search_start)
        if match_idx < 0:
            break
        match_end = match_idx + len(target_content)
        involved = [
            (pi, ri, cs, ce)
            for pi, ri, cs, ce in run_map
            if cs < match_end and ce > match_idx
        ]
        if involved:
            matches.append((match_idx, match_end, involved))
        search_start = match_end
        if not allow_multiple:
            break

    cross_paragraph = [
        (match_idx, match_end)
        for match_idx, match_end, involved in matches
        if involved[0][0] != involved[-1][0]
    ]
    if cross_paragraph:
        return ToolResult(
            content=(
                "DOCX 编辑被拒绝：当前安全编辑器只支持单个段落内的替换。"
                "这次匹配跨越了多个段落，继续写入可能产生额外空行或破坏段落对齐。"
                "请把 target_content 缩小到一个段落内后重试。"
            ),
            is_error=True,
            metadata={
                "path": str(path.resolve()),
                "match_count": count,
                "cross_paragraph_matches": len(cross_paragraph),
            },
        )

    # ── 执行替换（逐个匹配处理）──
    actual_count = 0
    comment_count = 0
    for match_idx, match_end, involved in reversed(matches):
        first_pi, first_ri, first_cs, first_ce = involved[0]
        last_pi, last_ri, last_cs, last_ce = involved[-1]
        first_para = doc.paragraphs[first_pi]
        first_run = first_para.runs[first_ri]

        # 第一段第一段：保留前缀 + 替换内容
        prefix = target_content[:0]  # 空
        if first_cs < match_idx:
            prefix = first_run.text[:match_idx - first_cs]

        # 最后一段最后一段：保留后缀
        last_para = doc.paragraphs[last_pi]
        last_run = last_para.runs[last_ri]
        suffix = ""
        if last_ce > match_end:
            suffix = last_run.text[match_end - last_cs:]

        runs = first_para.runs
        first_run.text = prefix
        if first_ri != last_ri:
            for idx in range(first_ri + 1, last_ri):
                runs[idx].text = ""
            last_run.text = ""

        replacement_run = _insert_docx_run_after(first_run, replacement_content)
        if suffix:
            _insert_docx_run_after(replacement_run, suffix, last_run)

        # ── 标色 ──
        if highlight:
            try:
                from docx.enum.text import WD_COLOR_INDEX
                _DOCX_COLOR_MAP = {
                    "yellow": WD_COLOR_INDEX.YELLOW,
                    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
                    "red": WD_COLOR_INDEX.RED,
                    "pink": WD_COLOR_INDEX.PINK,
                }
                color = _DOCX_COLOR_MAP.get(highlight)
                if color:
                    replacement_run.font.highlight_color = color
            except Exception as e:
                logger.warning("DOCX 标色失败: %s", e)

        # ── 批注 ──
        if comment:
            try:
                if hasattr(doc, "add_comment"):
                    doc.add_comment(
                        replacement_run,
                        text=comment,
                        author="Agent",
                        initials="AG",
                    )
                    comment_count += 1
                else:  # pragma: no cover - compatibility with older python-docx
                    logger.warning("python-docx 版本不支持批注，需要 >= 1.2.0")
            except Exception as e:
                logger.warning("DOCX 批注添加失败: %s", e)

        actual_count += 1

    # ── 原子写入 ──
    try:
        buf = io.BytesIO()
        doc.save(buf)
        _atomic_write_binary(path, buf.getvalue())
    except Exception as e:
        return ToolResult(content=f"保存 DOCX 失败: {e}", is_error=True)

    comment_status = ""
    if comment and comment_count:
        comment_status = f" 已添加批注 {comment_count} 处。"
    elif comment:
        comment_status = " 批注添加失败，已完成文本替换。"

    logger.info("file_edit DOCX成功: path=%s, replacements=%d", path, actual_count)
    return ToolResult(
        content=f"DOCX 替换成功: {path.name}，共 {actual_count} 处替换。"
                + (" 已标色。" if highlight else "")
                + comment_status,
        metadata={
            "path": str(path.resolve()),
            "format": "docx",
            "replacements": actual_count,
            "highlight": highlight,
            "comment_added": comment_count > 0,
            "comments_added": comment_count,
        },
    )




# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# XLSX 结构化读写辅助函数
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_XLSX_STRUCTURAL_OPS = {"append_rows", "upsert_rows", "delete_rows", "insert_rows"}
_XLSX_ROW_STRUCTURE_OPS = {"delete_rows", "insert_rows"}
_XLSX_VALUE_INPUTS = {"auto", "raw", "formula"}
_XLSX_CELL_KINDS = {"text", "number", "bool", "date", "formula", "blank"}


def _hash_json(data: Any) -> str:
    """生成短 hash，用于 workbook structure/content token。"""
    raw = json.dumps(data, ensure_ascii=False, sort_keys=True, default=str)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]


def _json_value(value: Any) -> Any:
    """把 openpyxl 值转换成 JSON metadata 友好的值。"""
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    return str(value)


def _metadata_scalar(value: Any) -> Any:
    if isinstance(value, dict):
        if "formula" in value:
            return value
        return value.get("value")
    return _json_value(value)


def _xlsx_cell_payload(cell, value_cell=None, render_mode: str = "values") -> Any:
    """返回单元格的结构化读取值。"""
    if render_mode == "both":
        formula_value = cell.value
        calculated = value_cell.value if value_cell is not None else None
        if isinstance(formula_value, str) and formula_value.startswith("="):
            return {"formula": formula_value, "value": _json_value(calculated)}
        return _json_value(calculated if value_cell is not None else formula_value)
    return _json_value(cell.value)


def _xlsx_cell_display(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, dict):
        formula = value.get("formula")
        calculated = value.get("value")
        if formula and calculated not in (None, ""):
            return f"{formula} => {calculated}"
        if formula:
            return str(formula)
        return str(value.get("value") or "")
    return str(value)


def _parse_a1_range(cell_range: str) -> tuple[int, int, int, int]:
    """解析 A1 range，返回 min_col, min_row, max_col, max_row。"""
    from openpyxl.utils.cell import range_boundaries

    if not isinstance(cell_range, str) or not cell_range.strip():
        raise ValueError("range 必须是非空 A1 字符串")
    text = cell_range.strip().replace("$", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    if not re.match(r"^[A-Za-z]+[1-9][0-9]*(?::[A-Za-z]+[1-9][0-9]*)?$", text):
        raise ValueError(f"不合法的 A1 range: {cell_range!r}")
    min_col, min_row, max_col, max_row = range_boundaries(text)
    if min_col > max_col or min_row > max_row:
        raise ValueError(f"不合法的 A1 range: {cell_range!r}")
    return min_col, min_row, max_col, max_row


def _parse_cell_address(address: str) -> tuple[int, int]:
    """解析单元格地址，返回 row, col。"""
    from openpyxl.utils.cell import coordinate_to_tuple

    if not isinstance(address, str) or not address.strip():
        raise ValueError("cell 必须是非空 A1 单元格地址")
    text = address.strip().replace("$", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    if not re.match(r"^[A-Za-z]+[1-9][0-9]*$", text):
        raise ValueError(f"不合法的单元格地址: {address!r}")
    return coordinate_to_tuple(text)


def _column_index(value: str | int) -> int:
    from openpyxl.utils.cell import column_index_from_string

    if isinstance(value, int):
        if value < 1:
            raise ValueError("列索引必须从 1 开始")
        return value
    if isinstance(value, str) and value.strip().isdigit():
        num = int(value.strip())
        if num < 1:
            raise ValueError("列索引必须从 1 开始")
        return num
    if isinstance(value, str):
        return column_index_from_string(value.strip())
    raise ValueError(f"不合法的列标识: {value!r}")


def _range_from_bounds(min_col: int, min_row: int, max_col: int, max_row: int) -> str:
    from openpyxl.utils.cell import get_column_letter

    start = f"{get_column_letter(min_col)}{min_row}"
    end = f"{get_column_letter(max_col)}{max_row}"
    return start if start == end else f"{start}:{end}"


def _bounds_overlap(a: tuple[int, int, int, int], b: tuple[int, int, int, int]) -> bool:
    amin_col, amin_row, amax_col, amax_row = a
    bmin_col, bmin_row, bmax_col, bmax_row = b
    return not (amax_col < bmin_col or bmax_col < amin_col or amax_row < bmin_row or bmax_row < amin_row)


def _bounds_contains(outer: tuple[int, int, int, int], inner: tuple[int, int, int, int]) -> bool:
    omin_col, omin_row, omax_col, omax_row = outer
    imin_col, imin_row, imax_col, imax_row = inner
    return omin_col <= imin_col and omin_row <= imin_row and omax_col >= imax_col and omax_row >= imax_row


def _merged_range_bounds(merged_range) -> tuple[int, int, int, int]:
    return merged_range.min_col, merged_range.min_row, merged_range.max_col, merged_range.max_row


def _is_merged_non_anchor(ws, row: int, col: int) -> str | None:
    for merged in ws.merged_cells.ranges:
        bounds = _merged_range_bounds(merged)
        if bounds[0] <= col <= bounds[2] and bounds[1] <= row <= bounds[3]:
            if row != merged.min_row or col != merged.min_col:
                return str(merged)
    return None


def _worksheet_tables(ws) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    try:
        names = list(ws.tables)
    except Exception:
        names = []
    for name in names:
        try:
            table = ws.tables[name]
            ref = getattr(table, "ref", table if isinstance(table, str) else None)
            if not ref:
                continue
            tables.append({
                "name": getattr(table, "name", name),
                "display_name": getattr(table, "displayName", name),
                "ref": ref,
                "totals_row_count": int(getattr(table, "totalsRowCount", 0) or 0),
            })
        except Exception:
            continue
    return tables


def _table_by_name(ws, table_name: str):
    if not table_name:
        return None
    try:
        if table_name in ws.tables:
            return ws.tables[table_name]
    except Exception:
        pass
    for table in _worksheet_tables(ws):
        if table.get("name") == table_name or table.get("display_name") == table_name:
            try:
                return ws.tables[table.get("name")]
            except Exception:
                return None
    return None


def _profile_workbook(wb) -> dict[str, Any]:
    sheets: list[dict[str, Any]] = []
    for ws in wb.worksheets:
        sheet_info: dict[str, Any] = {
            "name": ws.title,
            "max_row": ws.max_row or 0,
            "max_column": ws.max_column or 0,
            "dimension": ws.calculate_dimension(),
            "tables": _worksheet_tables(ws),
            "merged_ranges": [str(rng) for rng in ws.merged_cells.ranges],
            "charts": len(getattr(ws, "_charts", []) or []),
        }
        try:
            sheet_info["data_validations"] = len(ws.data_validations.dataValidation)
        except Exception:
            sheet_info["data_validations"] = 0
        try:
            sheet_info["conditional_formatting"] = len(ws.conditional_formatting)
        except Exception:
            sheet_info["conditional_formatting"] = 0
        sheets.append(sheet_info)

    defined_names: list[str] = []
    try:
        for defined_name in wb.defined_names:
            name = getattr(defined_name, "name", str(defined_name))
            defined_names.append(name)
    except Exception:
        pass

    return {
        "sheet_names": list(wb.sheetnames),
        "sheets": sheets,
        "defined_names": sorted(defined_names),
    }


def _content_token(ws, bounds: tuple[int, int, int, int] | None = None) -> str:
    if bounds is None:
        bounds = (1, 1, ws.max_column or 1, ws.max_row or 1)
    min_col, min_row, max_col, max_row = bounds
    rows: list[list[Any]] = []
    for row_idx in range(min_row, max_row + 1):
        row_values: list[Any] = []
        for col_idx in range(min_col, max_col + 1):
            row_values.append(_json_value(ws.cell(row=row_idx, column=col_idx).value))
        rows.append(row_values)
    return _hash_json({"sheet": ws.title, "range": _range_from_bounds(*bounds), "rows": rows})


def _resolve_worksheet_for_table_edit(wb, sheet_name: str | None, payload: dict[str, Any]) -> tuple[Any | None, str | None]:
    table_name = payload.get("table_name") if isinstance(payload, dict) else None

    if sheet_name is not None:
        if sheet_name not in wb.sheetnames:
            return None, (
                f"工作表 \"{sheet_name}\" 不存在。可用工作表: "
                f"{', '.join(repr(n) for n in wb.sheetnames)}"
            )
        ws = wb[sheet_name]
        if table_name and _table_by_name(ws, table_name) is None:
            return None, f"工作表 \"{sheet_name}\" 中不存在表格 {table_name!r}。"
        return ws, None

    if table_name:
        matches = [ws for ws in wb.worksheets if _table_by_name(ws, table_name) is not None]
        if len(matches) == 1:
            return matches[0], None
        if not matches:
            return None, f"工作簿中不存在表格 {table_name!r}。"
        return None, f"表格 {table_name!r} 在多个工作表中存在，请指定 sheet_name。"

    if len(wb.sheetnames) == 1:
        return wb[wb.sheetnames[0]], None
    return None, (
        f"文件包含 {len(wb.sheetnames)} 个工作表，必须指定 sheet_name。"
        f"可用工作表: {', '.join(repr(n) for n in wb.sheetnames)}"
    )


def _coerce_cell_value(value: Any, kind: str | None = None, value_input: str = "auto") -> tuple[Any, str | None]:
    """将 payload 值转换为 openpyxl 可写值，返回 (value, number_format)。"""
    value_input = (value_input or "auto").lower()
    if value_input not in _XLSX_VALUE_INPUTS:
        raise ValueError(f"不支持的 value_input: {value_input!r}。可选: auto, raw, formula")
    if kind is not None:
        kind = str(kind).lower()
        if kind not in _XLSX_CELL_KINDS:
            raise ValueError(f"不支持的 kind: {kind!r}。可选: {', '.join(sorted(_XLSX_CELL_KINDS))}")

    if kind == "blank":
        return None, None
    if value is None:
        return None, None

    if kind == "text":
        return str(value), None
    if kind == "formula" or value_input == "formula":
        if not isinstance(value, str) or not value.startswith("="):
            raise ValueError("formula 值必须是以 '=' 开头的字符串")
        return value, None
    if kind == "bool":
        if isinstance(value, bool):
            return value, None
        if isinstance(value, str) and value.strip().lower() in ("true", "false"):
            return value.strip().lower() == "true", None
        raise ValueError(f"无法将 {value!r} 转为 bool")
    if kind == "number":
        if isinstance(value, bool):
            raise ValueError("bool 不能作为 number 写入")
        if isinstance(value, (int, float)):
            return value, None
        if isinstance(value, str) and re.match(r"^[+-]?\d+(?:\.\d+)?$", value.strip()):
            text = value.strip()
            return (float(text), None) if "." in text else (int(text), None)
        raise ValueError(f"无法将 {value!r} 转为 number")
    if kind == "date":
        if isinstance(value, datetime):
            return value, "yyyy-mm-dd h:mm:ss"
        if isinstance(value, date):
            return value, "yyyy-mm-dd"
        if isinstance(value, str):
            text = value.strip()
            try:
                if "T" in text or " " in text:
                    return datetime.fromisoformat(text), "yyyy-mm-dd h:mm:ss"
                return date.fromisoformat(text), "yyyy-mm-dd"
            except ValueError as exc:
                raise ValueError(f"无法将 {value!r} 转为 date") from exc
        raise ValueError(f"无法将 {value!r} 转为 date")

    if value_input == "raw":
        if isinstance(value, (str, int, float, bool, datetime, date)):
            return value, None
        raise ValueError(f"raw 模式不支持写入 {type(value).__name__}")

    if isinstance(value, (int, float, bool, datetime, date)):
        return value, "yyyy-mm-dd" if isinstance(value, date) and not isinstance(value, datetime) else None
    if not isinstance(value, str):
        raise ValueError(f"不支持写入 {type(value).__name__} 类型")

    text = value.strip()
    if text == "":
        return None, None
    if text.startswith("="):
        return text, None
    lowered = text.lower()
    if lowered in ("true", "false"):
        return lowered == "true", None

    # 保守推断：前导零数字串、长数字串和混合字符串保持文本。
    if re.match(r"^[+-]?(?:0|[1-9]\d{0,14})$", text):
        signless = text[1:] if text[0] in "+-" else text
        if len(signless) > 1 and signless.startswith("0"):
            return value, None
        return int(text), None
    if re.match(r"^[+-]?(?:0|[1-9]\d*)\.\d+$", text):
        signless = text[1:] if text[0] in "+-" else text
        if len(signless.split(".", 1)[0]) > 1 and signless.startswith("0"):
            return value, None
        return float(text), None
    return value, None


def _cell_change(cell, new_value: Any) -> dict[str, Any]:
    return {
        "cell": cell.coordinate,
        "old_value": _json_value(cell.value),
        "new_value": _json_value(new_value),
    }



def _format_diff_value(value: Any) -> str:
    if value is None:
        return "<blank>"
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return repr(value)


def _compact_changes(changes: list[dict[str, Any]], limit: int = 50) -> list[dict[str, Any]]:
    """返回较小的 changes 列表，避免 metadata 过大。"""
    return changes[:limit]


def _build_xlsx_diff_preview(result: dict[str, Any], limit: int = 20) -> str:
    """构建可读 diff，用于判断改动位置和值是否符合预期。"""
    lines: list[str] = []
    changes = result.get("changes", []) or []
    value_changes = [
        change for change in changes
        if change.get("old_value") != change.get("new_value")
    ]

    if value_changes:
        lines.append("值变更 diff:")
        for change in value_changes[:limit]:
            old_text = _format_diff_value(change.get("old_value"))
            new_text = _format_diff_value(change.get("new_value"))
            lines.append(f"  - {change.get('cell')}: {old_text} -> {new_text}")
        if len(value_changes) > limit:
            lines.append(f"  ... 还有 {len(value_changes) - limit} 个单元格变更未展示")
    else:
        lines.append("值变更 diff: 无单元格值变化")

    row_actions = result.get("row_actions", []) or []
    if row_actions:
        lines.append("结构变更:")
        for action in row_actions[:limit]:
            details = [f"type={action.get('type')}", f"range={action.get('range')}"]
            for key in ("rows_appended", "rows_inserted", "rows_deleted"):
                if action.get(key) is not None:
                    details.append(f"{key}={action[key]}")
            if action.get("table_ref_after"):
                details.append(f"table_ref_after={action['table_ref_after']}")
            lines.append("  - " + ", ".join(details))

    format_actions = result.get("format_actions", []) or []
    if format_actions:
        lines.append("格式/批注结果:")
        for action in format_actions[:limit]:
            details = [f"range={action.get('range')}", f"cells={action.get('cells_formatted')}"]
            if action.get("fill"):
                details.append(f"fill=#{action['fill']}")
            if action.get("font"):
                details.append(f"font={action['font']}")
            if action.get("comment"):
                details.append(f"comment={_format_diff_value(action['comment'])}")
            lines.append("  - " + ", ".join(details))

    return "\n".join(lines)


def _cell_value_matches(actual: Any, expected: Any) -> bool:
    return _json_value(actual) == expected


def _color_matches(color_obj: Any, expected_rgb: str) -> bool:
    actual = getattr(color_obj, "rgb", None)
    if not actual:
        return False
    return str(actual).upper().endswith(expected_rgb.upper())


def _verify_xlsx_persisted(path: Path, sheet_name: str, result: dict[str, Any]) -> dict[str, Any]:
    """保存后重新打开工作簿，核对值、标色和批注是否实际落盘。"""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=False, data_only=False)
        ws = wb[sheet_name]
        checked = 0
        mismatches: list[str] = []

        for change in result.get("changes", []) or []:
            cell_ref = change.get("cell")
            if not cell_ref:
                continue
            checked += 1
            actual = ws[cell_ref].value
            expected = change.get("new_value")
            if not _cell_value_matches(actual, expected):
                mismatches.append(
                    f"{cell_ref}: expected {_format_diff_value(expected)}, got {_format_diff_value(_json_value(actual))}"
                )

        format_checked = 0
        for action in result.get("format_actions", []) or []:
            bounds = _parse_a1_range(action["range"])
            min_col, min_row, max_col, max_row = bounds
            for row_idx in range(min_row, max_row + 1):
                for col_idx in range(min_col, max_col + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if _is_merged_non_anchor(ws, row_idx, col_idx):
                        continue
                    format_checked += 1
                    if action.get("fill") and not _color_matches(cell.fill.fgColor, action["fill"]):
                        mismatches.append(f"{cell.coordinate}: fill 未落盘为 #{action['fill']}")
                    if action.get("comment"):
                        if cell.comment is None or cell.comment.text != str(action["comment"]):
                            mismatches.append(f"{cell.coordinate}: 批注未落盘或内容不一致")
                    font_payload = action.get("font") or {}
                    if "bold" in font_payload and bool(cell.font.bold) != bool(font_payload["bold"]):
                        mismatches.append(f"{cell.coordinate}: bold 字体状态不一致")
                    if "italic" in font_payload and bool(cell.font.italic) != bool(font_payload["italic"]):
                        mismatches.append(f"{cell.coordinate}: italic 字体状态不一致")
                    if "color" in font_payload and not _color_matches(cell.font.color, _normalize_color(font_payload["color"])):
                        mismatches.append(f"{cell.coordinate}: font color 未落盘为 #{_normalize_color(font_payload['color'])}")

        wb.close()
        return {
            "persisted": not mismatches,
            "cells_checked": checked,
            "format_cells_checked": format_checked,
            "mismatches": mismatches[:20],
            "mismatch_count": len(mismatches),
        }
    except Exception as exc:
        return {
            "persisted": False,
            "cells_checked": 0,
            "format_cells_checked": 0,
            "mismatches": [f"落盘校验失败: {type(exc).__name__}: {exc}"],
            "mismatch_count": 1,
        }


def _write_cell(cell, value: Any, kind: str | None = None, value_input: str = "auto") -> dict[str, Any]:
    new_value, number_format = _coerce_cell_value(value, kind=kind, value_input=value_input)
    change = _cell_change(cell, new_value)
    cell.value = new_value
    if number_format:
        cell.number_format = number_format
    return change


def _normalize_color(value: str) -> str:
    if not isinstance(value, str):
        raise ValueError("颜色必须是字符串")
    text = value.strip().lstrip("#").upper()
    if not re.match(r"^[0-9A-F]{6}$", text):
        raise ValueError(f"不合法的颜色: {value!r}，需要 RRGGBB 或 #RRGGBB")
    return text


def _find_merged_conflicts(ws, bounds: tuple[int, int, int, int]) -> list[str]:
    conflicts: list[str] = []
    for merged in ws.merged_cells.ranges:
        merged_bounds = _merged_range_bounds(merged)
        if _bounds_overlap(bounds, merged_bounds) and not _bounds_contains(bounds, merged_bounds):
            conflicts.append(f"目标区域 {_range_from_bounds(*bounds)} 与合并区域 {merged} 部分重叠")
    return conflicts


def _formula_cells(ws) -> list[str]:
    cells: list[str] = []
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row or 1, min_col=1, max_col=ws.max_column or 1):
        for cell in row:
            if isinstance(cell.value, str) and cell.value.startswith("="):
                cells.append(f"{ws.title}!{cell.coordinate}")
    return cells


def _analyze_range_impact(wb, ws, bounds_list: list[tuple[int, int, int, int]], operation: str) -> dict[str, Any]:
    """对目标区域做保守影响分析。"""
    warnings: list[str] = []
    errors: list[str] = []
    impacted_tables: list[str] = []
    affected_ranges = [_range_from_bounds(*bounds) for bounds in bounds_list]

    table_infos = _worksheet_tables(ws)
    for bounds in bounds_list:
        errors.extend(_find_merged_conflicts(ws, bounds))
        for table_info in table_infos:
            table_bounds = _parse_a1_range(table_info["ref"])
            if _bounds_overlap(bounds, table_bounds):
                impacted_tables.append(table_info["name"])
                if operation in _XLSX_ROW_STRUCTURE_OPS:
                    errors.append(
                        f"{operation} 会影响 native Excel table {table_info['name']} ({table_info['ref']})"
                    )
                elif operation not in {"append_rows", "update_rows_by_key", "upsert_rows"}:
                    warnings.append(
                        f"目标区域与 native Excel table {table_info['name']} ({table_info['ref']}) 重叠"
                    )

    if operation in _XLSX_ROW_STRUCTURE_OPS:
        formulas = _formula_cells(ws)
        if formulas:
            preview = ", ".join(formulas[:5])
            more = "..." if len(formulas) > 5 else ""
            errors.append(f"结构性行操作所在工作表包含公式单元格: {preview}{more}")
        if getattr(ws, "_charts", None):
            errors.append("结构性行操作所在工作表包含图表，当前编辑器不会自动修复图表数据源")
        try:
            if ws.data_validations and ws.data_validations.dataValidation:
                warnings.append("工作表包含数据验证，结构性行操作可能影响验证区域")
        except Exception:
            pass
        try:
            if len(ws.conditional_formatting):
                warnings.append("工作表包含条件格式，结构性行操作可能影响条件格式区域")
        except Exception:
            pass

    return {
        "affected_ranges": affected_ranges,
        "warnings": sorted(set(warnings)),
        "errors": sorted(set(errors)),
        "impacted_tables": sorted(set(impacted_tables)),
    }


def _table_headers(ws, table) -> list[str]:
    min_col, min_row, max_col, _ = _parse_a1_range(table.ref)
    headers: list[str] = []
    for col_idx in range(min_col, max_col + 1):
        value = ws.cell(row=min_row, column=col_idx).value
        headers.append(str(value).strip() if value is not None else "")
    return headers


def _validate_headers(headers: list[str]) -> str | None:
    seen: set[str] = set()
    for header in headers:
        if not header:
            return "表头包含空列名，无法按字段名定位。"
        if header in seen:
            return f"表头 {header!r} 重复，无法按字段名定位。"
        seen.add(header)
    return None


def _lookup_header(headers: list[str], key: str) -> int:
    if key in headers:
        return headers.index(key)
    lowered = {header.lower(): idx for idx, header in enumerate(headers)}
    idx = lowered.get(str(key).lower())
    if idx is not None:
        return idx
    raise ValueError(f"列 {key!r} 不存在。可用列: {headers}")


def _resolve_dataset(ws, payload: dict[str, Any]) -> dict[str, Any]:
    table_name = payload.get("table_name")
    if table_name:
        table = _table_by_name(ws, table_name)
        if table is None:
            raise ValueError(f"工作表 {ws.title!r} 中不存在表格 {table_name!r}")
        min_col, min_row, max_col, max_row = _parse_a1_range(table.ref)
        totals = int(getattr(table, "totalsRowCount", 0) or 0)
        headers = _table_headers(ws, table)
        header_error = _validate_headers(headers)
        if header_error:
            raise ValueError(header_error)
        return {
            "kind": "table",
            "table": table,
            "table_name": table_name,
            "headers": headers,
            "min_col": min_col,
            "max_col": max_col,
            "header_row": min_row,
            "data_min_row": min_row + 1,
            "data_max_row": max_row - totals,
            "append_row": max_row + 1,
            "totals_row_count": totals,
            "ref_bounds": (min_col, min_row, max_col, max_row),
        }

    header_row = int(payload.get("header_row", 1))
    start_col = _column_index(payload.get("start_col", "A"))
    end_col = _column_index(payload["end_col"]) if payload.get("end_col") else (ws.max_column or start_col)
    headers = []
    for col_idx in range(start_col, end_col + 1):
        value = ws.cell(row=header_row, column=col_idx).value
        headers.append(str(value).strip() if value is not None else "")
    header_error = _validate_headers(headers)
    if header_error:
        raise ValueError(header_error)

    data_min_row = header_row + 1
    data_max_row = max(ws.max_row or header_row, data_min_row - 1)
    append_row = data_max_row + 1
    return {
        "kind": "worksheet",
        "table": None,
        "table_name": None,
        "headers": headers,
        "min_col": start_col,
        "max_col": end_col,
        "header_row": header_row,
        "data_min_row": data_min_row,
        "data_max_row": data_max_row,
        "append_row": append_row,
        "totals_row_count": 0,
        "ref_bounds": (start_col, header_row, end_col, data_max_row),
    }


def _records_to_matrix(rows: list[Any], headers: list[str]) -> list[list[Any]]:
    matrix: list[list[Any]] = []
    for row in rows:
        if isinstance(row, dict):
            matrix.append([row.get(header) for header in headers])
        elif isinstance(row, list):
            if len(row) > len(headers):
                raise ValueError(f"行列数({len(row)})超过表头列数({len(headers)})")
            matrix.append(row + [None] * (len(headers) - len(row)))
        else:
            raise ValueError("rows 中的每一行必须是 object 或 array")
    return matrix


def _cell_bounds(row: int, col: int) -> tuple[int, int, int, int]:
    return col, row, col, row


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit_table — XLSX 专用编辑工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _copy_row_style(ws, source_row: int, target_row: int, max_col: int) -> None:
    if source_row < 1:
        return
    for col_idx in range(1, max_col + 1):
        source = ws.cell(row=source_row, column=col_idx)
        target = ws.cell(row=target_row, column=col_idx)
        if source.has_style:
            target._style = copy(source._style)
        if source.number_format:
            target.number_format = source.number_format
        if source.protection:
            target.protection = copy(source.protection)
        if source.alignment:
            target.alignment = copy(source.alignment)


def _plan_set_range(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    cell_range = payload.get("range")
    values = payload.get("values")
    if not cell_range:
        raise ValueError("set_range payload 必须包含 range")
    if not isinstance(values, list) or not values or not all(isinstance(row, list) for row in values):
        raise ValueError("set_range payload.values 必须是非空二维数组")

    row_lengths = {len(row) for row in values}
    if len(row_lengths) != 1:
        raise ValueError("set_range payload.values 必须是矩形二维数组")
    value_rows = len(values)
    value_cols = next(iter(row_lengths)) if row_lengths else 0
    if value_cols == 0:
        raise ValueError("set_range payload.values 至少需要一列")

    min_col, min_row, max_col, max_row = _parse_a1_range(cell_range)
    resize_range = bool(payload.get("resize_range", False))
    if resize_range:
        max_row = min_row + value_rows - 1
        max_col = min_col + value_cols - 1
    else:
        range_rows = max_row - min_row + 1
        range_cols = max_col - min_col + 1
        if range_rows != value_rows or range_cols != value_cols:
            raise ValueError(
                f"range 大小({range_rows}x{range_cols})与 values 大小({value_rows}x{value_cols})不一致"
            )

    value_input = payload.get("value_input", "auto")
    kind = payload.get("kind")
    changes: list[dict[str, Any]] = []
    for r_offset, row_values in enumerate(values):
        row_idx = min_row + r_offset
        for c_offset, raw_value in enumerate(row_values):
            col_idx = min_col + c_offset
            merged = _is_merged_non_anchor(ws, row_idx, col_idx)
            if merged:
                raise ValueError(f"不能写入合并区域 {merged} 的非左上角单元格")
            cell = ws.cell(row=row_idx, column=col_idx)
            new_value, _ = _coerce_cell_value(raw_value, kind=kind, value_input=value_input)
            changes.append(_cell_change(cell, new_value))
            if apply:
                _write_cell(cell, raw_value, kind=kind, value_input=value_input)

    bounds = (min_col, min_row, max_col, max_row)
    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": changes,
        "counts": {"cells_updated": value_rows * value_cols},
        "structural": False,
    }


def _plan_update_cells(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    cells = payload.get("cells")
    if not isinstance(cells, list) or not cells:
        raise ValueError("update_cells payload.cells 必须是非空数组")

    seen: set[str] = set()
    bounds: list[tuple[int, int, int, int]] = []
    changes: list[dict[str, Any]] = []
    value_input = payload.get("value_input", "auto")

    for item in cells:
        if not isinstance(item, dict):
            raise ValueError("update_cells 的每个 cell 项必须是 object")
        address = item.get("cell")
        if not address:
            raise ValueError("update_cells 的每个 cell 项必须包含 cell")
        row_idx, col_idx = _parse_cell_address(address)
        normalized = ws.cell(row=row_idx, column=col_idx).coordinate
        if normalized in seen:
            raise ValueError(f"重复更新单元格 {normalized}")
        seen.add(normalized)
        kind = item.get("kind")
        if "value" not in item and kind != "blank":
            raise ValueError(f"单元格 {normalized} 缺少 value")
        merged = _is_merged_non_anchor(ws, row_idx, col_idx)
        if merged:
            raise ValueError(f"不能写入合并区域 {merged} 的非左上角单元格")
        raw_value = item.get("value")
        cell = ws.cell(row=row_idx, column=col_idx)
        new_value, _ = _coerce_cell_value(raw_value, kind=kind, value_input=value_input)
        changes.append(_cell_change(cell, new_value))
        if apply:
            _write_cell(cell, raw_value, kind=kind, value_input=value_input)
        bounds.append(_cell_bounds(row_idx, col_idx))

    return {
        "bounds": bounds,
        "affected_ranges": [_range_from_bounds(*b) for b in bounds],
        "changes": changes,
        "counts": {"cells_updated": len(cells)},
        "structural": False,
    }


def _plan_clear_range(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    cell_range = payload.get("range")
    if not cell_range:
        raise ValueError("clear_range payload 必须包含 range")
    clear = payload.get("clear", "values")
    if clear != "values":
        raise ValueError("当前仅支持 clear='values'")

    bounds = _parse_a1_range(cell_range)
    min_col, min_row, max_col, max_row = bounds
    changes: list[dict[str, Any]] = []
    for row_idx in range(min_row, max_row + 1):
        for col_idx in range(min_col, max_col + 1):
            merged = _is_merged_non_anchor(ws, row_idx, col_idx)
            if merged:
                raise ValueError(f"不能清空合并区域 {merged} 的非左上角单元格")
            cell = ws.cell(row=row_idx, column=col_idx)
            changes.append(_cell_change(cell, None))
            if apply:
                cell.value = None

    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": changes,
        "counts": {"cells_cleared": (max_row - min_row + 1) * (max_col - min_col + 1)},
        "structural": False,
    }


def _plan_format_range(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    from openpyxl.comments import Comment as XlComment
    from openpyxl.styles import PatternFill

    cell_range = payload.get("range")
    if not cell_range:
        raise ValueError("format_range payload 必须包含 range")
    bounds = _parse_a1_range(cell_range)
    min_col, min_row, max_col, max_row = bounds

    fill = payload.get("fill")
    font_payload = payload.get("font") or {}
    comment = payload.get("comment")
    fill_style = None
    if fill:
        color = _normalize_color(fill)
        fill_style = PatternFill(start_color=color, end_color=color, fill_type="solid")
    if font_payload and not isinstance(font_payload, dict):
        raise ValueError("format_range font 必须是 object")

    formatted = 0
    for row_idx in range(min_row, max_row + 1):
        for col_idx in range(min_col, max_col + 1):
            if _is_merged_non_anchor(ws, row_idx, col_idx):
                continue
            cell = ws.cell(row=row_idx, column=col_idx)
            if apply:
                if fill_style:
                    cell.fill = fill_style
                if font_payload:
                    new_font = copy(cell.font)
                    if "bold" in font_payload:
                        new_font.bold = bool(font_payload["bold"])
                    if "italic" in font_payload:
                        new_font.italic = bool(font_payload["italic"])
                    if "color" in font_payload:
                        new_font.color = _normalize_color(font_payload["color"])
                    cell.font = new_font
                if comment:
                    cell.comment = XlComment(str(comment), "Agent")
            formatted += 1

    format_action = {
        "range": _range_from_bounds(*bounds),
        "cells_formatted": formatted,
    }
    if fill:
        format_action["fill"] = _normalize_color(fill)
    if font_payload:
        format_action["font"] = dict(font_payload)
    if comment:
        format_action["comment"] = str(comment)

    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": [],
        "format_actions": [format_action],
        "counts": {"cells_formatted": formatted},
        "structural": False,
    }


def _plan_append_rows(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    rows = payload.get("rows")
    if not isinstance(rows, list) or not rows:
        raise ValueError("append_rows payload.rows 必须是非空数组")
    dataset = _resolve_dataset(ws, payload)
    if dataset["totals_row_count"]:
        raise ValueError("目标 Excel table 包含 totals row，当前不支持 append_rows")

    matrix = _records_to_matrix(rows, dataset["headers"])
    append_row = dataset["append_row"]
    min_col = dataset["min_col"]
    max_col = dataset["max_col"]
    max_row = append_row + len(matrix) - 1
    bounds = (min_col, append_row, max_col, max_row)

    for row_idx in range(append_row, max_row + 1):
        for col_idx in range(min_col, max_col + 1):
            if ws.cell(row=row_idx, column=col_idx).value not in (None, ""):
                raise ValueError(f"追加区域 {_range_from_bounds(*bounds)} 覆盖了非空单元格")

    changes: list[dict[str, Any]] = []
    value_input = payload.get("value_input", "auto")
    for r_offset, row_values in enumerate(matrix):
        row_idx = append_row + r_offset
        for c_offset, raw_value in enumerate(row_values):
            col_idx = min_col + c_offset
            cell = ws.cell(row=row_idx, column=col_idx)
            new_value, _ = _coerce_cell_value(raw_value, value_input=value_input)
            changes.append(_cell_change(cell, new_value))
            if apply:
                _write_cell(cell, raw_value, value_input=value_input)

    new_table_ref = None
    if dataset["table"] is not None:
        new_table_ref = _range_from_bounds(min_col, dataset["header_row"], max_col, max_row)
    if apply and dataset["table"] is not None:
        dataset["table"].ref = new_table_ref

    row_action = {
        "type": "append_rows",
        "range": _range_from_bounds(*bounds),
        "rows_appended": len(matrix),
    }
    if dataset["table_name"]:
        row_action["table_name"] = dataset["table_name"]
    if new_table_ref:
        row_action["table_ref_after"] = new_table_ref

    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": changes,
        "row_actions": [row_action],
        "counts": {"rows_appended": len(matrix)},
        "structural": True,
    }


def _dataset_key_rows(ws, dataset: dict[str, Any], key_column: str) -> tuple[int, dict[Any, int], list[Any]]:
    key_offset = _lookup_header(dataset["headers"], key_column)
    key_col = dataset["min_col"] + key_offset
    key_rows: dict[Any, int] = {}
    duplicates: list[Any] = []
    for row_idx in range(dataset["data_min_row"], dataset["data_max_row"] + 1):
        value = ws.cell(row=row_idx, column=key_col).value
        if value in (None, ""):
            continue
        if value in key_rows:
            duplicates.append(value)
        else:
            key_rows[value] = row_idx
    return key_col, key_rows, duplicates


def _plan_update_rows_by_key(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    key_column = payload.get("key_column")
    updates = payload.get("updates")
    if not key_column:
        raise ValueError("update_rows_by_key payload 必须包含 key_column")
    if not isinstance(updates, list) or not updates:
        raise ValueError("update_rows_by_key payload.updates 必须是非空数组")

    missing = payload.get("missing", "error")
    if missing not in ("error", "ignore", "append"):
        raise ValueError("missing 可选: error, ignore, append")

    dataset = _resolve_dataset(ws, payload)
    if missing == "append" and dataset["totals_row_count"]:
        raise ValueError("目标 Excel table 包含 totals row，当前不支持 missing='append'")
    _, key_rows, duplicates = _dataset_key_rows(ws, dataset, str(key_column))
    if duplicates:
        raise ValueError(f"关键列 {key_column!r} 存在重复 key: {duplicates[:5]}")

    changes: list[dict[str, Any]] = []
    bounds: list[tuple[int, int, int, int]] = []
    appended_records: list[dict[str, Any]] = []
    rows_updated = 0
    rows_ignored = 0
    value_input = payload.get("value_input", "auto")

    for update in updates:
        if not isinstance(update, dict) or "key" not in update or not isinstance(update.get("values"), dict):
            raise ValueError("updates 中每一项必须包含 key 和 values object")
        key = update["key"]
        values = update["values"]
        row_idx = key_rows.get(key)
        if row_idx is None:
            if missing == "ignore":
                rows_ignored += 1
                continue
            if missing == "append":
                record = {str(key_column): key}
                record.update(values)
                appended_records.append(record)
                continue
            raise ValueError(f"关键列 {key_column!r} 未找到 key={key!r}")

        touched_cols: list[int] = []
        for header, raw_value in values.items():
            offset = _lookup_header(dataset["headers"], str(header))
            col_idx = dataset["min_col"] + offset
            cell = ws.cell(row=row_idx, column=col_idx)
            new_value, _ = _coerce_cell_value(raw_value, value_input=value_input)
            changes.append(_cell_change(cell, new_value))
            if apply:
                _write_cell(cell, raw_value, value_input=value_input)
            touched_cols.append(col_idx)
        if touched_cols:
            bounds.append((min(touched_cols), row_idx, max(touched_cols), row_idx))
            rows_updated += 1

    rows_appended = 0
    if appended_records:
        append_payload = dict(payload)
        append_payload["rows"] = appended_records
        append_payload.pop("updates", None)
        append_plan = _plan_append_rows(ws, append_payload, apply=apply, include_changes=include_changes)
        bounds.extend(append_plan["bounds"])
        changes.extend(append_plan["changes"])
        appended_row_actions = append_plan.get("row_actions", [])
        rows_appended = append_plan["counts"].get("rows_appended", 0)
    else:
        appended_row_actions = []

    return {
        "bounds": bounds,
        "affected_ranges": [_range_from_bounds(*b) for b in bounds],
        "changes": changes,
        "row_actions": appended_row_actions,
        "counts": {
            "rows_updated": rows_updated,
            "rows_appended": rows_appended,
            "rows_ignored": rows_ignored,
        },
        "structural": bool(appended_records),
    }


def _plan_upsert_rows(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    key_column = payload.get("key_column")
    rows = payload.get("rows")
    if not key_column:
        raise ValueError("upsert_rows payload 必须包含 key_column")
    if not isinstance(rows, list) or not rows:
        raise ValueError("upsert_rows payload.rows 必须是非空数组")
    updates = []
    for row in rows:
        if not isinstance(row, dict):
            raise ValueError("upsert_rows rows 中每一项必须是 object")
        if key_column not in row:
            raise ValueError(f"upsert row 缺少 key_column {key_column!r}")
        values = dict(row)
        key = values.pop(key_column)
        updates.append({"key": key, "values": values})
    update_payload = dict(payload)
    update_payload["updates"] = updates
    update_payload["missing"] = "append"
    return _plan_update_rows_by_key(ws, update_payload, apply=apply, include_changes=include_changes)


def _plan_delete_rows(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    rows = payload.get("rows")
    if rows is None and payload.get("keys") is not None:
        key_column = payload.get("key_column")
        if not key_column:
            raise ValueError("按 key 删除时必须提供 key_column")
        keys = payload.get("keys")
        if not isinstance(keys, list) or not keys:
            raise ValueError("delete_rows payload.keys 必须是非空数组")
        dataset = _resolve_dataset(ws, payload)
        _, key_rows, duplicates = _dataset_key_rows(ws, dataset, str(key_column))
        if duplicates:
            raise ValueError(f"关键列 {key_column!r} 存在重复 key: {duplicates[:5]}")
        rows = []
        for key in keys:
            row_idx = key_rows.get(key)
            if row_idx is None:
                raise ValueError(f"关键列 {key_column!r} 未找到 key={key!r}")
            rows.append(row_idx)
    if not isinstance(rows, list) or not rows:
        raise ValueError("delete_rows payload 必须包含非空 rows 或 keys")

    row_nums = sorted({int(row) for row in rows})
    if any(row < 1 for row in row_nums):
        raise ValueError("删除行号必须从 1 开始")
    max_col = ws.max_column or 1
    bounds = [(1, row, max_col, row) for row in row_nums]
    if apply:
        for row in sorted(row_nums, reverse=True):
            ws.delete_rows(row, 1)
    return {
        "bounds": bounds,
        "affected_ranges": [_range_from_bounds(*b) for b in bounds],
        "changes": [],
        "row_actions": [{
            "type": "delete_rows",
            "range": ",".join(_range_from_bounds(*b) for b in bounds),
            "rows_deleted": len(row_nums),
            "rows": row_nums,
        }],
        "counts": {"rows_deleted": len(row_nums)},
        "structural": True,
    }


def _plan_insert_rows(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    rows = payload.get("rows")
    if not isinstance(rows, list) or not rows:
        raise ValueError("insert_rows payload.rows 必须是非空数组")
    if payload.get("before_row") is not None:
        insert_at = int(payload["before_row"])
    elif payload.get("after_row") is not None:
        insert_at = int(payload["after_row"]) + 1
    else:
        raise ValueError("insert_rows payload 必须包含 before_row 或 after_row")
    if insert_at < 1:
        raise ValueError("插入行号必须从 1 开始")

    if all(isinstance(row, dict) for row in rows):
        dataset = _resolve_dataset(ws, payload)
        matrix = _records_to_matrix(rows, dataset["headers"])
    elif all(isinstance(row, list) for row in rows):
        matrix = rows
    else:
        raise ValueError("insert_rows rows 必须全部是 object 或全部是 array")

    value_input = payload.get("value_input", "auto")
    max_width = max(len(row) for row in matrix)
    bounds = (1, insert_at, max(max_width, ws.max_column or 1), insert_at + len(matrix) - 1)

    if apply:
        ws.insert_rows(insert_at, len(matrix))
        copy_style_from = payload.get("copy_style_from", "none")
        if copy_style_from not in ("above", "below", "none"):
            raise ValueError("copy_style_from 可选: above, below, none")
        if copy_style_from == "above":
            template_row = insert_at - 1
        elif copy_style_from == "below":
            template_row = insert_at + len(matrix)
        else:
            template_row = 0
        if template_row:
            for offset in range(len(matrix)):
                _copy_row_style(ws, template_row, insert_at + offset, bounds[2])
        for r_offset, row_values in enumerate(matrix):
            for c_offset, raw_value in enumerate(row_values):
                cell = ws.cell(row=insert_at + r_offset, column=1 + c_offset)
                _write_cell(cell, raw_value, value_input=value_input)

    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": [],
        "row_actions": [{
            "type": "insert_rows",
            "range": _range_from_bounds(*bounds),
            "rows_inserted": len(matrix),
            "before_row": insert_at,
        }],
        "counts": {"rows_inserted": len(matrix)},
        "structural": True,
    }


def _run_xlsx_operation(
    wb,
    ws,
    operation: str,
    payload: dict[str, Any],
    apply: bool,
    include_changes: bool,
) -> dict[str, Any]:
    if operation == "set_range":
        result = _plan_set_range(ws, payload, apply, include_changes)
    elif operation == "update_cells":
        result = _plan_update_cells(ws, payload, apply, include_changes)
    elif operation == "clear_range":
        result = _plan_clear_range(ws, payload, apply, include_changes)
    elif operation == "format_range":
        result = _plan_format_range(ws, payload, apply, include_changes)
    elif operation == "append_rows":
        result = _plan_append_rows(ws, payload, apply, include_changes)
    elif operation == "update_rows_by_key":
        result = _plan_update_rows_by_key(ws, payload, apply, include_changes)
    elif operation == "upsert_rows":
        result = _plan_upsert_rows(ws, payload, apply, include_changes)
    elif operation == "delete_rows":
        result = _plan_delete_rows(ws, payload, apply, include_changes)
    elif operation == "insert_rows":
        result = _plan_insert_rows(ws, payload, apply, include_changes)
    else:
        raise ValueError(
            "不支持的 operation: "
            f"{operation!r}。可选: set_range, update_cells, clear_range, format_range, "
            "append_rows, update_rows_by_key, upsert_rows, delete_rows, insert_rows"
        )
    result.setdefault("bounds", [])
    result.setdefault("affected_ranges", [])
    result.setdefault("changes", [])
    result.setdefault("format_actions", [])
    result.setdefault("row_actions", [])
    result.setdefault("counts", {})
    result.setdefault("structural", operation in _XLSX_STRUCTURAL_OPS)
    return result


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit_table — XLSX 结构化编辑工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@tool(name="file_edit_table",
      description=(
          "结构化编辑 XLSX 表格文件。使用 operation + payload 表达明确动作，"
          "不再支持 target_content/replacement_content 文本行匹配。\n\n"
          "支持 operation：\n"
          "- set_range: 按 A1 range 写入二维数组\n"
          "- update_cells: 按单元格地址批量更新\n"
          "- clear_range: 清空区域值并保留样式\n"
          "- format_range: 设置标色、字体和批注\n"
          "- append_rows: 向 native Excel table 或 header dataset 追加行\n"
          "- update_rows_by_key: 按关键列更新匹配行\n"
          "- upsert_rows: 按关键列更新并追加缺失行\n"
          "- delete_rows: 删除指定行或 key 匹配行\n"
          "- insert_rows: 在指定位置插入行\n\n"
          "默认 dry_run=True 只预览不保存。结构性写入需 dry_run=False 且 allow_structure_change=True。"
      ))
async def file_edit_table(
    path: str,
    operation: str,
    sheet_name: str | None = None,
    payload: dict = None,
    dry_run: bool = True,
    allow_structure_change: bool = False,
    expected_structure_token: str | None = None,
    expected_content_token: str | None = None,
    include_changes: bool = False,
) -> ToolResult:
    """
    结构化编辑 XLSX 表格文件。

    Args:
        path: 文件路径，仅支持 .xlsx
        operation: 编辑动作。可选 set_range/update_cells/clear_range/format_range/append_rows/update_rows_by_key/upsert_rows/delete_rows/insert_rows
        sheet_name: 工作表名。单 Sheet 自动选择；table_name 可唯一定位时也可省略
        payload: operation 对应的结构化参数
        dry_run: 是否只预览不保存，默认 True
        allow_structure_change: 是否允许插入/删除/追加等结构性写入，默认 False
        expected_structure_token: 可选结构 token，不匹配时拒绝写入
        expected_content_token: 可选内容 token，不匹配时拒绝写入
        include_changes: 是否在 metadata 中返回单元格级 changes
    """
    try:
        import openpyxl
    except ImportError:
        return ToolResult(
            content="缺少 openpyxl 库。请安装: pip install openpyxl",
            is_error=True,
        )

    operation = (operation or "").strip().lower()
    payload = payload or {}
    if not isinstance(payload, dict):
        return ToolResult(
            content="payload 必须是 object。",
            is_error=True,
            metadata={"ok": False, "operation": operation, "dry_run": dry_run},
        )

    error = _check_path_safety(path)
    if error:
        return ToolResult(content=error, is_error=True)

    target = Path(path)
    if not target.exists():
        return ToolResult(content=f"文件不存在: {path}", is_error=True)
    if not target.is_file():
        return ToolResult(content=f"不是文件: {path}", is_error=True)

    ext = target.suffix.lower()
    if ext == ".xls":
        return ToolResult(
            content="不支持旧版 .xls 格式。请转换为 .xlsx 后重试。",
            is_error=True,
        )
    if ext != ".xlsx":
        return ToolResult(
            content=f"file_edit_table 仅支持 .xlsx 文件，当前文件类型: {ext or '(无扩展名)'}。",
            is_error=True,
        )

    logger.info("file_edit_table 开始: path=%s, operation=%s, dry_run=%s", path, operation, dry_run)

    try:
        wb = openpyxl.load_workbook(str(path), data_only=False)
    except Exception as e:
        return ToolResult(content=f"打开 XLSX 失败: {e}", is_error=True)

    try:
        ws, sheet_error = _resolve_worksheet_for_table_edit(wb, sheet_name, payload)
        if sheet_error:
            wb.close()
            return ToolResult(
                content=sheet_error,
                is_error=True,
                metadata={"ok": False, "operation": operation, "dry_run": dry_run},
            )

        previous_profile = _profile_workbook(wb)
        previous_structure_token = _hash_json(previous_profile)

        # 先 dry-plan 一次：计算影响范围、校验 payload，不写入 workbook。
        plan = _run_xlsx_operation(
            wb, ws, operation, payload, apply=False, include_changes=include_changes
        )
        bounds = plan.get("bounds", [])
        previous_content_token = _hash_json([_content_token(ws, bound) for bound in bounds]) if bounds else _content_token(ws)

        if expected_structure_token and expected_structure_token != previous_structure_token:
            wb.close()
            return ToolResult(
                content="结构 token 不匹配，文件结构可能已变化，请重新 file_read 后再编辑。",
                is_error=True,
                metadata={
                    "ok": False,
                    "operation": operation,
                    "dry_run": dry_run,
                    "expected_structure_token": expected_structure_token,
                    "current_structure_token": previous_structure_token,
                },
            )
        if expected_content_token and expected_content_token != previous_content_token:
            wb.close()
            return ToolResult(
                content="内容 token 不匹配，目标区域可能已变化，请重新 file_read 后再编辑。",
                is_error=True,
                metadata={
                    "ok": False,
                    "operation": operation,
                    "dry_run": dry_run,
                    "expected_content_token": expected_content_token,
                    "current_content_token": previous_content_token,
                },
            )

        is_structural = bool(plan.get("structural")) or operation in _XLSX_STRUCTURAL_OPS
        if is_structural and not dry_run and not allow_structure_change:
            wb.close()
            return ToolResult(
                content=(
                    f"{operation} 属于结构性操作。请先 dry_run 预览，确认后设置 "
                    "dry_run=False 且 allow_structure_change=True。"
                ),
                is_error=True,
                metadata={
                    "ok": False,
                    "operation": operation,
                    "dry_run": dry_run,
                    "data": {
                        "sheet": ws.title,
                        "affected_ranges": plan.get("affected_ranges", []),
                        "requires_structure_change": True,
                    },
                },
            )

        impact = _analyze_range_impact(wb, ws, bounds, operation)
        warnings = sorted(set(plan.get("warnings", []) + impact.get("warnings", [])))
        if is_structural and dry_run and not allow_structure_change:
            warnings.append("该操作属于结构性变更；实际写入时需要 allow_structure_change=True")

        if impact.get("errors"):
            wb.close()
            message = "操作被阻止：" + "; ".join(impact["errors"][:5])
            return ToolResult(
                content=message,
                is_error=True,
                metadata={
                    "ok": False,
                    "operation": operation,
                    "dry_run": dry_run,
                    "error": {"type": "ImpactError", "message": message},
                    "data": {
                        "path": str(target.resolve()),
                        "sheet": ws.title,
                        "affected_ranges": plan.get("affected_ranges", []),
                        "warnings": warnings,
                        "impact": impact,
                    },
                },
            )

        if dry_run:
            diff_preview = _build_xlsx_diff_preview(plan)
            wb.close()
            message = (
                f"XLSX 表格编辑预览: {target.name}，operation={operation}，"
                f"将影响 {len(plan.get('affected_ranges', []))} 个区域。\n"
                f"{diff_preview}"
            )
            return ToolResult(
                content=message,
                metadata={
                    "ok": True,
                    "operation": operation,
                    "dry_run": True,
                    "data": {
                        "path": str(target.resolve()),
                        "sheet": ws.title,
                        "affected_ranges": plan.get("affected_ranges", []),
                        "counts": plan.get("counts", {}),
                        "warnings": warnings,
                        "diff_preview": diff_preview,
                        "changes": plan.get("changes", []) if include_changes else _compact_changes(plan.get("changes", [])),
                        "change_count": len(plan.get("changes", [])),
                        "format_actions": plan.get("format_actions", []),
                        "row_actions": plan.get("row_actions", []),
                        "requires_structure_change": is_structural,
                        "previous_structure_token": previous_structure_token,
                        "previous_content_token": previous_content_token,
                        "impact": impact,
                    },
                },
            )

        applied = _run_xlsx_operation(
            wb, ws, operation, payload, apply=True, include_changes=include_changes
        )
        new_profile = _profile_workbook(wb)
        new_structure_token = _hash_json(new_profile)
        new_content_token = _hash_json([_content_token(ws, bound) for bound in applied.get("bounds", [])]) if applied.get("bounds") else _content_token(ws)

        sheet_title = ws.title
        try:
            buf = io.BytesIO()
            wb.save(buf)
            wb.close()
            _atomic_write_binary(target, buf.getvalue())
        except Exception as e:
            return ToolResult(content=f"保存 XLSX 失败: {e}", is_error=True)

        persistence = _verify_xlsx_persisted(target, sheet_title, applied)
        diff_preview = _build_xlsx_diff_preview(applied)
        counts = applied.get("counts", {})
        count_summary = ", ".join(f"{k}={v}" for k, v in counts.items() if v) or "已完成"
        persisted_text = "落盘校验通过" if persistence.get("persisted") else "落盘校验失败"
        message = (
            f"XLSX 表格编辑成功: {target.name}，operation={operation}，{count_summary}。"
            f" {persisted_text}。\n{diff_preview}"
        )
        if warnings:
            message += f"\n警告: {'; '.join(warnings[:5])}"
        if not persistence.get("persisted"):
            message += "\n校验问题: " + "; ".join(persistence.get("mismatches", [])[:5])

        logger.info("file_edit_table 成功: path=%s, operation=%s", target, operation)
        return ToolResult(
            content=message,
            metadata={
                "ok": True,
                "operation": operation,
                "dry_run": False,
                "data": {
                    "path": str(target.resolve()),
                    "sheet": sheet_title,
                    "affected_ranges": applied.get("affected_ranges", []),
                    "counts": counts,
                    "warnings": warnings,
                    "diff_preview": diff_preview,
                    "persistence_verification": persistence,
                    "changes": applied.get("changes", []) if include_changes else _compact_changes(applied.get("changes", [])),
                    "change_count": len(applied.get("changes", [])),
                    "format_actions": applied.get("format_actions", []),
                    "row_actions": applied.get("row_actions", []),
                    "previous_structure_token": previous_structure_token,
                    "new_structure_token": new_structure_token,
                    "previous_content_token": previous_content_token,
                    "new_content_token": new_content_token,
                },
            },
        )
    except Exception as e:
        try:
            wb.close()
        except Exception:
            pass
        return ToolResult(
            content=f"XLSX 表格编辑失败: {type(e).__name__}: {e}",
            is_error=True,
            metadata={"ok": False, "operation": operation, "dry_run": dry_run},
        )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit — 纯文本 + DOCX 编辑
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@tool(name="file_edit",
      description=(
          "精确编辑已有文件。支持纯文本文件和 DOCX 两种类型。\n"
          "通过 target_content 精确匹配原始内容，替换为 replacement_content。\n"
          "支持对修改位置进行标色（highlight）和添加批注（comment）。\n\n"
          "使用建议：\n"
          "- 修改已有文件优先使用 file_edit，而非 file_write\n"
          "- target_content 必须精确匹配原文（含缩进和空白）\n"
          "- 匹配失败时先 file_read 确认最新内容再重试\n"
          "- 不支持 .doc 旧格式\n"
          "- 编辑 XLSX 请使用 file_edit_table"
      ))
async def file_edit(
    path: str,
    target_content: str,
    replacement_content: str,
    start_line: int | None = None,
    end_line: int | None = None,
    allow_multiple: bool = False,
    highlight: str | None = None,
    comment: str | None = None,
) -> ToolResult:
    """
    精确编辑已有文件（纯文本 / DOCX）。

    Args:
        path: 文件路径，支持文本/DOCX
        target_content: 要被替换的原始文本（精确匹配，必填）
        replacement_content: 替换后的新文本（必填）
        start_line: 纯文本文件搜索范围起始行（1-based）
        end_line: 纯文本文件搜索范围结束行
        allow_multiple: 是否允许替换多个匹配（默认 False）
        highlight: 标色颜色: "yellow"/"green"/"red"/"pink"/None
        comment: 批注内容，None=不添加批注
    """
    # ── highlight 值校验 ──
    if highlight is not None and highlight not in ("yellow", "green", "red", "pink"):
        return ToolResult(
            content=f"不支持的 highlight 颜色: {highlight!r}。可选: yellow, green, red, pink",
            is_error=True,
        )

    logger.info("file_edit 开始: path=%s", path)

    # ── 文件存在性检查 ──
    target = Path(path)
    if not target.exists():
        return ToolResult(content=f"文件不存在: {path}", is_error=True)
    if not target.is_file():
        return ToolResult(content=f"不是文件: {path}", is_error=True)

    # ── 文件类型检测 ──
    ext = target.suffix.lower()

    # 拦截旧格式
    if ext == ".doc":
        return ToolResult(
            content="不支持旧版 .doc 格式。请转换为 .docx 后重试。",
            is_error=True,
        )
    if ext == ".xls":
        return ToolResult(
            content="不支持旧版 .xls 格式。请转换为 .xlsx 后重试。",
            is_error=True,
        )

    # 拦截 XLSX，引导使用 file_edit_table
    if ext in (".xlsx",):
        return ToolResult(
            content="XLSX 文件请使用 file_edit_table 工具编辑。file_edit 仅支持纯文本和 DOCX 格式。",
            is_error=True,
        )

    # ── 路由到对应编辑器 ──
    file_type = _detect_file_type(target)

    if file_type == "text":
        return await _edit_text(
            target, target_content, replacement_content,
            start_line, end_line, allow_multiple,
            highlight, comment,
        )
    elif file_type == "docx":
        return await _edit_docx(
            target, target_content, replacement_content,
            allow_multiple, highlight, comment,
        )
    else:
        return ToolResult(
            content=f"不支持的文件类型: {ext or '(无扩展名)'}。file_edit 支持纯文本、DOCX 格式。编辑 XLSX 请使用 file_edit_table。",
            is_error=True,
        )

