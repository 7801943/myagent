"""file_edit 与 file_edit_table 工具入口。"""
import io
import logging
import os
from copy import copy, deepcopy
from pathlib import Path
from typing import Any

from myagent.tools.api import tool, ToolResult
from myagent.tools.builtin._file_common import (
    _XLSX_STRUCTURAL_OPS,
    _analyze_range_impact,
    _build_xlsx_diff_preview,
    _cell_bounds,
    _cell_change,
    _check_path_safety,
    _coerce_cell_value,
    _column_index,
    _compact_changes,
    _content_token,
    _detect_encoding,
    _detect_file_type,
    _hash_json,
    _is_merged_non_anchor,
    _lookup_header,
    _normalize_color,
    _parse_a1_range,
    _parse_cell_address,
    _profile_workbook,
    _range_from_bounds,
    _records_to_matrix,
    _resolve_dataset,
    _resolve_worksheet_for_table_edit,
    _verify_xlsx_persisted,
    _write_cell,
)

logger = logging.getLogger(__name__)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit — 精确搜索替换编辑
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# ── 注释风格映射 ──
_COMMENT_STYLES: dict[str, tuple[str, str]] = {
    # (prefix, suffix) — 行注释用 (prefix, "")，块注释用 (prefix, suffix)
    ".css":  ("/* ", " */"),
    ".html": ("<!-- ", " -->"),
    ".xml":  ("<!-- ", " -->"),
    ".md":   ("<!-- ", " -->"),
    ".ini":  ("; ", ""),
    ".toml": ("# ", ""),
    ".yaml": ("# ", ""),
    ".yml":  ("# ", ""),
    ".conf": ("# ", ""),
}

# JSON 等不支持注释的格式
_NO_COMMENT_EXTENSIONS = {".json", ".csv", ".tsv"}


def _get_comment_style(path: Path) -> tuple[str, str] | None:
    """根据扩展名返回注释风格 (prefix, suffix)，不支持则返回 None。"""
    ext = path.suffix.lower()
    if ext in _NO_COMMENT_EXTENSIONS:
        return None
    return _COMMENT_STYLES.get(ext, ("# ", ""))  # 默认用 # 注释


def _build_diff_preview(
    original: str,
    replaced: str,
    match_pos: int,
    match_line_count: int = 1,
) -> str:
    """构建替换位置 ±3 行的 diff 预览。

    Args:
        original: 原始文本（用于定位匹配行号）。
        replaced: 替换后的文本（用于展示内容）。
        match_pos: 匹配在 original 中的字符偏移。
        match_line_count: 匹配区域占几行（target_content 跨多行时需传入）。
    """
    # 找到匹配所在的行号
    before_text = original[:match_pos]
    match_line = before_text.count("\n")
    match_end_line = match_line + max(1, match_line_count)

    # 计算新文本中替换区域的行号（近似）
    repl_lines = replaced.split("\n")

    ctx_start = max(0, match_line - 3)
    ctx_end = min(len(repl_lines), match_end_line + 3)

    preview_lines = []
    for i in range(ctx_start, ctx_end):
        marker = ">" if match_line <= i < match_end_line else " "
        preview_lines.append(f"  {marker} {i + 1:>{len(str(ctx_end))}} | {repl_lines[i]}")

    return "\n".join(preview_lines)


def _atomic_write_text(path: Path, content: str) -> None:
    """原子写入纯文本文件。"""
    import tempfile
    tmp = tempfile.NamedTemporaryFile(
        mode="w", encoding="utf-8",
        dir=str(path.parent), suffix=".tmp", delete=False,
    )
    try:
        tmp.write(content)
        tmp.close()
        os.replace(tmp.name, str(path))
    except Exception:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass
        raise


def _atomic_write_binary(path: Path, data: bytes) -> None:
    """原子写入二进制文件。"""
    import tempfile
    tmp = tempfile.NamedTemporaryFile(
        dir=str(path.parent), suffix=".tmp", delete=False,
    )
    try:
        tmp.write(data)
        tmp.close()
        os.replace(tmp.name, str(path))
    except Exception:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass
        raise


# ── 纯文本编辑 ──

async def _edit_text(
    path: Path,
    target_content: str,
    replacement_content: str,
    start_line: int | None,
    end_line: int | None,
    allow_multiple: bool,
    highlight: str | None,
    comment: str | None,
) -> ToolResult:
    """纯文本文件的精确搜索替换。"""
    enc = _detect_encoding(path)
    if enc == "binary":
        return ToolResult(content="文件似乎是二进制格式，无法进行文本编辑。", is_error=True)

    try:
        with open(path, "r", encoding=enc, errors="replace") as f:
            original = f.read()
    except Exception as e:
        return ToolResult(content=f"读取文件失败: {e}", is_error=True)

    # ── 行号范围约束 ──
    search_text = original
    offset = 0  # 原始文本中的字符偏移
    lines = original.split("\n")
    total_lines = len(lines)
    s = 1
    e = total_lines
    if start_line is not None or end_line is not None:
        s = max(1, start_line or 1)
        e = min(end_line or total_lines, total_lines) if end_line else total_lines
        if s > total_lines:
            return ToolResult(
                content=f"文件共 {total_lines} 行，start_line={s} 超出范围。",
                is_error=True,
            )
        # 计算偏移量
        offset = sum(len(lines[i]) + 1 for i in range(s - 1))
        search_text = "\n".join(lines[s - 1:e])

    # ── 精确匹配 ──
    count = search_text.count(target_content)
    if count == 0:
        hint = ""
        if start_line is not None:
            hint = f"（在行 {start_line}-{end_line or '末尾'} 范围内）"
        return ToolResult(
            content=f"未找到目标内容{hint}。请使用 file_read 确认文件最新内容后重试。",
            is_error=True,
            metadata={"path": str(path.resolve()), "match_count": 0},
        )
    if count > 1 and not allow_multiple:
        return ToolResult(
            content=f"找到 {count} 个匹配，请提供更精确的 target_content，"
                    f"或使用 start_line/end_line 缩小范围，或设置 allow_multiple=True。",
            is_error=True,
            metadata={"path": str(path.resolve()), "match_count": count},
        )

    # ── 执行替换 ──
    # 替换次数恒等于匹配次数 count（之前用 new_search_text.count(replacement_content) 重算是错误的，
    # 因为 replacement_content 可能正好在原文中也出现，导致计数偏大；删除场景下 "".count("") 更荒谬）
    if start_line is not None or end_line is not None:
        new_search_text = search_text.replace(target_content, replacement_content)
        new_lines = lines[:s - 1] + new_search_text.split("\n") + lines[e:]
        replaced = "\n".join(new_lines)
    else:
        replaced = original.replace(target_content, replacement_content)
    actual_count = count

    # ── 批注（在替换位置上方添加注释）──
    comment_applied = False
    comment_hint = ""
    if comment:
        style = _get_comment_style(path)
        if style is None:
            comment_hint = "（该文件类型不支持注释，已忽略 comment 参数）"
        else:
            prefix, suffix = style
            comment_line = f"{prefix}[AGENT COMMENT] {comment}{suffix}\n"
            # 在第一个替换位置上方插入注释
            idx = replaced.find(replacement_content)
            if idx >= 0:
                # 找到该行的起始位置
                line_start = replaced.rfind("\n", 0, idx)
                line_start = line_start + 1 if line_start >= 0 else 0
                replaced = replaced[:line_start] + comment_line + replaced[line_start:]
                comment_applied = True

    # ── 原子写入 ──
    try:
        _atomic_write_text(path, replaced)
    except Exception as e:
        return ToolResult(content=f"写入文件失败: {e}", is_error=True)

    # ── 构建返回信息 ──
    # 真正的匹配位置：offset 是 search_text 起始偏移，需要再加上 target 在 search_text 中的相对位置。
    # 之前直接用 offset 作为 match_pos 会导致 diff 高亮行号偏上（Issue 8）。
    relative_pos = search_text.find(target_content)
    match_pos = (offset + relative_pos) if relative_pos >= 0 else original.find(target_content)
    # 匹配占多少行（target_content 跨行时高亮多行）
    match_line_count = target_content.count("\n") + 1
    diff_preview = _build_diff_preview(original, replaced, max(0, match_pos), match_line_count)

    msg = f"替换成功: {path.name}，共 {actual_count} 处替换。"
    if highlight:
        msg += "（纯文本文件不支持标色，已忽略 highlight 参数）"
    if comment_hint:
        msg += comment_hint
    elif comment_applied:
        msg += " 已添加批注注释。"

    logger.info("file_edit 纯文本成功: path=%s, replacements=%d", path, actual_count)
    return ToolResult(
        content=msg + f"\n\n预览:\n{diff_preview}",
        metadata={
            "path": str(path.resolve()),
            "format": "text",
            "replacements": actual_count,
            "highlight": highlight,
            "comment_added": comment_applied,
        },
    )


# ── DOCX 编辑 ──

def _insert_docx_run_after(anchor_run: Any, text: str, template_run: Any | None = None) -> Any:
    """Insert a run after `anchor_run`, copying character formatting from `template_run`."""
    from docx.oxml import OxmlElement
    from docx.text.run import Run

    template_run = template_run or anchor_run
    new_r = OxmlElement("w:r")
    r_pr = template_run._r.rPr
    if r_pr is not None:
        new_r.append(deepcopy(r_pr))
    anchor_run._r.addnext(new_r)

    new_run = Run(new_r, anchor_run._parent)
    new_run.text = text
    return new_run


def _normalize_docx_edit_text(
    target_content: str,
    replacement_content: str,
) -> tuple[str, str]:
    """Tolerate a single trailing newline copied from file_read's paragraph output."""
    if (
        target_content.endswith("\n")
        and replacement_content.endswith("\n")
        and "\n" not in target_content[:-1]
        and "\n" not in replacement_content[:-1]
    ):
        return target_content[:-1], replacement_content[:-1]
    return target_content, replacement_content


async def _edit_docx(
    path: Path,
    target_content: str,
    replacement_content: str,
    allow_multiple: bool,
    highlight: str | None,
    comment: str | None,
) -> ToolResult:
    """DOCX 文件的精确搜索替换，保留格式。"""
    try:
        from docx import Document
    except ImportError:
        return ToolResult(
            content="缺少 python-docx 库。请安装: pip install python-docx",
            is_error=True,
        )

    try:
        doc = Document(str(path))
    except Exception as e:
        return ToolResult(content=f"打开 DOCX 失败: {e}", is_error=True)

    target_content, replacement_content = _normalize_docx_edit_text(
        target_content,
        replacement_content,
    )

    # ── 构建虚拟文本流及 Run 映射 ──
    virtual_parts: list[str] = []
    # 每个 entry: (para_idx, run_idx, char_start, char_end)
    run_map: list[tuple[int, int, int, int]] = []
    char_pos = 0

    for pi, para in enumerate(doc.paragraphs):
        if pi > 0:
            virtual_parts.append("\n")
            char_pos += 1
        for ri, run in enumerate(para.runs):
            text = run.text
            start = char_pos
            end = char_pos + len(text)
            if text:
                run_map.append((pi, ri, start, end))
                virtual_parts.append(text)
                char_pos = end

    virtual_text = "".join(virtual_parts)

    # ── 精确匹配 ──
    count = virtual_text.count(target_content)
    if count == 0:
        return ToolResult(
            content="未找到目标内容。请使用 file_read 确认文档最新内容后重试。",
            is_error=True,
            metadata={"path": str(path.resolve()), "match_count": 0},
        )
    if count > 1 and not allow_multiple:
        return ToolResult(
            content=f"找到 {count} 个匹配，请提供更精确的 target_content 或设置 allow_multiple=True。",
            is_error=True,
            metadata={"path": str(path.resolve()), "match_count": count},
        )

    matches: list[tuple[int, int, list[tuple[int, int, int, int]]]] = []
    search_start = 0
    while search_start < len(virtual_text):
        match_idx = virtual_text.find(target_content, search_start)
        if match_idx < 0:
            break
        match_end = match_idx + len(target_content)
        involved = [
            (pi, ri, cs, ce)
            for pi, ri, cs, ce in run_map
            if cs < match_end and ce > match_idx
        ]
        if involved:
            matches.append((match_idx, match_end, involved))
        search_start = match_end
        if not allow_multiple:
            break

    cross_paragraph = [
        (match_idx, match_end)
        for match_idx, match_end, involved in matches
        if involved[0][0] != involved[-1][0]
    ]
    if cross_paragraph:
        return ToolResult(
            content=(
                "DOCX 编辑被拒绝：当前安全编辑器只支持单个段落内的替换。"
                "这次匹配跨越了多个段落，继续写入可能产生额外空行或破坏段落对齐。"
                "请把 target_content 缩小到一个段落内后重试。"
            ),
            is_error=True,
            metadata={
                "path": str(path.resolve()),
                "match_count": count,
                "cross_paragraph_matches": len(cross_paragraph),
            },
        )

    # ── 执行替换（逐个匹配处理）──
    actual_count = 0
    comment_count = 0
    for match_idx, match_end, involved in reversed(matches):
        first_pi, first_ri, first_cs, first_ce = involved[0]
        last_pi, last_ri, last_cs, last_ce = involved[-1]
        first_para = doc.paragraphs[first_pi]
        first_run = first_para.runs[first_ri]

        # 第一段第一段：保留前缀 + 替换内容
        prefix = target_content[:0]  # 空
        if first_cs < match_idx:
            prefix = first_run.text[:match_idx - first_cs]

        # 最后一段最后一段：保留后缀
        last_para = doc.paragraphs[last_pi]
        last_run = last_para.runs[last_ri]
        suffix = ""
        if last_ce > match_end:
            suffix = last_run.text[match_end - last_cs:]

        runs = first_para.runs
        first_run.text = prefix
        if first_ri != last_ri:
            for idx in range(first_ri + 1, last_ri):
                runs[idx].text = ""
            last_run.text = ""

        replacement_run = _insert_docx_run_after(first_run, replacement_content)
        if suffix:
            _insert_docx_run_after(replacement_run, suffix, last_run)

        # ── 标色 ──
        if highlight:
            try:
                from docx.enum.text import WD_COLOR_INDEX
                _DOCX_COLOR_MAP = {
                    "yellow": WD_COLOR_INDEX.YELLOW,
                    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
                    "red": WD_COLOR_INDEX.RED,
                    "pink": WD_COLOR_INDEX.PINK,
                }
                color = _DOCX_COLOR_MAP.get(highlight)
                if color:
                    replacement_run.font.highlight_color = color
            except Exception as e:
                logger.warning("DOCX 标色失败: %s", e)

        # ── 批注 ──
        if comment:
            try:
                if hasattr(doc, "add_comment"):
                    doc.add_comment(
                        replacement_run,
                        text=comment,
                        author="Agent",
                        initials="AG",
                    )
                    comment_count += 1
                else:  # pragma: no cover - compatibility with older python-docx
                    logger.warning("python-docx 版本不支持批注，需要 >= 1.2.0")
            except Exception as e:
                logger.warning("DOCX 批注添加失败: %s", e)

        actual_count += 1

    # ── 原子写入 ──
    try:
        buf = io.BytesIO()
        doc.save(buf)
        _atomic_write_binary(path, buf.getvalue())
    except Exception as e:
        return ToolResult(content=f"保存 DOCX 失败: {e}", is_error=True)

    comment_status = ""
    if comment and comment_count:
        comment_status = f" 已添加批注 {comment_count} 处。"
    elif comment:
        comment_status = " 批注添加失败，已完成文本替换。"

    logger.info("file_edit DOCX成功: path=%s, replacements=%d", path, actual_count)
    return ToolResult(
        content=f"DOCX 替换成功: {path.name}，共 {actual_count} 处替换。"
                + (" 已标色。" if highlight else "")
                + comment_status,
        metadata={
            "path": str(path.resolve()),
            "format": "docx",
            "replacements": actual_count,
            "highlight": highlight,
            "comment_added": comment_count > 0,
            "comments_added": comment_count,
        },
    )




# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit_table — XLSX 专用编辑工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _copy_row_style(ws, source_row: int, target_row: int, max_col: int) -> None:
    if source_row < 1:
        return
    for col_idx in range(1, max_col + 1):
        source = ws.cell(row=source_row, column=col_idx)
        target = ws.cell(row=target_row, column=col_idx)
        if source.has_style:
            target._style = copy(source._style)
        if source.number_format:
            target.number_format = source.number_format
        if source.protection:
            target.protection = copy(source.protection)
        if source.alignment:
            target.alignment = copy(source.alignment)


def _plan_set_range(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    cell_range = payload.get("range")
    values = payload.get("values")
    if not cell_range:
        raise ValueError("set_range payload 必须包含 range")
    if not isinstance(values, list) or not values or not all(isinstance(row, list) for row in values):
        raise ValueError("set_range payload.values 必须是非空二维数组")

    row_lengths = {len(row) for row in values}
    if len(row_lengths) != 1:
        raise ValueError("set_range payload.values 必须是矩形二维数组")
    value_rows = len(values)
    value_cols = next(iter(row_lengths)) if row_lengths else 0
    if value_cols == 0:
        raise ValueError("set_range payload.values 至少需要一列")

    min_col, min_row, max_col, max_row = _parse_a1_range(cell_range)
    resize_range = bool(payload.get("resize_range", False))
    if resize_range:
        max_row = min_row + value_rows - 1
        max_col = min_col + value_cols - 1
    else:
        range_rows = max_row - min_row + 1
        range_cols = max_col - min_col + 1
        if range_rows != value_rows or range_cols != value_cols:
            raise ValueError(
                f"range 大小({range_rows}x{range_cols})与 values 大小({value_rows}x{value_cols})不一致"
            )

    value_input = payload.get("value_input", "auto")
    kind = payload.get("kind")
    changes: list[dict[str, Any]] = []
    for r_offset, row_values in enumerate(values):
        row_idx = min_row + r_offset
        for c_offset, raw_value in enumerate(row_values):
            col_idx = min_col + c_offset
            merged = _is_merged_non_anchor(ws, row_idx, col_idx)
            if merged:
                raise ValueError(f"不能写入合并区域 {merged} 的非左上角单元格")
            cell = ws.cell(row=row_idx, column=col_idx)
            new_value, _ = _coerce_cell_value(raw_value, kind=kind, value_input=value_input)
            changes.append(_cell_change(cell, new_value))
            if apply:
                _write_cell(cell, raw_value, kind=kind, value_input=value_input)

    bounds = (min_col, min_row, max_col, max_row)
    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": changes,
        "counts": {"cells_updated": value_rows * value_cols},
        "structural": False,
    }


def _plan_update_cells(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    cells = payload.get("cells")
    if not isinstance(cells, list) or not cells:
        raise ValueError("update_cells payload.cells 必须是非空数组")

    seen: set[str] = set()
    bounds: list[tuple[int, int, int, int]] = []
    changes: list[dict[str, Any]] = []
    value_input = payload.get("value_input", "auto")

    for item in cells:
        if not isinstance(item, dict):
            raise ValueError("update_cells 的每个 cell 项必须是 object")
        address = item.get("cell")
        if not address:
            raise ValueError("update_cells 的每个 cell 项必须包含 cell")
        row_idx, col_idx = _parse_cell_address(address)
        normalized = ws.cell(row=row_idx, column=col_idx).coordinate
        if normalized in seen:
            raise ValueError(f"重复更新单元格 {normalized}")
        seen.add(normalized)
        kind = item.get("kind")
        if "value" not in item and kind != "blank":
            raise ValueError(f"单元格 {normalized} 缺少 value")
        merged = _is_merged_non_anchor(ws, row_idx, col_idx)
        if merged:
            raise ValueError(f"不能写入合并区域 {merged} 的非左上角单元格")
        raw_value = item.get("value")
        cell = ws.cell(row=row_idx, column=col_idx)
        new_value, _ = _coerce_cell_value(raw_value, kind=kind, value_input=value_input)
        changes.append(_cell_change(cell, new_value))
        if apply:
            _write_cell(cell, raw_value, kind=kind, value_input=value_input)
        bounds.append(_cell_bounds(row_idx, col_idx))

    return {
        "bounds": bounds,
        "affected_ranges": [_range_from_bounds(*b) for b in bounds],
        "changes": changes,
        "counts": {"cells_updated": len(cells)},
        "structural": False,
    }


def _plan_clear_range(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    cell_range = payload.get("range")
    if not cell_range:
        raise ValueError("clear_range payload 必须包含 range")
    clear = payload.get("clear", "values")
    if clear != "values":
        raise ValueError("当前仅支持 clear='values'")

    bounds = _parse_a1_range(cell_range)
    min_col, min_row, max_col, max_row = bounds
    changes: list[dict[str, Any]] = []
    for row_idx in range(min_row, max_row + 1):
        for col_idx in range(min_col, max_col + 1):
            merged = _is_merged_non_anchor(ws, row_idx, col_idx)
            if merged:
                raise ValueError(f"不能清空合并区域 {merged} 的非左上角单元格")
            cell = ws.cell(row=row_idx, column=col_idx)
            changes.append(_cell_change(cell, None))
            if apply:
                cell.value = None

    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": changes,
        "counts": {"cells_cleared": (max_row - min_row + 1) * (max_col - min_col + 1)},
        "structural": False,
    }


def _plan_format_range(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    from openpyxl.comments import Comment as XlComment
    from openpyxl.styles import PatternFill

    cell_range = payload.get("range")
    if not cell_range:
        raise ValueError("format_range payload 必须包含 range")
    bounds = _parse_a1_range(cell_range)
    min_col, min_row, max_col, max_row = bounds

    fill = payload.get("fill")
    font_payload = payload.get("font") or {}
    comment = payload.get("comment")
    fill_style = None
    if fill:
        color = _normalize_color(fill)
        fill_style = PatternFill(start_color=color, end_color=color, fill_type="solid")
    if font_payload and not isinstance(font_payload, dict):
        raise ValueError("format_range font 必须是 object")

    formatted = 0
    for row_idx in range(min_row, max_row + 1):
        for col_idx in range(min_col, max_col + 1):
            if _is_merged_non_anchor(ws, row_idx, col_idx):
                continue
            cell = ws.cell(row=row_idx, column=col_idx)
            if apply:
                if fill_style:
                    cell.fill = fill_style
                if font_payload:
                    new_font = copy(cell.font)
                    if "bold" in font_payload:
                        new_font.bold = bool(font_payload["bold"])
                    if "italic" in font_payload:
                        new_font.italic = bool(font_payload["italic"])
                    if "color" in font_payload:
                        new_font.color = _normalize_color(font_payload["color"])
                    cell.font = new_font
                if comment:
                    cell.comment = XlComment(str(comment), "Agent")
            formatted += 1

    format_action = {
        "range": _range_from_bounds(*bounds),
        "cells_formatted": formatted,
    }
    if fill:
        format_action["fill"] = _normalize_color(fill)
    if font_payload:
        format_action["font"] = dict(font_payload)
    if comment:
        format_action["comment"] = str(comment)

    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": [],
        "format_actions": [format_action],
        "counts": {"cells_formatted": formatted},
        "structural": False,
    }


def _plan_append_rows(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    rows = payload.get("rows")
    if not isinstance(rows, list) or not rows:
        raise ValueError("append_rows payload.rows 必须是非空数组")
    dataset = _resolve_dataset(ws, payload)
    if dataset["totals_row_count"]:
        raise ValueError("目标 Excel table 包含 totals row，当前不支持 append_rows")

    matrix = _records_to_matrix(rows, dataset["headers"])
    append_row = dataset["append_row"]
    min_col = dataset["min_col"]
    max_col = dataset["max_col"]
    max_row = append_row + len(matrix) - 1
    bounds = (min_col, append_row, max_col, max_row)

    for row_idx in range(append_row, max_row + 1):
        for col_idx in range(min_col, max_col + 1):
            if ws.cell(row=row_idx, column=col_idx).value not in (None, ""):
                raise ValueError(f"追加区域 {_range_from_bounds(*bounds)} 覆盖了非空单元格")

    changes: list[dict[str, Any]] = []
    value_input = payload.get("value_input", "auto")
    for r_offset, row_values in enumerate(matrix):
        row_idx = append_row + r_offset
        for c_offset, raw_value in enumerate(row_values):
            col_idx = min_col + c_offset
            cell = ws.cell(row=row_idx, column=col_idx)
            new_value, _ = _coerce_cell_value(raw_value, value_input=value_input)
            changes.append(_cell_change(cell, new_value))
            if apply:
                _write_cell(cell, raw_value, value_input=value_input)

    new_table_ref = None
    if dataset["table"] is not None:
        new_table_ref = _range_from_bounds(min_col, dataset["header_row"], max_col, max_row)
    if apply and dataset["table"] is not None:
        dataset["table"].ref = new_table_ref

    row_action = {
        "type": "append_rows",
        "range": _range_from_bounds(*bounds),
        "rows_appended": len(matrix),
    }
    if dataset["table_name"]:
        row_action["table_name"] = dataset["table_name"]
    if new_table_ref:
        row_action["table_ref_after"] = new_table_ref

    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": changes,
        "row_actions": [row_action],
        "counts": {"rows_appended": len(matrix)},
        "structural": True,
    }


def _dataset_key_rows(ws, dataset: dict[str, Any], key_column: str) -> tuple[int, dict[Any, int], list[Any]]:
    key_offset = _lookup_header(dataset["headers"], key_column)
    key_col = dataset["min_col"] + key_offset
    key_rows: dict[Any, int] = {}
    duplicates: list[Any] = []
    for row_idx in range(dataset["data_min_row"], dataset["data_max_row"] + 1):
        value = ws.cell(row=row_idx, column=key_col).value
        if value in (None, ""):
            continue
        if value in key_rows:
            duplicates.append(value)
        else:
            key_rows[value] = row_idx
    return key_col, key_rows, duplicates


def _plan_update_rows_by_key(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    key_column = payload.get("key_column")
    updates = payload.get("updates")
    if not key_column:
        raise ValueError("update_rows_by_key payload 必须包含 key_column")
    if not isinstance(updates, list) or not updates:
        raise ValueError("update_rows_by_key payload.updates 必须是非空数组")

    missing = payload.get("missing", "error")
    if missing not in ("error", "ignore", "append"):
        raise ValueError("missing 可选: error, ignore, append")

    dataset = _resolve_dataset(ws, payload)
    if missing == "append" and dataset["totals_row_count"]:
        raise ValueError("目标 Excel table 包含 totals row，当前不支持 missing='append'")
    _, key_rows, duplicates = _dataset_key_rows(ws, dataset, str(key_column))
    if duplicates:
        raise ValueError(f"关键列 {key_column!r} 存在重复 key: {duplicates[:5]}")

    changes: list[dict[str, Any]] = []
    bounds: list[tuple[int, int, int, int]] = []
    appended_records: list[dict[str, Any]] = []
    rows_updated = 0
    rows_ignored = 0
    value_input = payload.get("value_input", "auto")

    for update in updates:
        if not isinstance(update, dict) or "key" not in update or not isinstance(update.get("values"), dict):
            raise ValueError("updates 中每一项必须包含 key 和 values object")
        key = update["key"]
        values = update["values"]
        row_idx = key_rows.get(key)
        if row_idx is None:
            if missing == "ignore":
                rows_ignored += 1
                continue
            if missing == "append":
                record = {str(key_column): key}
                record.update(values)
                appended_records.append(record)
                continue
            raise ValueError(f"关键列 {key_column!r} 未找到 key={key!r}")

        touched_cols: list[int] = []
        for header, raw_value in values.items():
            offset = _lookup_header(dataset["headers"], str(header))
            col_idx = dataset["min_col"] + offset
            cell = ws.cell(row=row_idx, column=col_idx)
            new_value, _ = _coerce_cell_value(raw_value, value_input=value_input)
            changes.append(_cell_change(cell, new_value))
            if apply:
                _write_cell(cell, raw_value, value_input=value_input)
            touched_cols.append(col_idx)
        if touched_cols:
            bounds.append((min(touched_cols), row_idx, max(touched_cols), row_idx))
            rows_updated += 1

    rows_appended = 0
    if appended_records:
        append_payload = dict(payload)
        append_payload["rows"] = appended_records
        append_payload.pop("updates", None)
        append_plan = _plan_append_rows(ws, append_payload, apply=apply, include_changes=include_changes)
        bounds.extend(append_plan["bounds"])
        changes.extend(append_plan["changes"])
        appended_row_actions = append_plan.get("row_actions", [])
        rows_appended = append_plan["counts"].get("rows_appended", 0)
    else:
        appended_row_actions = []

    return {
        "bounds": bounds,
        "affected_ranges": [_range_from_bounds(*b) for b in bounds],
        "changes": changes,
        "row_actions": appended_row_actions,
        "counts": {
            "rows_updated": rows_updated,
            "rows_appended": rows_appended,
            "rows_ignored": rows_ignored,
        },
        "structural": bool(appended_records),
    }


def _plan_upsert_rows(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    key_column = payload.get("key_column")
    rows = payload.get("rows")
    if not key_column:
        raise ValueError("upsert_rows payload 必须包含 key_column")
    if not isinstance(rows, list) or not rows:
        raise ValueError("upsert_rows payload.rows 必须是非空数组")
    updates = []
    for row in rows:
        if not isinstance(row, dict):
            raise ValueError("upsert_rows rows 中每一项必须是 object")
        if key_column not in row:
            raise ValueError(f"upsert row 缺少 key_column {key_column!r}")
        values = dict(row)
        key = values.pop(key_column)
        updates.append({"key": key, "values": values})
    update_payload = dict(payload)
    update_payload["updates"] = updates
    update_payload["missing"] = "append"
    return _plan_update_rows_by_key(ws, update_payload, apply=apply, include_changes=include_changes)


def _plan_delete_rows(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    rows = payload.get("rows")
    if rows is None and payload.get("keys") is not None:
        key_column = payload.get("key_column")
        if not key_column:
            raise ValueError("按 key 删除时必须提供 key_column")
        keys = payload.get("keys")
        if not isinstance(keys, list) or not keys:
            raise ValueError("delete_rows payload.keys 必须是非空数组")
        dataset = _resolve_dataset(ws, payload)
        _, key_rows, duplicates = _dataset_key_rows(ws, dataset, str(key_column))
        if duplicates:
            raise ValueError(f"关键列 {key_column!r} 存在重复 key: {duplicates[:5]}")
        rows = []
        for key in keys:
            row_idx = key_rows.get(key)
            if row_idx is None:
                raise ValueError(f"关键列 {key_column!r} 未找到 key={key!r}")
            rows.append(row_idx)
    if not isinstance(rows, list) or not rows:
        raise ValueError("delete_rows payload 必须包含非空 rows 或 keys")

    row_nums = sorted({int(row) for row in rows})
    if any(row < 1 for row in row_nums):
        raise ValueError("删除行号必须从 1 开始")
    max_col = ws.max_column or 1
    bounds = [(1, row, max_col, row) for row in row_nums]
    if apply:
        for row in sorted(row_nums, reverse=True):
            ws.delete_rows(row, 1)
    return {
        "bounds": bounds,
        "affected_ranges": [_range_from_bounds(*b) for b in bounds],
        "changes": [],
        "row_actions": [{
            "type": "delete_rows",
            "range": ",".join(_range_from_bounds(*b) for b in bounds),
            "rows_deleted": len(row_nums),
            "rows": row_nums,
        }],
        "counts": {"rows_deleted": len(row_nums)},
        "structural": True,
    }


def _plan_insert_rows(ws, payload: dict[str, Any], apply: bool, include_changes: bool) -> dict[str, Any]:
    rows = payload.get("rows")
    if not isinstance(rows, list) or not rows:
        raise ValueError("insert_rows payload.rows 必须是非空数组")
    if payload.get("before_row") is not None:
        insert_at = int(payload["before_row"])
    elif payload.get("after_row") is not None:
        insert_at = int(payload["after_row"]) + 1
    else:
        raise ValueError("insert_rows payload 必须包含 before_row 或 after_row")
    if insert_at < 1:
        raise ValueError("插入行号必须从 1 开始")

    if all(isinstance(row, dict) for row in rows):
        dataset = _resolve_dataset(ws, payload)
        matrix = _records_to_matrix(rows, dataset["headers"])
    elif all(isinstance(row, list) for row in rows):
        matrix = rows
    else:
        raise ValueError("insert_rows rows 必须全部是 object 或全部是 array")

    value_input = payload.get("value_input", "auto")
    max_width = max(len(row) for row in matrix)
    bounds = (1, insert_at, max(max_width, ws.max_column or 1), insert_at + len(matrix) - 1)

    if apply:
        ws.insert_rows(insert_at, len(matrix))
        copy_style_from = payload.get("copy_style_from", "none")
        if copy_style_from not in ("above", "below", "none"):
            raise ValueError("copy_style_from 可选: above, below, none")
        if copy_style_from == "above":
            template_row = insert_at - 1
        elif copy_style_from == "below":
            template_row = insert_at + len(matrix)
        else:
            template_row = 0
        if template_row:
            for offset in range(len(matrix)):
                _copy_row_style(ws, template_row, insert_at + offset, bounds[2])
        for r_offset, row_values in enumerate(matrix):
            for c_offset, raw_value in enumerate(row_values):
                cell = ws.cell(row=insert_at + r_offset, column=1 + c_offset)
                _write_cell(cell, raw_value, value_input=value_input)

    return {
        "bounds": [bounds],
        "affected_ranges": [_range_from_bounds(*bounds)],
        "changes": [],
        "row_actions": [{
            "type": "insert_rows",
            "range": _range_from_bounds(*bounds),
            "rows_inserted": len(matrix),
            "before_row": insert_at,
        }],
        "counts": {"rows_inserted": len(matrix)},
        "structural": True,
    }


def _run_xlsx_operation(
    wb,
    ws,
    operation: str,
    payload: dict[str, Any],
    apply: bool,
    include_changes: bool,
) -> dict[str, Any]:
    if operation == "set_range":
        result = _plan_set_range(ws, payload, apply, include_changes)
    elif operation == "update_cells":
        result = _plan_update_cells(ws, payload, apply, include_changes)
    elif operation == "clear_range":
        result = _plan_clear_range(ws, payload, apply, include_changes)
    elif operation == "format_range":
        result = _plan_format_range(ws, payload, apply, include_changes)
    elif operation == "append_rows":
        result = _plan_append_rows(ws, payload, apply, include_changes)
    elif operation == "update_rows_by_key":
        result = _plan_update_rows_by_key(ws, payload, apply, include_changes)
    elif operation == "upsert_rows":
        result = _plan_upsert_rows(ws, payload, apply, include_changes)
    elif operation == "delete_rows":
        result = _plan_delete_rows(ws, payload, apply, include_changes)
    elif operation == "insert_rows":
        result = _plan_insert_rows(ws, payload, apply, include_changes)
    else:
        raise ValueError(
            "不支持的 operation: "
            f"{operation!r}。可选: set_range, update_cells, clear_range, format_range, "
            "append_rows, update_rows_by_key, upsert_rows, delete_rows, insert_rows"
        )
    result.setdefault("bounds", [])
    result.setdefault("affected_ranges", [])
    result.setdefault("changes", [])
    result.setdefault("format_actions", [])
    result.setdefault("row_actions", [])
    result.setdefault("counts", {})
    result.setdefault("structural", operation in _XLSX_STRUCTURAL_OPS)
    return result


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit_table — XLSX 结构化编辑工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@tool(name="file_edit_table",
      description=(
          "结构化编辑 XLSX 表格文件。使用 operation + payload 表达明确动作，"
          "不再支持 target_content/replacement_content 文本行匹配。\n\n"
          "支持 operation：\n"
          "- set_range: 按 A1 range 写入二维数组\n"
          "- update_cells: 按单元格地址批量更新\n"
          "- clear_range: 清空区域值并保留样式\n"
          "- format_range: 设置标色、字体和批注\n"
          "- append_rows: 向 native Excel table 或 header dataset 追加行\n"
          "- update_rows_by_key: 按关键列更新匹配行\n"
          "- upsert_rows: 按关键列更新并追加缺失行\n"
          "- delete_rows: 删除指定行或 key 匹配行\n"
          "- insert_rows: 在指定位置插入行\n\n"
          "默认 dry_run=True 只预览不保存。结构性写入需 dry_run=False 且 allow_structure_change=True。"
      ))
async def file_edit_table(
    path: str,
    operation: str,
    sheet_name: str | None = None,
    payload: dict = None,
    dry_run: bool = True,
    allow_structure_change: bool = False,
    expected_structure_token: str | None = None,
    expected_content_token: str | None = None,
    include_changes: bool = False,
) -> ToolResult:
    """
    结构化编辑 XLSX 表格文件。

    Args:
        path: 文件路径，仅支持 .xlsx。可传绝对路径、workspace 可见路径或相对路径；在会话工作区中会由工具层解析到允许的真实路径
        operation: 编辑动作。可选 set_range/update_cells/clear_range/format_range/append_rows/update_rows_by_key/upsert_rows/delete_rows/insert_rows
        sheet_name: 工作表名。单 Sheet 自动选择；table_name 可唯一定位时也可省略
        payload: operation 对应的结构化参数
        dry_run: 是否只预览不保存，默认 True
        allow_structure_change: 是否允许插入/删除/追加等结构性写入，默认 False
        expected_structure_token: 可选结构 token，不匹配时拒绝写入
        expected_content_token: 可选内容 token，不匹配时拒绝写入
        include_changes: 是否在 metadata 中返回单元格级 changes
    """
    try:
        import openpyxl
    except ImportError:
        return ToolResult(
            content="缺少 openpyxl 库。请安装: pip install openpyxl",
            is_error=True,
        )

    operation = (operation or "").strip().lower()
    payload = payload or {}
    if not isinstance(payload, dict):
        return ToolResult(
            content="payload 必须是 object。",
            is_error=True,
            metadata={"ok": False, "operation": operation, "dry_run": dry_run},
        )

    error = _check_path_safety(path)
    if error:
        return ToolResult(content=error, is_error=True)

    target = Path(path)
    if not target.exists():
        return ToolResult(content=f"文件不存在: {path}", is_error=True)
    if not target.is_file():
        return ToolResult(content=f"不是文件: {path}", is_error=True)

    ext = target.suffix.lower()
    if ext == ".xls":
        return ToolResult(
            content="不支持旧版 .xls 格式。请转换为 .xlsx 后重试。",
            is_error=True,
        )
    if ext != ".xlsx":
        return ToolResult(
            content=f"file_edit_table 仅支持 .xlsx 文件，当前文件类型: {ext or '(无扩展名)'}。",
            is_error=True,
        )

    logger.info("file_edit_table 开始: path=%s, operation=%s, dry_run=%s", path, operation, dry_run)

    try:
        wb = openpyxl.load_workbook(str(path), data_only=False)
    except Exception as e:
        return ToolResult(content=f"打开 XLSX 失败: {e}", is_error=True)

    try:
        ws, sheet_error = _resolve_worksheet_for_table_edit(wb, sheet_name, payload)
        if sheet_error:
            wb.close()
            return ToolResult(
                content=sheet_error,
                is_error=True,
                metadata={"ok": False, "operation": operation, "dry_run": dry_run},
            )

        previous_profile = _profile_workbook(wb)
        previous_structure_token = _hash_json(previous_profile)

        # 先 dry-plan 一次：计算影响范围、校验 payload，不写入 workbook。
        plan = _run_xlsx_operation(
            wb, ws, operation, payload, apply=False, include_changes=include_changes
        )
        bounds = plan.get("bounds", [])
        previous_content_token = _hash_json([_content_token(ws, bound) for bound in bounds]) if bounds else _content_token(ws)

        if expected_structure_token and expected_structure_token != previous_structure_token:
            wb.close()
            return ToolResult(
                content="结构 token 不匹配，文件结构可能已变化，请重新 file_read 后再编辑。",
                is_error=True,
                metadata={
                    "ok": False,
                    "operation": operation,
                    "dry_run": dry_run,
                    "expected_structure_token": expected_structure_token,
                    "current_structure_token": previous_structure_token,
                },
            )
        if expected_content_token and expected_content_token != previous_content_token:
            wb.close()
            return ToolResult(
                content="内容 token 不匹配，目标区域可能已变化，请重新 file_read 后再编辑。",
                is_error=True,
                metadata={
                    "ok": False,
                    "operation": operation,
                    "dry_run": dry_run,
                    "expected_content_token": expected_content_token,
                    "current_content_token": previous_content_token,
                },
            )

        is_structural = bool(plan.get("structural")) or operation in _XLSX_STRUCTURAL_OPS
        if is_structural and not dry_run and not allow_structure_change:
            wb.close()
            return ToolResult(
                content=(
                    f"{operation} 属于结构性操作。请先 dry_run 预览，确认后设置 "
                    "dry_run=False 且 allow_structure_change=True。"
                ),
                is_error=True,
                metadata={
                    "ok": False,
                    "operation": operation,
                    "dry_run": dry_run,
                    "data": {
                        "sheet": ws.title,
                        "affected_ranges": plan.get("affected_ranges", []),
                        "requires_structure_change": True,
                    },
                },
            )

        impact = _analyze_range_impact(wb, ws, bounds, operation)
        warnings = sorted(set(plan.get("warnings", []) + impact.get("warnings", [])))
        if is_structural and dry_run and not allow_structure_change:
            warnings.append("该操作属于结构性变更；实际写入时需要 allow_structure_change=True")

        if impact.get("errors"):
            wb.close()
            message = "操作被阻止：" + "; ".join(impact["errors"][:5])
            return ToolResult(
                content=message,
                is_error=True,
                metadata={
                    "ok": False,
                    "operation": operation,
                    "dry_run": dry_run,
                    "error": {"type": "ImpactError", "message": message},
                    "data": {
                        "path": str(target.resolve()),
                        "sheet": ws.title,
                        "affected_ranges": plan.get("affected_ranges", []),
                        "warnings": warnings,
                        "impact": impact,
                    },
                },
            )

        if dry_run:
            diff_preview = _build_xlsx_diff_preview(plan)
            wb.close()
            message = (
                f"XLSX 表格编辑预览: {target.name}，operation={operation}，"
                f"将影响 {len(plan.get('affected_ranges', []))} 个区域。\n"
                f"{diff_preview}"
            )
            return ToolResult(
                content=message,
                metadata={
                    "ok": True,
                    "operation": operation,
                    "dry_run": True,
                    "data": {
                        "path": str(target.resolve()),
                        "sheet": ws.title,
                        "affected_ranges": plan.get("affected_ranges", []),
                        "counts": plan.get("counts", {}),
                        "warnings": warnings,
                        "diff_preview": diff_preview,
                        "changes": plan.get("changes", []) if include_changes else _compact_changes(plan.get("changes", [])),
                        "change_count": len(plan.get("changes", [])),
                        "format_actions": plan.get("format_actions", []),
                        "row_actions": plan.get("row_actions", []),
                        "requires_structure_change": is_structural,
                        "previous_structure_token": previous_structure_token,
                        "previous_content_token": previous_content_token,
                        "impact": impact,
                    },
                },
            )

        applied = _run_xlsx_operation(
            wb, ws, operation, payload, apply=True, include_changes=include_changes
        )
        new_profile = _profile_workbook(wb)
        new_structure_token = _hash_json(new_profile)
        new_content_token = _hash_json([_content_token(ws, bound) for bound in applied.get("bounds", [])]) if applied.get("bounds") else _content_token(ws)

        sheet_title = ws.title
        try:
            buf = io.BytesIO()
            wb.save(buf)
            wb.close()
            _atomic_write_binary(target, buf.getvalue())
        except Exception as e:
            return ToolResult(content=f"保存 XLSX 失败: {e}", is_error=True)

        persistence = _verify_xlsx_persisted(target, sheet_title, applied)
        diff_preview = _build_xlsx_diff_preview(applied)
        counts = applied.get("counts", {})
        count_summary = ", ".join(f"{k}={v}" for k, v in counts.items() if v) or "已完成"
        persisted_text = "落盘校验通过" if persistence.get("persisted") else "落盘校验失败"
        message = (
            f"XLSX 表格编辑成功: {target.name}，operation={operation}，{count_summary}。"
            f" {persisted_text}。\n{diff_preview}"
        )
        if warnings:
            message += f"\n警告: {'; '.join(warnings[:5])}"
        if not persistence.get("persisted"):
            message += "\n校验问题: " + "; ".join(persistence.get("mismatches", [])[:5])

        logger.info("file_edit_table 成功: path=%s, operation=%s", target, operation)
        return ToolResult(
            content=message,
            metadata={
                "ok": True,
                "operation": operation,
                "dry_run": False,
                "data": {
                    "path": str(target.resolve()),
                    "sheet": sheet_title,
                    "affected_ranges": applied.get("affected_ranges", []),
                    "counts": counts,
                    "warnings": warnings,
                    "diff_preview": diff_preview,
                    "persistence_verification": persistence,
                    "changes": applied.get("changes", []) if include_changes else _compact_changes(applied.get("changes", [])),
                    "change_count": len(applied.get("changes", [])),
                    "format_actions": applied.get("format_actions", []),
                    "row_actions": applied.get("row_actions", []),
                    "previous_structure_token": previous_structure_token,
                    "new_structure_token": new_structure_token,
                    "previous_content_token": previous_content_token,
                    "new_content_token": new_content_token,
                },
            },
        )
    except Exception as e:
        try:
            wb.close()
        except Exception:
            pass
        return ToolResult(
            content=f"XLSX 表格编辑失败: {type(e).__name__}: {e}",
            is_error=True,
            metadata={"ok": False, "operation": operation, "dry_run": dry_run},
        )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit — 纯文本 + DOCX 编辑
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@tool(name="file_edit",
      description=(
          "精确编辑已有文件。支持纯文本文件和 DOCX 两种类型。\n"
          "通过 target_content 精确匹配原始内容，替换为 replacement_content。\n"
          "支持对修改位置进行标色（highlight）和添加批注（comment）。\n\n"
          "使用建议：\n"
          "- 修改已有文件优先使用 file_edit，而非 file_write\n"
          "- target_content 必须精确匹配原文（含缩进和空白）\n"
          "- 匹配失败时先 file_read 确认最新内容再重试\n"
          "- 不支持 .doc 旧格式\n"
          "- 编辑 XLSX 请使用 file_edit_table"
      ))
async def file_edit(
    path: str,
    target_content: str,
    replacement_content: str,
    start_line: int | None = None,
    end_line: int | None = None,
    allow_multiple: bool = False,
    highlight: str | None = None,
    comment: str | None = None,
) -> ToolResult:
    """
    精确编辑已有文件（纯文本 / DOCX）。

    Args:
        path: 文件路径，支持文本/DOCX。可传绝对路径、workspace 可见路径或相对路径；在会话工作区中会由工具层解析到允许的真实路径
        target_content: 要被替换的原始文本（精确匹配，必填）
        replacement_content: 替换后的新文本（必填）
        start_line: 纯文本文件搜索范围起始行（1-based）
        end_line: 纯文本文件搜索范围结束行
        allow_multiple: 是否允许替换多个匹配（默认 False）
        highlight: 标色颜色: "yellow"/"green"/"red"/"pink"/None
        comment: 批注内容，None=不添加批注
    """
    # ── highlight 值校验 ──
    if highlight is not None and highlight not in ("yellow", "green", "red", "pink"):
        return ToolResult(
            content=f"不支持的 highlight 颜色: {highlight!r}。可选: yellow, green, red, pink",
            is_error=True,
        )

    logger.info("file_edit 开始: path=%s", path)

    # ── 文件存在性检查 ──
    target = Path(path)
    if not target.exists():
        return ToolResult(content=f"文件不存在: {path}", is_error=True)
    if not target.is_file():
        return ToolResult(content=f"不是文件: {path}", is_error=True)

    # ── 文件类型检测 ──
    ext = target.suffix.lower()

    # 拦截旧格式
    if ext == ".doc":
        return ToolResult(
            content="不支持旧版 .doc 格式。请转换为 .docx 后重试。",
            is_error=True,
        )
    if ext == ".xls":
        return ToolResult(
            content="不支持旧版 .xls 格式。请转换为 .xlsx 后重试。",
            is_error=True,
        )

    # 拦截 XLSX，引导使用 file_edit_table
    if ext in (".xlsx",):
        return ToolResult(
            content="XLSX 文件请使用 file_edit_table 工具编辑。file_edit 仅支持纯文本和 DOCX 格式。",
            is_error=True,
        )

    # ── 路由到对应编辑器 ──
    file_type = _detect_file_type(target)

    if file_type == "text":
        return await _edit_text(
            target, target_content, replacement_content,
            start_line, end_line, allow_multiple,
            highlight, comment,
        )
    elif file_type == "docx":
        return await _edit_docx(
            target, target_content, replacement_content,
            allow_multiple, highlight, comment,
        )
    else:
        return ToolResult(
            content=f"不支持的文件类型: {ext or '(无扩展名)'}。file_edit 支持纯文本、DOCX 格式。编辑 XLSX 请使用 file_edit_table。",
            is_error=True,
        )
