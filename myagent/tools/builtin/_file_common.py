"""
文件工具共享实现。

包含路径安全、文件类型检测、读取解析器、DOCX/XLSX 公共辅助函数。
"""
import base64
import csv
import hashlib
import io
import json
import logging
import mimetypes
import os
import re
from copy import copy, deepcopy
from datetime import date, datetime
from pathlib import Path
from typing import Any

from myagent.tools.api import ToolResult

logger = logging.getLogger(__name__)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 安全策略
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_DENIED_PATHS = {
    "/etc", "/root",
    "/sys", "/proc/sys", "/boot", "/dev",
}

_MAX_BASE64_BYTES = 20 * 1024 * 1024  # 20MB base64 输出上限
_PDF_DEFAULT_MAX_PAGES = 3            # 未指定页码范围时，默认最多渲染页数
_PDF_SCAN_FALLBACK_RATIO = 0.5        # 无文本页占比超过此阈值时，自动回退到 base64 渲染


def _check_path_safety(path: str) -> str | None:
    """检查路径是否被安全策略禁止。返回错误信息或 None。"""
    try:
        resolved = str(Path(path).resolve())
    except Exception:
        return f"路径解析失败: {path}"
    for denied_path in _DENIED_PATHS:
        if resolved.startswith(denied_path):
            return f"路径被安全策略禁止: {path} (匹配: {denied_path})"
    return None


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 文件类型检测
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_TEXT_EXTENSIONS = {
    ".txt", ".md", ".json", ".yaml", ".yml", ".toml",
    ".xml", ".html", ".css",
    ".ini", ".cfg", ".conf", ".env", ".gitignore",
    ".log", ".csv", ".tsv",
}

_SPREADSHEET_EXTENSIONS = {".xlsx", ".xls"}
_DOC_EXTENSIONS = {".docx", ".doc"}
_PDF_EXTENSIONS = {".pdf"}

_IMAGE_EXTENSIONS = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".svg": "image/svg+xml",
    ".ico": "image/x-icon",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
}


def _detect_file_type(path: Path) -> str:
    """
    根据扩展名检测文件类型。
    返回: "text" | "csv" | "xlsx" | "docx" | "pdf" | "image" | "binary"
    """
    ext = path.suffix.lower()

    if ext == ".csv":
        return "csv"
    if ext == ".tsv":
        return "csv"
    if ext in _SPREADSHEET_EXTENSIONS:
        return "xlsx"
    if ext in _DOC_EXTENSIONS:
        return "docx"
    if ext in _PDF_EXTENSIONS:
        return "pdf"
    if ext in _IMAGE_EXTENSIONS:
        return "image"
    if ext in _TEXT_EXTENSIONS:
        return "text"

    # 无扩展名或未知扩展名：尝试检测是否为文本
    if not ext:
        try:
            mime, _ = mimetypes.guess_type(str(path))
            if mime and mime.startswith("text/"):
                return "text"
        except Exception:
            pass

    return "binary"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 编码检测
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _detect_encoding(path: Path, hint: str | None = None) -> str:
    """检测文件编码。hint 优先，然后尝试 chardet，最后回退 utf-8。"""
    if hint:
        return hint

    # 尝试读取前 8KB 用于检测
    try:
        raw = path.open("rb").read(8192)
    except Exception:
        return "utf-8"

    # 检测 BOM
    if raw.startswith(b"\xef\xbb\xbf"):
        return "utf-8-sig"
    if raw.startswith(b"\xff\xfe"):
        return "utf-16-le"
    if raw.startswith(b"\xfe\xff"):
        return "utf-16-be"

    # 检测 null 字节（二进制文件）
    if b"\x00" in raw[:1024]:
        return "binary"

    # 尝试 utf-8
    try:
        raw.decode("utf-8")
        return "utf-8"
    except UnicodeDecodeError:
        pass

    # chardet 检测
    try:
        import chardet
        result = chardet.detect(raw)
        if result and result.get("encoding"):
            return result["encoding"]
    except ImportError:
        pass

    # 回退 latin-1（永远不会失败）
    return "latin-1"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 文本行号格式化
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _format_lines(
    lines: list[str],
    start_line: int,
    end_line: int | None,
    total_lines: int | None = None,
) -> tuple[str, dict[str, Any]]:
    """
    格式化行列表，始终添加行号。
    返回 (格式化后的文本, 元数据)。
    """
    display_total = total_lines or len(lines)
    width = len(str(display_total))

    output_parts: list[str] = []
    actual_count = 0

    for i, line in enumerate(lines):
        line_num = start_line + i
        if end_line is not None and line_num > end_line:
            break

        # 确保行以换行结尾
        if not line.endswith("\n"):
            line += "\n"

        output_parts.append(f"{line_num:>{width}} | {line}")
        actual_count += 1

    content = "".join(output_parts)

    meta = {
        "lines_output": actual_count,
        "start_line": start_line,
        "total_lines": total_lines,
    }

    return content, meta


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 各格式解析器
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

async def _read_text(
    path: Path,
    start_line: int | None,
    end_line: int | None,
    encoding: str | None,
) -> ToolResult:
    """读取纯文本文件。"""
    enc = _detect_encoding(path, encoding)
    if enc == "binary":
        return ToolResult(
            content=f"文件 {path.name} 似乎是二进制文件，请使用 output_format=\"base64\" 读取。",
            is_error=True,
        )

    try:
        with open(path, "r", encoding=enc, errors="replace") as f:
            all_lines = f.readlines()
    except PermissionError:
        return ToolResult(content=f"权限不足，无法读取: {path}", is_error=True)
    except Exception as e:
        return ToolResult(content=f"读取文件失败: {e}", is_error=True)

    total = len(all_lines)
    s = max(1, start_line or 1)
    e = min(end_line or total, total) if end_line else total

    if s > total:
        return ToolResult(
            content=f"文件共 {total} 行，start_line={s} 超出范围。",
            is_error=True,
            metadata={"total_lines": total},
        )

    sliced = all_lines[s - 1:e]
    content, meta = _format_lines(sliced, s, e, total)
    meta["path"] = str(path.resolve())
    meta["encoding"] = enc

    logger.info("file_read 纯文本完成: %s, 返回行数=%d", path, meta["lines_output"])
    return ToolResult(content=content, metadata=meta)


async def _read_csv(
    path: Path,
    start_line: int | None,
    end_line: int | None,
    encoding: str | None,
) -> ToolResult:
    """读取 CSV/TSV 文件。使用 csv 模块解析，按行输出。"""
    enc = _detect_encoding(path, encoding)
    if enc == "binary":
        enc = "utf-8"

    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","

    try:
        with open(path, "r", encoding=enc, errors="replace", newline="") as f:
            # 使用 csv.reader 处理各种引号和转义
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
    except PermissionError:
        return ToolResult(content=f"权限不足，无法读取: {path}", is_error=True)
    except Exception as e:
        return ToolResult(content=f"读取 CSV 失败: {e}", is_error=True)

    if not rows:
        return ToolResult(content="CSV 文件为空。", metadata={"path": str(path.resolve())})

    total = len(rows)
    s = max(1, start_line or 1)
    e = min(end_line or total, total) if end_line else total

    if s > total:
        return ToolResult(
            content=f"文件共 {total} 行，start_line={s} 超出范围。",
            is_error=True,
            metadata={"total_lines": total},
        )

    # 格式化每行为 | 分隔的列
    formatted_lines: list[str] = []
    for row in rows[s - 1:e]:
        formatted_lines.append(" | ".join(str(cell) for cell in row) + "\n")

    content, meta = _format_lines(formatted_lines, s, e, total)
    meta["path"] = str(path.resolve())
    meta["encoding"] = enc
    meta["format"] = "csv"

    logger.info("file_read CSV完成: %s, 返回行数=%d", path, meta["lines_output"])
    return ToolResult(content=content, metadata=meta)


async def _read_xlsx(
    path: Path,
    sheet_name: str | None,
    start_line: int | None,
    end_line: int | None,
    xlsx_range: str | None = None,
    render_mode: str = "values",
    row_mode: str = "text",
    include_tables: bool = True,
    include_merges: bool = True,
) -> ToolResult:
    """读取 XLSX/XLS 文件。多 Sheet 时先列出 Sheet 列表。"""
    try:
        import openpyxl
    except ImportError:
        return ToolResult(
            content="缺少 openpyxl 库，无法解析 XLSX。请安装: pip install openpyxl",
            is_error=True,
        )

    render_mode = (render_mode or "values").lower()
    row_mode = (row_mode or "text").lower()
    if render_mode not in ("values", "formulas", "both"):
        return ToolResult(
            content=f"不支持的 render_mode: {render_mode}。可选: values, formulas, both",
            is_error=True,
        )
    if row_mode not in ("text", "arrays", "objects"):
        return ToolResult(
            content=f"不支持的 row_mode: {row_mode}。可选: text, arrays, objects",
            is_error=True,
        )

    try:
        if render_mode == "both":
            wb = openpyxl.load_workbook(str(path), read_only=False, data_only=False)
            value_wb = openpyxl.load_workbook(str(path), read_only=False, data_only=True)
        else:
            wb = openpyxl.load_workbook(
                str(path), read_only=False, data_only=(render_mode == "values")
            )
            value_wb = None
    except Exception as e:
        return ToolResult(content=f"打开 XLSX 失败: {e}", is_error=True)

    try:
        sheet_names = wb.sheetnames

        # ── 未指定 sheet_name：列出所有工作表 ──
        if sheet_name is None:
            if len(sheet_names) == 1:
                ws = wb[sheet_names[0]]
                value_ws = value_wb[sheet_names[0]] if value_wb else None
                return await _format_xlsx_sheet(
                    ws, sheet_names[0], wb, path, start_line, end_line,
                    xlsx_range=xlsx_range, render_mode=render_mode,
                    row_mode=row_mode, value_ws=value_ws,
                    include_tables=include_tables, include_merges=include_merges,
                )

            # 多个 Sheet：列出信息
            lines = [f"文件: {path.name} (XLSX 工作簿)\n"]
            lines.append(f"包含 {len(sheet_names)} 个工作表：\n")
            for i, name in enumerate(sheet_names, 1):
                ws = wb[name]
                row_count = ws.max_row or 0
                col_count = ws.max_column or 0
                table_count = len(_worksheet_tables(ws))
                table_hint = f", {table_count} 个表格" if table_count else ""
                lines.append(f"  [{i}] \"{name}\" ({row_count} 行 × {col_count} 列{table_hint})\n")
            lines.append(
                "\n请指定 sheet_name 参数选择要读取的工作表。\n"
                f"示例: file_read(path=\"{path}\", sheet_name=\"{sheet_names[0]}\")\n"
            )

            profile = _profile_workbook(wb)
            wb.close()
            if value_wb:
                value_wb.close()
            return ToolResult(
                content="".join(lines),
                metadata={
                    "path": str(path.resolve()),
                    "format": "xlsx",
                    "sheet_count": len(sheet_names),
                    "sheet_names": sheet_names,
                    "workbook": profile,
                    "structure_token": _hash_json(profile),
                },
            )

        # ── 指定了 sheet_name ──
        if sheet_name not in sheet_names:
            wb.close()
            if value_wb:
                value_wb.close()
            return ToolResult(
                content=(
                    f"工作表 \"{sheet_name}\" 不存在。\n"
                    f"可用工作表: {', '.join(repr(n) for n in sheet_names)}"
                ),
                is_error=True,
                metadata={"sheet_names": sheet_names},
            )

        ws = wb[sheet_name]
        value_ws = value_wb[sheet_name] if value_wb else None
        return await _format_xlsx_sheet(
            ws, sheet_name, wb, path, start_line, end_line,
            xlsx_range=xlsx_range, render_mode=render_mode,
            row_mode=row_mode, value_ws=value_ws,
            include_tables=include_tables, include_merges=include_merges,
        )
    except Exception as e:
        wb.close()
        if 'value_wb' in locals() and value_wb:
            value_wb.close()
        return ToolResult(content=f"读取 XLSX 失败: {e}", is_error=True)


async def _format_xlsx_sheet(
    ws, sheet_name: str, wb, path: Path,
    start_line: int | None, end_line: int | None,
    xlsx_range: str | None = None,
    render_mode: str = "values",
    row_mode: str = "text",
    value_ws=None,
    include_tables: bool = True,
    include_merges: bool = True,
) -> ToolResult:
    """格式化一个 XLSX 工作表的内容。"""
    total = ws.max_row or 0
    col_count = ws.max_column or 0

    if total == 0:
        metadata = {"path": str(path.resolve()), "format": "xlsx", "sheet": sheet_name}
        metadata["workbook"] = _profile_workbook(wb)
        metadata["structure_token"] = _hash_json(metadata["workbook"])
        wb.close()
        if value_ws is not None:
            value_ws.parent.close()
        return ToolResult(content=f"工作表 \"{sheet_name}\" 为空。", metadata=metadata)

    if xlsx_range:
        min_col, min_row, max_col, max_row = _parse_a1_range(xlsx_range)
        s, e = min_row, max_row
    else:
        s = max(1, start_line or 1)
        e = min(end_line or total, total) if end_line else total
        min_col, min_row, max_col, max_row = 1, s, col_count, e

    if s > total:
        wb.close()
        if value_ws is not None:
            value_ws.parent.close()
        return ToolResult(
            content=f"工作表共 {total} 行，start_line={s} 超出范围。",
            is_error=True,
            metadata={"total_lines": total, "sheet": sheet_name},
        )

    display_lines: list[str] = []
    rows: list[list[Any]] = []
    for row_idx in range(min_row, max_row + 1):
        display_cells: list[str] = []
        row_values: list[Any] = []
        for col_idx in range(min_col, max_col + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            value_cell = value_ws.cell(row=row_idx, column=col_idx) if value_ws else None
            payload = _xlsx_cell_payload(cell, value_cell, render_mode)
            row_values.append(payload)
            display_cells.append(_xlsx_cell_display(payload))
        rows.append(row_values)
        display_lines.append(" | ".join(display_cells) + "\n")

    content, meta = _format_lines(display_lines, min_row, max_row, total)
    selected_range = _range_from_bounds(min_col, min_row, max_col, max_row)
    profile = _profile_workbook(wb)
    sheet_profile = next(
        (sheet for sheet in profile.get("sheets", []) if sheet.get("name") == sheet_name),
        {},
    )

    metadata_rows: Any = rows
    headers: list[str] | None = None
    if row_mode == "objects" and rows:
        headers = [str(_metadata_scalar(v) or "") for v in rows[0]]
        metadata_rows = []
        for row in rows[1:]:
            metadata_rows.append({headers[i]: _metadata_scalar(row[i]) for i in range(len(headers)) if headers[i]})
    elif row_mode == "arrays":
        metadata_rows = [[_metadata_scalar(v) for v in row] for row in rows]

    meta.update({
        "path": str(path.resolve()),
        "format": "xlsx",
        "sheet": sheet_name,
        "columns": col_count,
        "range": selected_range,
        "render_mode": render_mode,
        "row_mode": row_mode,
        "rows": metadata_rows,
        "workbook": profile,
        "structure_token": _hash_json(profile),
        "content_token": _content_token(ws, (min_col, min_row, max_col, max_row)),
    })
    if headers is not None:
        meta["headers"] = headers
    if include_tables:
        meta["tables"] = sheet_profile.get("tables", [])
    if include_merges:
        meta["merged_ranges"] = sheet_profile.get("merged_ranges", [])

    wb.close()
    if value_ws is not None:
        value_ws.parent.close()

    logger.info("file_read XLSX完成: %s, sheet=%s, 返回行数=%d", path, sheet_name, meta["lines_output"])
    return ToolResult(content=content, metadata=meta)


async def _read_pdf_text(
    path: Path,
    start_line: int | None,
    end_line: int | None,
) -> ToolResult:
    """使用 pdfplumber 提取 PDF 文本内容。"""
    # 抑制 pdfminer 的 DEBUG 日志洪水。
    # pdfminer 在解析 PDF 时会为每个 token/操作输出大量 DEBUG 日志
    # （nexttoken / add_results / nextobject / do_keyword / exec 等），
    # 一个普通 PDF 可能产生数千~数万条日志，导致：
    #   1. 子进程 stderr 被海量日志填满，产生严重 I/O 瓶颈
    #   2. 主进程 _drain_stderr 逐行读取并以 INFO 级别转发，进一步加剧延迟
    #   3. PDF 解析看似"卡住"，实际大部分时间花在日志 I/O 上
    # 将 pdfminer 相关 logger 提升到 WARNING 级别，彻底消除底层噪音。
    for _logger_name in ("pdfminer", "pdfminer.psparser", "pdfminer.pdfinterp",
                         "pdfminer.pdfpage", "pdfminer.converter",
                         "pdfminer.layout", "pdfminer.utils"):
        logging.getLogger(_logger_name).setLevel(logging.WARNING)

    try:
        import pdfplumber
    except ImportError:
        return ToolResult(
            content="缺少 pdfplumber 库，无法解析 PDF。请安装: pip install pdfplumber",
            is_error=True,
        )

    try:
        all_lines: list[str] = []
        page_count = 0

        with pdfplumber.open(str(path)) as pdf:
            page_count = len(pdf.pages)
            start_page = max(1, start_line or 1)
            end_page = min(end_line or page_count, page_count) if end_line else page_count

            if start_page > page_count:
                return ToolResult(
                    content=f"PDF 共 {page_count} 页，start_line={start_page} 超出范围。",
                    is_error=True,
                    metadata={"page_count": page_count},
                )

            for page_num in range(start_page, end_page + 1):
                page = pdf.pages[page_num - 1]
                text = page.extract_text() or ""
                # 添加页码标记
                all_lines.append(f"--- 第 {page_num} 页 ---\n")
                if text.strip():
                    for line in text.split("\n"):
                        all_lines.append(line + "\n")
                else:
                    all_lines.append("[此页无可提取文本，可能是扫描图片]\n")

        if not all_lines:
            return ToolResult(
                content="PDF 未能提取到任何文本。可能是扫描件/图片 PDF，"
                        "请使用 output_format=\"base64\" 以图片形式读取。",
                metadata={"page_count": page_count},
            )

        # ── 扫描页检测：如果大部分页无文本，自动回退到 base64 渲染 ──
        pages_read_count = end_page - start_page + 1
        empty_page_count = sum(
            1 for line in all_lines
            if line.strip() == "[此页无可提取文本，可能是扫描图片]"
        )
        # 每个 empty page 产生 2 行（页码标题 + 提示），实际空页数 = empty_page_count
        if (
            pages_read_count > 0
            and empty_page_count / pages_read_count > _PDF_SCAN_FALLBACK_RATIO
        ):
            logger.debug(
                "PDF 扫描页回退: 读取 %d 页中 %d 页无文本(%.0f%% > 阈值 %.0f%%)，回退到 base64 渲染",
                pages_read_count, empty_page_count,
                empty_page_count / pages_read_count * 100,
                _PDF_SCAN_FALLBACK_RATIO * 100,
            )
            return await _read_pdf_base64(path, start_line, end_line)

        # 把所有行视为一个整体，用页码范围代替行号范围
        total = len(all_lines)
        s = 1  # 已经只取了目标页的内容

        content, meta = _format_lines(all_lines, s, None, total)
        meta["path"] = str(path.resolve())
        meta["format"] = "pdf"
        meta["page_count"] = page_count
        meta["pages_read"] = f"{start_page}-{end_page}"

        logger.info("file_read PDF文本完成: %s, 读取页数=%d", path, end_page - start_page + 1)
        return ToolResult(content=content, metadata=meta)

    except Exception as e:
        return ToolResult(content=f"解析 PDF 失败: {e}", is_error=True)


async def _read_pdf_base64(
    path: Path,
    start_line: int | None,
    end_line: int | None,
) -> ToolResult:
    """使用 PyMuPDF 将 PDF 页渲染为 JPEG 图片并返回 base64。"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ToolResult(
            content="缺少 PyMuPDF 库，无法渲染 PDF。请安装: pip install PyMuPDF",
            is_error=True,
        )

    try:
        doc = fitz.open(str(path))
        page_count = len(doc)

        # 未指定范围时，默认只渲染前几页，避免生成过大响应
        if start_line is None and end_line is None:
            start_page = 1
            end_page = min(_PDF_DEFAULT_MAX_PAGES, page_count)
        else:
            start_page = max(1, start_line or 1)
            end_page = min(end_line or page_count, page_count) if end_line else page_count

        if start_page > page_count:
            doc.close()
            return ToolResult(
                content=f"PDF 共 {page_count} 页，start_line={start_page} 超出范围。",
                is_error=True,
                metadata={"page_count": page_count},
            )

        content_blocks: list[dict[str, Any]] = []
        total_bytes = 0

        for page_num in range(start_page, end_page + 1):
            page = doc[page_num - 1]
            # 渲染为 150 DPI 的 JPEG
            pix = page.get_pixmap(dpi=150)
            logger.debug(
                "PDF 第 %d/%d 页渲染: %dx%d (dpi=150, quality=75)",
                page_num, page_count, pix.width, pix.height,
            )
            img_bytes = pix.tobytes("jpeg", jpg_quality=75)
            b64_data = base64.b64encode(img_bytes).decode("ascii")
            total_bytes += len(img_bytes)

            content_blocks.append({
                "type": "image_base64",
                "data": b64_data,
                "media_type": "image/jpeg",
                "page": page_num,
            })

            if total_bytes > _MAX_BASE64_BYTES:
                # 超过大小限制，截断
                remaining = end_page - page_num
                desc = (
                    f"PDF 共 {page_count} 页，已渲染第 {start_page}-{page_num} 页为图片。"
                    f"（剩余 {remaining} 页因大小限制未输出，"
                    f"可调整 start_line/end_line 读取其他页）"
                )
                doc.close()
                logger.info("file_read PDF图片完成(截断): %s, 渲染页数=%d", path, page_num - start_page + 1)
                return ToolResult(
                    content=desc,
                    content_blocks=content_blocks,
                    metadata={
                        "path": str(path.resolve()),
                        "format": "pdf",
                        "page_count": page_count,
                        "pages_rendered": list(range(start_page, page_num + 1)),
                        "truncated": True,
                    },
                )

        doc.close()

        pages_hint = ""
        if page_count > end_page:
            pages_hint = (
                f"（仅渲染了前 {end_page - start_page + 1} 页，"
                f"共 {page_count} 页。可指定 start_line/end_line 读取更多页）"
            )
        desc = (
            f"PDF 共 {page_count} 页，已渲染第 {start_page}-{end_page} 页为图片"
            f"（共 {end_page - start_page + 1} 张）。{pages_hint}"
        )
        logger.info("file_read PDF图片完成: %s, 渲染页数=%d", path, end_page - start_page + 1)
        return ToolResult(
            content=desc,
            content_blocks=content_blocks,
            metadata={
                "path": str(path.resolve()),
                "format": "pdf",
                "page_count": page_count,
                "pages_rendered": list(range(start_page, end_page + 1)),
            },
        )

    except Exception as e:
        return ToolResult(content=f"渲染 PDF 失败: {e}", is_error=True)


async def _read_docx(
    path: Path,
    start_line: int | None,
    end_line: int | None,
) -> ToolResult:
    """读取 DOCX 文件。"""
    ext = path.suffix.lower()

    if ext == ".doc":
        # .doc 格式：python-docx 有限支持，提示用户
        return ToolResult(
            content=(
                "不支持旧版 .doc 格式。请将文件转换为 .docx 格式后重试。\n"
                "转换方法: 使用 LibreOffice 命令 `libreoffice --headless --convert-to docx <file.doc>`"
            ),
            is_error=True,
        )

    try:
        from docx import Document
    except ImportError:
        return ToolResult(
            content="缺少 python-docx 库，无法解析 DOCX。请安装: pip install python-docx",
            is_error=True,
        )

    try:
        doc = Document(str(path))
    except Exception as e:
        return ToolResult(content=f"打开 DOCX 失败: {e}", is_error=True)

    lines = _extract_docx_lines(doc)

    total = len(lines)
    if total == 0:
        return ToolResult(
            content="DOCX 文件内容为空。",
            metadata={"path": str(path.resolve()), "format": "docx"},
        )

    s = max(1, start_line or 1)
    e = min(end_line or total, total) if end_line else total

    if s > total:
        return ToolResult(
            content=f"文档共 {total} 行，start_line={s} 超出范围。",
            is_error=True,
            metadata={"total_lines": total},
        )

    sliced = lines[s - 1:e]
    content, meta = _format_lines(sliced, s, e, total)
    meta["path"] = str(path.resolve())
    meta["format"] = "docx"

    logger.info("file_read DOCX完成: %s, 返回行数=%d", path, meta["lines_output"])
    return ToolResult(content=content, metadata=meta)


def _iter_docx_body_blocks(doc: Any):
    """按 DOCX body XML 顺序产出段落和表格。"""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _format_docx_table_cell(text: str) -> str:
    """把单元格内多段文本压成 Markdown 表格单元格。"""
    compact = " / ".join(part.strip() for part in text.splitlines() if part.strip())
    return compact.replace("|", r"\|")


def _extract_docx_lines(doc: Any) -> list[str]:
    """按文档内容流提取 DOCX 段落和表格文本。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    lines: list[str] = []
    table_index = 0

    for block in _iter_docx_body_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text
            lines.append(text + "\n" if text.strip() else "\n")
            continue

        if isinstance(block, Table):
            table_index += 1
            lines.append(f"[表格 {table_index}]\n")
            for row in block.rows:
                cells = [_format_docx_table_cell(cell.text) for cell in row.cells]
                row_text = " | ".join(cells)
                lines.append(f"| {row_text} |\n")
            lines.append(f"[/表格 {table_index}]\n")

    return lines


# 大图片自动缩放阈值（超过此尺寸则按比例缩小）
_IMAGE_RESIZE_THRESHOLD = 2048  # 像素（长边）
_IMAGE_JPEG_QUALITY = 80


def _resize_image_if_needed(
    data: bytes, media_type: str
) -> tuple[bytes, str, int, int, bool]:
    """
    如果图片尺寸过大，按比例缩小以减少传输体积。

    Returns:
        (处理后的bytes, media_type, width, height, was_resized)
    """
    try:
        from PIL import Image
    except ImportError:
        return data, media_type, 0, 0, False

    try:
        img = Image.open(io.BytesIO(data))
        orig_width, orig_height = img.size
        was_resized = False

        max_dim = max(orig_width, orig_height)
        if max_dim > _IMAGE_RESIZE_THRESHOLD:
            scale = _IMAGE_RESIZE_THRESHOLD / max_dim
            new_width = int(orig_width * scale)
            new_height = int(orig_height * scale)
            logger.debug(
                "图像缩放: %dx%d -> %dx%d (阈值=%d, 缩放比=%.3f)",
                orig_width, orig_height, new_width, new_height,
                _IMAGE_RESIZE_THRESHOLD, scale,
            )
            img = img.resize((new_width, new_height), Image.LANCZOS)
            was_resized = True
        else:
            logger.debug("图像无需缩放: %dx%d (长边=%d, 阈值=%d)",
                         orig_width, orig_height, max_dim, _IMAGE_RESIZE_THRESHOLD)

        buf = io.BytesIO()
        if media_type == "image/jpeg" or was_resized:
            if img.mode in ("RGBA", "LA", "P"):
                img = img.convert("RGB")
            img.save(buf, format="JPEG", quality=_IMAGE_JPEG_QUALITY)
            output_media_type = "image/jpeg"
        else:
            img.save(buf, format="PNG")
            output_media_type = "image/png"

        final_width, final_height = img.size
        return buf.getvalue(), output_media_type, final_width, final_height, was_resized

    except Exception as e:
        logger.warning("图像处理异常，返回原始数据: %s", e)
        return data, media_type, 0, 0, False


async def _read_image_base64(path: Path) -> ToolResult:
    """读取图片文件并返回 base64 编码。大图自动缩放以减少传输体积。"""
    ext = path.suffix.lower()
    media_type = _IMAGE_EXTENSIONS.get(ext)

    if not media_type:
        mime, _ = mimetypes.guess_type(str(path))
        media_type = mime or "application/octet-stream"

    try:
        file_size = path.stat().st_size
        data = path.read_bytes()
    except PermissionError:
        return ToolResult(content=f"权限不足，无法读取: {path}", is_error=True)
    except Exception as e:
        return ToolResult(content=f"读取图片失败: {e}", is_error=True)

    # 自动缩放处理
    processed_data, final_media_type, width, height, was_resized = (
        _resize_image_if_needed(data, media_type)
    )

    # 最终大小检查
    if len(processed_data) > _MAX_BASE64_BYTES:
        return ToolResult(
            content=f"图片文件过大 ({file_size / 1024 / 1024:.1f}MB，"
                    f"缩放后仍超过 {_MAX_BASE64_BYTES / 1024 / 1024:.0f}MB 限制)。",
            is_error=True,
            metadata={
                "path": str(path.resolve()),
                "original_size": file_size,
                "dimensions": f"{width}x{height}" if width else "unknown",
            },
        )

    b64_data = base64.b64encode(processed_data).decode("ascii")

    size_info = f"{file_size} 字节"
    if was_resized:
        size_info += f"（已自动缩放，缩放后 {len(processed_data)} 字节）"
    dim_info = f"，尺寸 {width}x{height}" if width else ""

    logger.info("file_read 图片完成: %s, 尺寸=%s", path, f"{width}x{height}" if width else "unknown")
    return ToolResult(
        content=f"图片: {path.name} ({final_media_type}, {size_info}{dim_info})",
        content_blocks=[{
            "type": "image_base64",
            "data": b64_data,
            "media_type": final_media_type,
        }],
        metadata={
            "path": str(path.resolve()),
            "format": "image",
            "media_type": final_media_type,
            "original_size": file_size,
            "encoded_size": len(processed_data),
            "was_resized": was_resized,
            **({"width": width, "height": height} if width else {}),
        },
    )


async def _read_binary_base64(path: Path) -> ToolResult:
    """将任意二进制文件以 base64 编码返回。"""
    try:
        file_size = path.stat().st_size
        if file_size > _MAX_BASE64_BYTES:
            return ToolResult(
                content=f"文件过大 ({file_size / 1024 / 1024:.1f}MB > "
                        f"{_MAX_BASE64_BYTES / 1024 / 1024:.0f}MB 限制)。",
                is_error=True,
            )

        data = path.read_bytes()
    except PermissionError:
        return ToolResult(content=f"权限不足，无法读取: {path}", is_error=True)
    except Exception as e:
        return ToolResult(content=f"读取文件失败: {e}", is_error=True)

    b64_data = base64.b64encode(data).decode("ascii")

    mime, _ = mimetypes.guess_type(str(path))
    media_type = mime or "application/octet-stream"

    logger.info("file_read 二进制完成: %s, size=%d", path, file_size)
    return ToolResult(
        content=f"二进制文件: {path.name} ({media_type}, {file_size} 字节)",
        content_blocks=[{
            "type": "image_base64",
            "data": b64_data,
            "media_type": media_type,
        }],
        metadata={
            "path": str(path.resolve()),
            "format": "binary",
            "media_type": media_type,
            "file_size": file_size,
        },
    )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 主入口：file_read
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# XLSX 结构化读写辅助函数
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

_XLSX_STRUCTURAL_OPS = {"append_rows", "upsert_rows", "delete_rows", "insert_rows"}
_XLSX_ROW_STRUCTURE_OPS = {"delete_rows", "insert_rows"}
_XLSX_VALUE_INPUTS = {"auto", "raw", "formula"}
_XLSX_CELL_KINDS = {"text", "number", "bool", "date", "formula", "blank"}


def _hash_json(data: Any) -> str:
    """生成短 hash，用于 workbook structure/content token。"""
    raw = json.dumps(data, ensure_ascii=False, sort_keys=True, default=str)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]


def _json_value(value: Any) -> Any:
    """把 openpyxl 值转换成 JSON metadata 友好的值。"""
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    return str(value)


def _metadata_scalar(value: Any) -> Any:
    if isinstance(value, dict):
        if "formula" in value:
            return value
        return value.get("value")
    return _json_value(value)


def _xlsx_cell_payload(cell, value_cell=None, render_mode: str = "values") -> Any:
    """返回单元格的结构化读取值。"""
    if render_mode == "both":
        formula_value = cell.value
        calculated = value_cell.value if value_cell is not None else None
        if isinstance(formula_value, str) and formula_value.startswith("="):
            return {"formula": formula_value, "value": _json_value(calculated)}
        return _json_value(calculated if value_cell is not None else formula_value)
    return _json_value(cell.value)


def _xlsx_cell_display(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, dict):
        formula = value.get("formula")
        calculated = value.get("value")
        if formula and calculated not in (None, ""):
            return f"{formula} => {calculated}"
        if formula:
            return str(formula)
        return str(value.get("value") or "")
    return str(value)


def _parse_a1_range(cell_range: str) -> tuple[int, int, int, int]:
    """解析 A1 range，返回 min_col, min_row, max_col, max_row。"""
    from openpyxl.utils.cell import range_boundaries

    if not isinstance(cell_range, str) or not cell_range.strip():
        raise ValueError("range 必须是非空 A1 字符串")
    text = cell_range.strip().replace("$", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    if not re.match(r"^[A-Za-z]+[1-9][0-9]*(?::[A-Za-z]+[1-9][0-9]*)?$", text):
        raise ValueError(f"不合法的 A1 range: {cell_range!r}")
    min_col, min_row, max_col, max_row = range_boundaries(text)
    if min_col > max_col or min_row > max_row:
        raise ValueError(f"不合法的 A1 range: {cell_range!r}")
    return min_col, min_row, max_col, max_row


def _parse_cell_address(address: str) -> tuple[int, int]:
    """解析单元格地址，返回 row, col。"""
    from openpyxl.utils.cell import coordinate_to_tuple

    if not isinstance(address, str) or not address.strip():
        raise ValueError("cell 必须是非空 A1 单元格地址")
    text = address.strip().replace("$", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    if not re.match(r"^[A-Za-z]+[1-9][0-9]*$", text):
        raise ValueError(f"不合法的单元格地址: {address!r}")
    return coordinate_to_tuple(text)


def _column_index(value: str | int) -> int:
    from openpyxl.utils.cell import column_index_from_string

    if isinstance(value, int):
        if value < 1:
            raise ValueError("列索引必须从 1 开始")
        return value
    if isinstance(value, str) and value.strip().isdigit():
        num = int(value.strip())
        if num < 1:
            raise ValueError("列索引必须从 1 开始")
        return num
    if isinstance(value, str):
        return column_index_from_string(value.strip())
    raise ValueError(f"不合法的列标识: {value!r}")


def _range_from_bounds(min_col: int, min_row: int, max_col: int, max_row: int) -> str:
    from openpyxl.utils.cell import get_column_letter

    start = f"{get_column_letter(min_col)}{min_row}"
    end = f"{get_column_letter(max_col)}{max_row}"
    return start if start == end else f"{start}:{end}"


def _bounds_overlap(a: tuple[int, int, int, int], b: tuple[int, int, int, int]) -> bool:
    amin_col, amin_row, amax_col, amax_row = a
    bmin_col, bmin_row, bmax_col, bmax_row = b
    return not (amax_col < bmin_col or bmax_col < amin_col or amax_row < bmin_row or bmax_row < amin_row)


def _bounds_contains(outer: tuple[int, int, int, int], inner: tuple[int, int, int, int]) -> bool:
    omin_col, omin_row, omax_col, omax_row = outer
    imin_col, imin_row, imax_col, imax_row = inner
    return omin_col <= imin_col and omin_row <= imin_row and omax_col >= imax_col and omax_row >= imax_row


def _merged_range_bounds(merged_range) -> tuple[int, int, int, int]:
    return merged_range.min_col, merged_range.min_row, merged_range.max_col, merged_range.max_row


def _is_merged_non_anchor(ws, row: int, col: int) -> str | None:
    for merged in ws.merged_cells.ranges:
        bounds = _merged_range_bounds(merged)
        if bounds[0] <= col <= bounds[2] and bounds[1] <= row <= bounds[3]:
            if row != merged.min_row or col != merged.min_col:
                return str(merged)
    return None


def _worksheet_tables(ws) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    try:
        names = list(ws.tables)
    except Exception:
        names = []
    for name in names:
        try:
            table = ws.tables[name]
            ref = getattr(table, "ref", table if isinstance(table, str) else None)
            if not ref:
                continue
            tables.append({
                "name": getattr(table, "name", name),
                "display_name": getattr(table, "displayName", name),
                "ref": ref,
                "totals_row_count": int(getattr(table, "totalsRowCount", 0) or 0),
            })
        except Exception:
            continue
    return tables


def _table_by_name(ws, table_name: str):
    if not table_name:
        return None
    try:
        if table_name in ws.tables:
            return ws.tables[table_name]
    except Exception:
        pass
    for table in _worksheet_tables(ws):
        if table.get("name") == table_name or table.get("display_name") == table_name:
            try:
                return ws.tables[table.get("name")]
            except Exception:
                return None
    return None


def _profile_workbook(wb) -> dict[str, Any]:
    sheets: list[dict[str, Any]] = []
    for ws in wb.worksheets:
        sheet_info: dict[str, Any] = {
            "name": ws.title,
            "max_row": ws.max_row or 0,
            "max_column": ws.max_column or 0,
            "dimension": ws.calculate_dimension(),
            "tables": _worksheet_tables(ws),
            "merged_ranges": [str(rng) for rng in ws.merged_cells.ranges],
            "charts": len(getattr(ws, "_charts", []) or []),
        }
        try:
            sheet_info["data_validations"] = len(ws.data_validations.dataValidation)
        except Exception:
            sheet_info["data_validations"] = 0
        try:
            sheet_info["conditional_formatting"] = len(ws.conditional_formatting)
        except Exception:
            sheet_info["conditional_formatting"] = 0
        sheets.append(sheet_info)

    defined_names: list[str] = []
    try:
        for defined_name in wb.defined_names:
            name = getattr(defined_name, "name", str(defined_name))
            defined_names.append(name)
    except Exception:
        pass

    return {
        "sheet_names": list(wb.sheetnames),
        "sheets": sheets,
        "defined_names": sorted(defined_names),
    }


def _content_token(ws, bounds: tuple[int, int, int, int] | None = None) -> str:
    if bounds is None:
        bounds = (1, 1, ws.max_column or 1, ws.max_row or 1)
    min_col, min_row, max_col, max_row = bounds
    rows: list[list[Any]] = []
    for row_idx in range(min_row, max_row + 1):
        row_values: list[Any] = []
        for col_idx in range(min_col, max_col + 1):
            row_values.append(_json_value(ws.cell(row=row_idx, column=col_idx).value))
        rows.append(row_values)
    return _hash_json({"sheet": ws.title, "range": _range_from_bounds(*bounds), "rows": rows})


def _resolve_worksheet_for_table_edit(wb, sheet_name: str | None, payload: dict[str, Any]) -> tuple[Any | None, str | None]:
    table_name = payload.get("table_name") if isinstance(payload, dict) else None

    if sheet_name is not None:
        if sheet_name not in wb.sheetnames:
            return None, (
                f"工作表 \"{sheet_name}\" 不存在。可用工作表: "
                f"{', '.join(repr(n) for n in wb.sheetnames)}"
            )
        ws = wb[sheet_name]
        if table_name and _table_by_name(ws, table_name) is None:
            return None, f"工作表 \"{sheet_name}\" 中不存在表格 {table_name!r}。"
        return ws, None

    if table_name:
        matches = [ws for ws in wb.worksheets if _table_by_name(ws, table_name) is not None]
        if len(matches) == 1:
            return matches[0], None
        if not matches:
            return None, f"工作簿中不存在表格 {table_name!r}。"
        return None, f"表格 {table_name!r} 在多个工作表中存在，请指定 sheet_name。"

    if len(wb.sheetnames) == 1:
        return wb[wb.sheetnames[0]], None
    return None, (
        f"文件包含 {len(wb.sheetnames)} 个工作表，必须指定 sheet_name。"
        f"可用工作表: {', '.join(repr(n) for n in wb.sheetnames)}"
    )


def _coerce_cell_value(value: Any, kind: str | None = None, value_input: str = "auto") -> tuple[Any, str | None]:
    """将 payload 值转换为 openpyxl 可写值，返回 (value, number_format)。"""
    value_input = (value_input or "auto").lower()
    if value_input not in _XLSX_VALUE_INPUTS:
        raise ValueError(f"不支持的 value_input: {value_input!r}。可选: auto, raw, formula")
    if kind is not None:
        kind = str(kind).lower()
        if kind not in _XLSX_CELL_KINDS:
            raise ValueError(f"不支持的 kind: {kind!r}。可选: {', '.join(sorted(_XLSX_CELL_KINDS))}")

    if kind == "blank":
        return None, None
    if value is None:
        return None, None

    if kind == "text":
        return str(value), None
    if kind == "formula" or value_input == "formula":
        if not isinstance(value, str) or not value.startswith("="):
            raise ValueError("formula 值必须是以 '=' 开头的字符串")
        return value, None
    if kind == "bool":
        if isinstance(value, bool):
            return value, None
        if isinstance(value, str) and value.strip().lower() in ("true", "false"):
            return value.strip().lower() == "true", None
        raise ValueError(f"无法将 {value!r} 转为 bool")
    if kind == "number":
        if isinstance(value, bool):
            raise ValueError("bool 不能作为 number 写入")
        if isinstance(value, (int, float)):
            return value, None
        if isinstance(value, str) and re.match(r"^[+-]?\d+(?:\.\d+)?$", value.strip()):
            text = value.strip()
            return (float(text), None) if "." in text else (int(text), None)
        raise ValueError(f"无法将 {value!r} 转为 number")
    if kind == "date":
        if isinstance(value, datetime):
            return value, "yyyy-mm-dd h:mm:ss"
        if isinstance(value, date):
            return value, "yyyy-mm-dd"
        if isinstance(value, str):
            text = value.strip()
            try:
                if "T" in text or " " in text:
                    return datetime.fromisoformat(text), "yyyy-mm-dd h:mm:ss"
                return date.fromisoformat(text), "yyyy-mm-dd"
            except ValueError as exc:
                raise ValueError(f"无法将 {value!r} 转为 date") from exc
        raise ValueError(f"无法将 {value!r} 转为 date")

    if value_input == "raw":
        if isinstance(value, (str, int, float, bool, datetime, date)):
            return value, None
        raise ValueError(f"raw 模式不支持写入 {type(value).__name__}")

    if isinstance(value, (int, float, bool, datetime, date)):
        return value, "yyyy-mm-dd" if isinstance(value, date) and not isinstance(value, datetime) else None
    if not isinstance(value, str):
        raise ValueError(f"不支持写入 {type(value).__name__} 类型")

    text = value.strip()
    if text == "":
        return None, None
    if text.startswith("="):
        return text, None
    lowered = text.lower()
    if lowered in ("true", "false"):
        return lowered == "true", None

    # 保守推断：前导零数字串、长数字串和混合字符串保持文本。
    if re.match(r"^[+-]?(?:0|[1-9]\d{0,14})$", text):
        signless = text[1:] if text[0] in "+-" else text
        if len(signless) > 1 and signless.startswith("0"):
            return value, None
        return int(text), None
    if re.match(r"^[+-]?(?:0|[1-9]\d*)\.\d+$", text):
        signless = text[1:] if text[0] in "+-" else text
        if len(signless.split(".", 1)[0]) > 1 and signless.startswith("0"):
            return value, None
        return float(text), None
    return value, None


def _cell_change(cell, new_value: Any) -> dict[str, Any]:
    return {
        "cell": cell.coordinate,
        "old_value": _json_value(cell.value),
        "new_value": _json_value(new_value),
    }



def _format_diff_value(value: Any) -> str:
    if value is None:
        return "<blank>"
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return repr(value)


def _compact_changes(changes: list[dict[str, Any]], limit: int = 50) -> list[dict[str, Any]]:
    """返回较小的 changes 列表，避免 metadata 过大。"""
    return changes[:limit]


def _build_xlsx_diff_preview(result: dict[str, Any], limit: int = 20) -> str:
    """构建可读 diff，用于判断改动位置和值是否符合预期。"""
    lines: list[str] = []
    changes = result.get("changes", []) or []
    value_changes = [
        change for change in changes
        if change.get("old_value") != change.get("new_value")
    ]

    if value_changes:
        lines.append("值变更 diff:")
        for change in value_changes[:limit]:
            old_text = _format_diff_value(change.get("old_value"))
            new_text = _format_diff_value(change.get("new_value"))
            lines.append(f"  - {change.get('cell')}: {old_text} -> {new_text}")
        if len(value_changes) > limit:
            lines.append(f"  ... 还有 {len(value_changes) - limit} 个单元格变更未展示")
    else:
        lines.append("值变更 diff: 无单元格值变化")

    row_actions = result.get("row_actions", []) or []
    if row_actions:
        lines.append("结构变更:")
        for action in row_actions[:limit]:
            details = [f"type={action.get('type')}", f"range={action.get('range')}"]
            for key in ("rows_appended", "rows_inserted", "rows_deleted"):
                if action.get(key) is not None:
                    details.append(f"{key}={action[key]}")
            if action.get("table_ref_after"):
                details.append(f"table_ref_after={action['table_ref_after']}")
            lines.append("  - " + ", ".join(details))

    format_actions = result.get("format_actions", []) or []
    if format_actions:
        lines.append("格式/批注结果:")
        for action in format_actions[:limit]:
            details = [f"range={action.get('range')}", f"cells={action.get('cells_formatted')}"]
            if action.get("fill"):
                details.append(f"fill=#{action['fill']}")
            if action.get("font"):
                details.append(f"font={action['font']}")
            if action.get("comment"):
                details.append(f"comment={_format_diff_value(action['comment'])}")
            lines.append("  - " + ", ".join(details))

    return "\n".join(lines)


def _cell_value_matches(actual: Any, expected: Any) -> bool:
    return _json_value(actual) == expected


def _color_matches(color_obj: Any, expected_rgb: str) -> bool:
    actual = getattr(color_obj, "rgb", None)
    if not actual:
        return False
    return str(actual).upper().endswith(expected_rgb.upper())


def _verify_xlsx_persisted(path: Path, sheet_name: str, result: dict[str, Any]) -> dict[str, Any]:
    """保存后重新打开工作簿，核对值、标色和批注是否实际落盘。"""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=False, data_only=False)
        ws = wb[sheet_name]
        checked = 0
        mismatches: list[str] = []

        for change in result.get("changes", []) or []:
            cell_ref = change.get("cell")
            if not cell_ref:
                continue
            checked += 1
            actual = ws[cell_ref].value
            expected = change.get("new_value")
            if not _cell_value_matches(actual, expected):
                mismatches.append(
                    f"{cell_ref}: expected {_format_diff_value(expected)}, got {_format_diff_value(_json_value(actual))}"
                )

        format_checked = 0
        for action in result.get("format_actions", []) or []:
            bounds = _parse_a1_range(action["range"])
            min_col, min_row, max_col, max_row = bounds
            for row_idx in range(min_row, max_row + 1):
                for col_idx in range(min_col, max_col + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if _is_merged_non_anchor(ws, row_idx, col_idx):
                        continue
                    format_checked += 1
                    if action.get("fill") and not _color_matches(cell.fill.fgColor, action["fill"]):
                        mismatches.append(f"{cell.coordinate}: fill 未落盘为 #{action['fill']}")
                    if action.get("comment"):
                        if cell.comment is None or cell.comment.text != str(action["comment"]):
                            mismatches.append(f"{cell.coordinate}: 批注未落盘或内容不一致")
                    font_payload = action.get("font") or {}
                    if "bold" in font_payload and bool(cell.font.bold) != bool(font_payload["bold"]):
                        mismatches.append(f"{cell.coordinate}: bold 字体状态不一致")
                    if "italic" in font_payload and bool(cell.font.italic) != bool(font_payload["italic"]):
                        mismatches.append(f"{cell.coordinate}: italic 字体状态不一致")
                    if "color" in font_payload and not _color_matches(cell.font.color, _normalize_color(font_payload["color"])):
                        mismatches.append(f"{cell.coordinate}: font color 未落盘为 #{_normalize_color(font_payload['color'])}")

        wb.close()
        return {
            "persisted": not mismatches,
            "cells_checked": checked,
            "format_cells_checked": format_checked,
            "mismatches": mismatches[:20],
            "mismatch_count": len(mismatches),
        }
    except Exception as exc:
        return {
            "persisted": False,
            "cells_checked": 0,
            "format_cells_checked": 0,
            "mismatches": [f"落盘校验失败: {type(exc).__name__}: {exc}"],
            "mismatch_count": 1,
        }


def _write_cell(cell, value: Any, kind: str | None = None, value_input: str = "auto") -> dict[str, Any]:
    new_value, number_format = _coerce_cell_value(value, kind=kind, value_input=value_input)
    change = _cell_change(cell, new_value)
    cell.value = new_value
    if number_format:
        cell.number_format = number_format
    return change


def _normalize_color(value: str) -> str:
    if not isinstance(value, str):
        raise ValueError("颜色必须是字符串")
    text = value.strip().lstrip("#").upper()
    if not re.match(r"^[0-9A-F]{6}$", text):
        raise ValueError(f"不合法的颜色: {value!r}，需要 RRGGBB 或 #RRGGBB")
    return text


def _find_merged_conflicts(ws, bounds: tuple[int, int, int, int]) -> list[str]:
    conflicts: list[str] = []
    for merged in ws.merged_cells.ranges:
        merged_bounds = _merged_range_bounds(merged)
        if _bounds_overlap(bounds, merged_bounds) and not _bounds_contains(bounds, merged_bounds):
            conflicts.append(f"目标区域 {_range_from_bounds(*bounds)} 与合并区域 {merged} 部分重叠")
    return conflicts


def _formula_cells(ws) -> list[str]:
    cells: list[str] = []
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row or 1, min_col=1, max_col=ws.max_column or 1):
        for cell in row:
            if isinstance(cell.value, str) and cell.value.startswith("="):
                cells.append(f"{ws.title}!{cell.coordinate}")
    return cells


def _analyze_range_impact(wb, ws, bounds_list: list[tuple[int, int, int, int]], operation: str) -> dict[str, Any]:
    """对目标区域做保守影响分析。"""
    warnings: list[str] = []
    errors: list[str] = []
    impacted_tables: list[str] = []
    affected_ranges = [_range_from_bounds(*bounds) for bounds in bounds_list]

    table_infos = _worksheet_tables(ws)
    for bounds in bounds_list:
        errors.extend(_find_merged_conflicts(ws, bounds))
        for table_info in table_infos:
            table_bounds = _parse_a1_range(table_info["ref"])
            if _bounds_overlap(bounds, table_bounds):
                impacted_tables.append(table_info["name"])
                if operation in _XLSX_ROW_STRUCTURE_OPS:
                    errors.append(
                        f"{operation} 会影响 native Excel table {table_info['name']} ({table_info['ref']})"
                    )
                elif operation not in {"append_rows", "update_rows_by_key", "upsert_rows"}:
                    warnings.append(
                        f"目标区域与 native Excel table {table_info['name']} ({table_info['ref']}) 重叠"
                    )

    if operation in _XLSX_ROW_STRUCTURE_OPS:
        formulas = _formula_cells(ws)
        if formulas:
            preview = ", ".join(formulas[:5])
            more = "..." if len(formulas) > 5 else ""
            errors.append(f"结构性行操作所在工作表包含公式单元格: {preview}{more}")
        if getattr(ws, "_charts", None):
            errors.append("结构性行操作所在工作表包含图表，当前编辑器不会自动修复图表数据源")
        try:
            if ws.data_validations and ws.data_validations.dataValidation:
                warnings.append("工作表包含数据验证，结构性行操作可能影响验证区域")
        except Exception:
            pass
        try:
            if len(ws.conditional_formatting):
                warnings.append("工作表包含条件格式，结构性行操作可能影响条件格式区域")
        except Exception:
            pass

    return {
        "affected_ranges": affected_ranges,
        "warnings": sorted(set(warnings)),
        "errors": sorted(set(errors)),
        "impacted_tables": sorted(set(impacted_tables)),
    }


def _table_headers(ws, table) -> list[str]:
    min_col, min_row, max_col, _ = _parse_a1_range(table.ref)
    headers: list[str] = []
    for col_idx in range(min_col, max_col + 1):
        value = ws.cell(row=min_row, column=col_idx).value
        headers.append(str(value).strip() if value is not None else "")
    return headers


def _validate_headers(headers: list[str]) -> str | None:
    seen: set[str] = set()
    for header in headers:
        if not header:
            return "表头包含空列名，无法按字段名定位。"
        if header in seen:
            return f"表头 {header!r} 重复，无法按字段名定位。"
        seen.add(header)
    return None


def _lookup_header(headers: list[str], key: str) -> int:
    if key in headers:
        return headers.index(key)
    lowered = {header.lower(): idx for idx, header in enumerate(headers)}
    idx = lowered.get(str(key).lower())
    if idx is not None:
        return idx
    raise ValueError(f"列 {key!r} 不存在。可用列: {headers}")


def _resolve_dataset(ws, payload: dict[str, Any]) -> dict[str, Any]:
    table_name = payload.get("table_name")
    if table_name:
        table = _table_by_name(ws, table_name)
        if table is None:
            raise ValueError(f"工作表 {ws.title!r} 中不存在表格 {table_name!r}")
        min_col, min_row, max_col, max_row = _parse_a1_range(table.ref)
        totals = int(getattr(table, "totalsRowCount", 0) or 0)
        headers = _table_headers(ws, table)
        header_error = _validate_headers(headers)
        if header_error:
            raise ValueError(header_error)
        return {
            "kind": "table",
            "table": table,
            "table_name": table_name,
            "headers": headers,
            "min_col": min_col,
            "max_col": max_col,
            "header_row": min_row,
            "data_min_row": min_row + 1,
            "data_max_row": max_row - totals,
            "append_row": max_row + 1,
            "totals_row_count": totals,
            "ref_bounds": (min_col, min_row, max_col, max_row),
        }

    header_row = int(payload.get("header_row", 1))
    start_col = _column_index(payload.get("start_col", "A"))
    end_col = _column_index(payload["end_col"]) if payload.get("end_col") else (ws.max_column or start_col)
    headers = []
    for col_idx in range(start_col, end_col + 1):
        value = ws.cell(row=header_row, column=col_idx).value
        headers.append(str(value).strip() if value is not None else "")
    header_error = _validate_headers(headers)
    if header_error:
        raise ValueError(header_error)

    data_min_row = header_row + 1
    data_max_row = max(ws.max_row or header_row, data_min_row - 1)
    append_row = data_max_row + 1
    return {
        "kind": "worksheet",
        "table": None,
        "table_name": None,
        "headers": headers,
        "min_col": start_col,
        "max_col": end_col,
        "header_row": header_row,
        "data_min_row": data_min_row,
        "data_max_row": data_max_row,
        "append_row": append_row,
        "totals_row_count": 0,
        "ref_bounds": (start_col, header_row, end_col, data_max_row),
    }


def _records_to_matrix(rows: list[Any], headers: list[str]) -> list[list[Any]]:
    matrix: list[list[Any]] = []
    for row in rows:
        if isinstance(row, dict):
            matrix.append([row.get(header) for header in headers])
        elif isinstance(row, list):
            if len(row) > len(headers):
                raise ValueError(f"行列数({len(row)})超过表头列数({len(headers)})")
            matrix.append(row + [None] * (len(headers) - len(row)))
        else:
            raise ValueError("rows 中的每一行必须是 object 或 array")
    return matrix


def _cell_bounds(row: int, col: int) -> tuple[int, int, int, int]:
    return col, row, col, row


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# file_edit_table — XLSX 专用编辑工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
