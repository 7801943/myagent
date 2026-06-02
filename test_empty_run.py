from docx import Document
from docx.oxml import OxmlElement
from docx.text.run import Run
import copy

def _insert_docx_run_after(anchor_run, text, template_run=None):
    template_run = template_run or anchor_run
    new_r = OxmlElement("w:r")
    r_pr = template_run._r.rPr
    if r_pr is not None:
        new_r.append(copy.deepcopy(r_pr))
    anchor_run._r.addnext(new_r)
    new_run = Run(new_r, anchor_run._parent)
    new_run.text = text
    return new_run

doc = Document()
p = doc.add_paragraph("Hello ")
run1 = p.add_run("world")

run2 = _insert_docx_run_after(run1, "")
try:
    doc.add_comment(run2, text="This is a comment", author="Agent", initials="AG")
    doc.save("test2.docx")
    print("Success")
except Exception as e:
    print(f"Error: {e}")
